# -*- coding: utf-8 -*-
"""
Generisanje akademskog rada: Agenti zasnovani na Velikim jezičkim modelima (LLM-based Agents) i planiranje zadataka
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUTPUT_DIR, "images")
os.makedirs(IMG_DIR, exist_ok=True)

# ============================================================
# DIAGRAM GENERATION
# ============================================================

def create_multi_agent_architecture_diagram():
    """Creates the Planner-Manager-Worker-Tester architecture diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(8, 8))
    ax.set_xlim(-1, 11)
    ax.set_ylim(-0.5, 11)
    ax.axis('off')
    ax.set_aspect('equal')

    # Layer positions (center_x, center_y, width, height, label, facecolor)
    layers = {
        'user':    (5.0, 10.0, 2.2, 0.9, 'Korisnik', '#E8EAF6'),
        'planner': (5.0, 8.0, 2.6, 0.9, 'Planer Agent', '#C8E6C9'),
        'manager': (5.0, 6.0, 2.6, 0.9, 'Menadžer Agent', '#BBDEFB'),
        'w1':      (1.5, 3.5, 2.2, 0.9, 'Radnik 1', '#FFF9C4'),
        'w2':      (5.0, 3.5, 2.2, 0.9, 'Radnik 2', '#FFF9C4'),
        'wn':      (8.5, 3.5, 2.2, 0.9, 'Radnik N', '#FFF9C4'),
        'tester':  (5.0, 1.0, 2.6, 0.9, 'Tester Agent', '#FFCCBC'),
    }

    # Draw boxes
    for key, (cx, cy, w, h, label, fc) in layers.items():
        rect = FancyBboxPatch((cx - w/2, cy - h/2), w, h, boxstyle="round,pad=0.1",
                              edgecolor='#37474F', facecolor=fc, linewidth=1.8)
        ax.add_patch(rect)
        ax.text(cx, cy, label, ha='center', va='center',
                fontsize=10, fontweight='bold', color='#212121')

    # Helper to get edge point of a box
    def box_edge(key, side):
        cx, cy, w, h = layers[key][0], layers[key][1], layers[key][2], layers[key][3]
        if side == 'top': return (cx, cy + h/2)
        if side == 'bottom': return (cx, cy - h/2)
        if side == 'left': return (cx - w/2, cy)
        if side == 'right': return (cx + w/2, cy)

    # Vertical arrows (simple top-to-bottom connections)
    vert_connections = [
        ('user', 'planner'),
        ('planner', 'manager'),
        ('w2', 'tester'),
    ]
    for src, dst in vert_connections:
        x1, y1 = box_edge(src, 'bottom')
        x2, y2 = box_edge(dst, 'top')
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.5,
                                   shrinkA=3, shrinkB=3))

    # Manager to workers (fan-out) - use bottom-left, bottom-center, bottom-right of manager
    mcx, mcy, mw, mh = 5.0, 6.0, 2.6, 0.9
    worker_keys = ['w1', 'w2', 'wn']
    mgr_exit_xs = [mcx - mw/4, mcx, mcx + mw/4]
    for wk, mx in zip(worker_keys, mgr_exit_xs):
        x1, y1 = mx, mcy - mh/2
        x2, y2 = box_edge(wk, 'top')
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.4,
                                   shrinkA=3, shrinkB=3))

    # Workers to tester (fan-in)
    for wk in ['w1', 'wn']:
        x1, y1 = box_edge(wk, 'bottom')
        tcx, tcy, tw, th = 5.0, 1.0, 2.6, 0.9
        # Aim at left/right portion of tester top
        if wk == 'w1':
            x2 = tcx - tw/4
        else:
            x2 = tcx + tw/4
        y2 = tcy + th/2
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.4,
                                   shrinkA=3, shrinkB=3))

    # Feedback: L-shaped arrow from tester right side up to manager right side
    tx_r, ty_r = box_edge('tester', 'right')
    mx_r, my_r = box_edge('manager', 'right')
    # Route: go right from tester, then straight up, then left into manager
    route_x = 10.2  # x position of the vertical segment (clear of all boxes)
    # Draw path segments as lines
    ax.plot([tx_r, route_x], [ty_r, ty_r], color='#C62828', lw=1.5, linestyle='dashed', zorder=2)
    ax.plot([route_x, route_x], [ty_r, my_r], color='#C62828', lw=1.5, linestyle='dashed', zorder=2)
    # Final segment with arrow into manager
    ax.annotate('', xy=(mx_r, my_r), xytext=(route_x, my_r),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5,
                               shrinkA=0, shrinkB=3))
    # Label next to vertical segment
    ax.text(route_x + 0.4, (ty_r + my_r) / 2, 'Povratna\ninformacija', ha='left', va='center',
            fontsize=8, fontstyle='italic', color='#C62828')

    # Ellipsis dots between Radnik 2 and Radnik N
    for dx in [6.3, 6.8, 7.3]:
        ax.plot(dx, 3.5, '.', color='#666666', markersize=8)

    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    path = os.path.join(IMG_DIR, "multi_agent_architecture.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_llm_agent_components_diagram():
    """Creates a diagram showing LLM agent components."""
    fig, ax = plt.subplots(1, 1, figsize=(8, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')

    # Central LLM circle
    circle = plt.Circle((5, 4), 1.3, fill=True, facecolor='#E3F2FD',
                        edgecolor='#1565C0', linewidth=2.5)
    ax.add_patch(circle)
    ax.text(5, 4, 'LLM\n(Jezgro)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#1565C0')

    # Components positioned evenly around the circle
    components = [
        (2.0, 6.5, 'Memorija', '#E8F5E9', '#2E7D32'),
        (5.0, 7.3, 'Planiranje', '#FFF3E0', '#E65100'),
        (8.0, 6.5, 'Alati', '#F3E5F5', '#6A1B9A'),
        (2.0, 1.5, 'Percepcija', '#E0F7FA', '#00695C'),
        (8.0, 1.5, 'Akcija', '#FCE4EC', '#AD1457'),
    ]

    for (cx, cy, label, fcolor, ecolor) in components:
        rect = FancyBboxPatch((cx - 0.9, cy - 0.35), 1.8, 0.7, boxstyle="round,pad=0.08",
                              edgecolor=ecolor, facecolor=fcolor, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(cx, cy, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color=ecolor)

    # Connections from each component to the LLM circle edge
    center = np.array([5.0, 4.0])
    radius = 1.3
    for (cx, cy, _, _, ecolor) in components:
        comp = np.array([cx, cy])
        direction = center - comp
        dist = np.linalg.norm(direction)
        unit = direction / dist
        # Start point: from edge of box (approx 0.5 away from center of component)
        start = comp + unit * 0.9
        # End point: at circle edge
        end = center - unit * radius
        ax.annotate('', xy=(end[0], end[1]), xytext=(start[0], start[1]),
                    arrowprops=dict(arrowstyle='<->', color=ecolor, lw=1.3,
                                   shrinkA=0, shrinkB=0))

    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "llm_agent_components.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_task_planning_flowchart():
    """Creates a task planning and execution flowchart."""
    fig, ax = plt.subplots(1, 1, figsize=(11, 6))
    ax.set_xlim(-0.5, 13)
    ax.set_ylim(-1, 7)
    ax.axis('off')

    # Main flow steps - positioned with generous spacing
    steps = [
        (1.5, 4.0, 'Zahtev\nkorisnika', '#E8EAF6'),
        (4.0, 4.0, 'Dekompozicija\nzadatka', '#C8E6C9'),
        (6.5, 4.0, 'Alokacija\nagenata', '#BBDEFB'),
        (9.0, 4.0, 'Paralelno\nizvršavanje', '#FFF9C4'),
        (11.5, 4.0, 'Agregacija\nrezultata', '#FFCCBC'),
    ]

    box_w, box_h = 2.0, 1.1
    for i, (cx, cy, label, color) in enumerate(steps):
        rect = FancyBboxPatch((cx - box_w/2, cy - box_h/2), box_w, box_h,
                              boxstyle="round,pad=0.1",
                              edgecolor='#37474F', facecolor=color, linewidth=1.6)
        ax.add_patch(rect)
        ax.text(cx, cy, label, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='#212121')
        # Arrow to next step
        if i < len(steps) - 1:
            next_cx = steps[i+1][0]
            ax.annotate('', xy=(next_cx - box_w/2, cy),
                        xytext=(cx + box_w/2, cy),
                        arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.4,
                                       shrinkA=4, shrinkB=4))

    # Sub-tasks below "Paralelno izvršavanje" - with proper spacing
    sub_tasks = ['Zadatak A', 'Zadatak B', 'Zadatak C']
    sub_y_start = 2.0
    for j, st_label in enumerate(sub_tasks):
        sy = sub_y_start - j * 0.8
        rect = FancyBboxPatch((8.2, sy - 0.3), 1.6, 0.55,
                              boxstyle="round,pad=0.06",
                              edgecolor='#666666', facecolor='#FFFDE7', linewidth=1.0)
        ax.add_patch(rect)
        ax.text(9.0, sy, st_label, ha='center', va='center',
                fontsize=7.5, color='#333333')

    # Arrow from parallel box bottom to sub-tasks top
    ax.annotate('', xy=(9.0, sub_y_start + 0.3), xytext=(9.0, 4.0 - box_h/2),
                arrowprops=dict(arrowstyle='->', color='#455A64', lw=1.2,
                               shrinkA=4, shrinkB=4))

    # Brace or bracket visual for parallel
    ax.plot([8.0, 8.0], [sub_y_start - 2*0.8 - 0.3, sub_y_start + 0.3],
            color='#888888', lw=1.0, linestyle='-')
    ax.text(7.6, sub_y_start - 0.8, '||', ha='center', va='center',
            fontsize=12, color='#666666', fontweight='bold')

    # Feedback loop: curved dashed arrow from Agregacija top to Dekompozicija top
    ax.annotate('',
                xy=(4.0, 4.0 + box_h/2 + 0.1),
                xytext=(11.5, 4.0 + box_h/2 + 0.1),
                arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.4,
                               connectionstyle="arc3,rad=0.3",
                               shrinkA=4, shrinkB=4, linestyle='dashed'))
    ax.text(7.75, 6.2, 'Iterativno poboljšanje (feedback)', ha='center',
            fontsize=8, fontstyle='italic', color='#C62828')

    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    path = os.path.join(IMG_DIR, "task_planning_flow.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_parallel_execution_diagram():
    """Diagram showing parallel vs sequential execution."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4), gridspec_kw={'wspace': 0.3})

    for ax in [ax1, ax2]:
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 5.5)
        ax.axis('off')

    # === Sequential ===
    ax1.set_title('Sekvencijalno izvršavanje', fontsize=10, fontweight='bold',
                  color='#333333', pad=10)

    colors_seq = ['#BBDEFB', '#C8E6C9', '#FFF9C4']
    labels = ['Zadatak A', 'Zadatak B', 'Zadatak C']
    # Stacked horizontally end-to-end
    widths = [2.2, 2.2, 2.2]
    x_start = 1.0
    y_center = 3.0
    bar_h = 0.9

    for i, (w, label, color) in enumerate(zip(widths, labels, colors_seq)):
        x = x_start + sum(widths[:i])
        rect = FancyBboxPatch((x, y_center - bar_h/2), w, bar_h,
                              boxstyle="round,pad=0.03",
                              edgecolor='#333333', facecolor=color, linewidth=1.2)
        ax1.add_patch(rect)
        ax1.text(x + w/2, y_center, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color='#333333')

    # Time arrow
    ax1.annotate('', xy=(x_start + sum(widths) + 0.3, 1.5), xytext=(x_start - 0.3, 1.5),
                arrowprops=dict(arrowstyle='->', color='#666666', lw=1.0))
    ax1.text(x_start + sum(widths)/2, 1.0, 'Vreme →', ha='center', fontsize=7, color='#666666')
    ax1.text(5, 0.3, 'Ukupno: T(A) + T(B) + T(C)', ha='center',
            fontsize=8, fontstyle='italic', color='#444444')

    # === Parallel ===
    ax2.set_title('Paralelno izvršavanje', fontsize=10, fontweight='bold',
                  color='#333333', pad=10)

    y_positions = [4.0, 3.0, 2.0]
    par_w = 5.0
    x_par = 2.5

    for i, (y, label, color) in enumerate(zip(y_positions, labels, colors_seq)):
        rect = FancyBboxPatch((x_par, y - bar_h/2), par_w, bar_h,
                              boxstyle="round,pad=0.03",
                              edgecolor='#333333', facecolor=color, linewidth=1.2)
        ax2.add_patch(rect)
        ax2.text(x_par + par_w/2, y, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color='#333333')

    # Time arrow
    ax2.annotate('', xy=(x_par + par_w + 0.3, 0.8), xytext=(x_par - 0.3, 0.8),
                arrowprops=dict(arrowstyle='->', color='#666666', lw=1.0))
    ax2.text(x_par + par_w/2, 0.3, 'Ukupno: max(T(A), T(B), T(C))', ha='center',
            fontsize=8, fontstyle='italic', color='#444444')

    fig.patch.set_facecolor('white')
    ax1.set_facecolor('white')
    ax2.set_facecolor('white')
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "parallel_vs_sequential.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_communication_patterns_diagram():
    """Diagram showing agent communication patterns."""
    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(11, 4), gridspec_kw={'wspace': 0.1})

    configs = [
        (ax1, 'Centralizovano', 'star'),
        (ax2, 'Decentralizovano', 'mesh'),
        (ax3, 'Hijerarhijsko', 'tree')
    ]

    for ax, title, pattern in configs:
        ax.set_xlim(0, 6)
        ax.set_ylim(0, 6)
        ax.axis('off')
        ax.set_title(title, fontsize=10, fontweight='bold', color='#333333', pad=8)

        if pattern == 'star':
            center = (3, 3)
            nodes = [(1.2, 5.0), (4.8, 5.0), (1.2, 1.0), (4.8, 1.0)]
            # Draw lines first (behind nodes)
            for n in nodes:
                ax.plot([center[0], n[0]], [center[1], n[1]],
                       color='#666666', lw=1.2, zorder=1)
            # Manager node
            c = plt.Circle(center, 0.45, facecolor='#BBDEFB', edgecolor='#1565C0',
                          lw=2, zorder=3)
            ax.add_patch(c)
            ax.text(*center, 'M', ha='center', va='center', fontsize=9,
                   fontweight='bold', color='#1565C0', zorder=4)
            # Agent nodes
            for n in nodes:
                c = plt.Circle(n, 0.35, facecolor='#C8E6C9', edgecolor='#2E7D32',
                              lw=1.5, zorder=3)
                ax.add_patch(c)
                ax.text(*n, 'A', ha='center', va='center', fontsize=8,
                       color='#2E7D32', zorder=4)

        elif pattern == 'mesh':
            nodes = [(1.5, 4.5), (4.5, 4.5), (1.5, 1.5), (4.5, 1.5)]
            # Draw all connections first
            for i, n1 in enumerate(nodes):
                for n2 in nodes[i+1:]:
                    ax.plot([n1[0], n2[0]], [n1[1], n2[1]],
                           color='#666666', lw=1.0, zorder=1)
            # Draw nodes on top
            for n in nodes:
                c = plt.Circle(n, 0.35, facecolor='#FFF9C4', edgecolor='#F57F17',
                              lw=1.5, zorder=3)
                ax.add_patch(c)
                ax.text(*n, 'A', ha='center', va='center', fontsize=8,
                       color='#F57F17', fontweight='bold', zorder=4)

        elif pattern == 'tree':
            top = (3, 5)
            mid = [(1.5, 3.2), (4.5, 3.2)]
            bottom = [(0.7, 1.2), (2.3, 1.2), (3.7, 1.2), (5.3, 1.2)]
            # Draw lines first
            for m in mid:
                ax.plot([top[0], m[0]], [top[1], m[1]], color='#666666', lw=1.2, zorder=1)
            for i, b in enumerate(bottom):
                parent = mid[0] if i < 2 else mid[1]
                ax.plot([parent[0], b[0]], [parent[1], b[1]],
                       color='#666666', lw=1.0, zorder=1)
            # Top node
            c = plt.Circle(top, 0.4, facecolor='#FFCCBC', edgecolor='#BF360C',
                          lw=2, zorder=3)
            ax.add_patch(c)
            ax.text(*top, 'P', ha='center', va='center', fontsize=9,
                   fontweight='bold', color='#BF360C', zorder=4)
            # Mid nodes
            for m in mid:
                c = plt.Circle(m, 0.35, facecolor='#BBDEFB', edgecolor='#1565C0',
                              lw=1.5, zorder=3)
                ax.add_patch(c)
                ax.text(*m, 'M', ha='center', va='center', fontsize=8,
                       color='#1565C0', fontweight='bold', zorder=4)
            # Bottom nodes
            for b in bottom:
                c = plt.Circle(b, 0.3, facecolor='#C8E6C9', edgecolor='#2E7D32',
                              lw=1.2, zorder=3)
                ax.add_patch(c)
                ax.text(*b, 'R', ha='center', va='center', fontsize=7,
                       color='#2E7D32', zorder=4)

    fig.patch.set_facecolor('white')
    for a in [ax1, ax2, ax3]:
        a.set_facecolor('white')
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "communication_patterns.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_cost_quality_analysis_diagram():
    """Creates a comparative diagram: multi-agent vs single-agent token cost and quality."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 5), gridspec_kw={'wspace': 0.35})

    # === Left: Token cost comparison (bar chart) ===
    ax1.set_facecolor('white')

    categories = ['Planiranje', 'Koordinacija', 'Izvršavanje\n(5 zadataka)', 'Verifikacija', 'UKUPNO']
    # Multi-agent: planner(strong) + manager(medium) + 5 workers(weak) + tester(medium)
    # Single-agent: one strong model doing everything with growing context
    multi_agent_tokens = [4000, 3000, 12500, 3500, 23000]  # sum of specialized calls
    single_agent_tokens = [0, 0, 45000, 0, 45000]  # one massive context window

    x = np.arange(len(categories))
    width = 0.35

    bars1 = ax1.bar(x - width/2, [t/1000 for t in multi_agent_tokens], width,
                    label='Multi-agent', color='#4CAF50', edgecolor='#2E7D32', linewidth=0.8)
    bars2 = ax1.bar(x + width/2, [t/1000 for t in single_agent_tokens], width,
                    label='Mono-agent', color='#FF7043', edgecolor='#BF360C', linewidth=0.8)

    ax1.set_ylabel('Tokeni (×1000)', fontsize=9, color='#333333')
    ax1.set_title('Potrošnja tokena po fazi', fontsize=10, fontweight='bold', color='#333333', pad=10)
    ax1.set_xticks(x)
    ax1.set_xticklabels(categories, fontsize=7.5, color='#333333')
    ax1.legend(fontsize=8, loc='upper left')
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.tick_params(colors='#333333')
    ax1.yaxis.grid(True, alpha=0.3)
    ax1.set_axisbelow(True)

    # Add value labels on bars
    for bar in bars1:
        h = bar.get_height()
        if h > 0:
            ax1.text(bar.get_x() + bar.get_width()/2., h + 0.3,
                    f'{h:.1f}k', ha='center', va='bottom', fontsize=6.5, color='#2E7D32')
    for bar in bars2:
        h = bar.get_height()
        if h > 0:
            ax1.text(bar.get_x() + bar.get_width()/2., h + 0.3,
                    f'{h:.1f}k', ha='center', va='bottom', fontsize=6.5, color='#BF360C')

    # === Right: Quality vs Cost scatter/quadrant ===
    ax2.set_facecolor('white')
    ax2.set_xlim(0, 10)
    ax2.set_ylim(0, 10)

    # Quadrant lines
    ax2.axhline(y=5, color='#BDBDBD', linestyle='--', lw=0.8)
    ax2.axvline(x=5, color='#BDBDBD', linestyle='--', lw=0.8)

    # Points: (cost, quality)
    # Multi-agent approach with model tiering
    ax2.scatter([3.5], [8.2], s=200, c='#4CAF50', edgecolors='#2E7D32', linewidth=1.5, zorder=5)
    ax2.text(3.5, 7.3, 'Multi-agent\n(tiered models)', ha='center', fontsize=7.5,
            color='#2E7D32', fontweight='bold')

    # Single strong agent
    ax2.scatter([7.5], [7.8], s=200, c='#FF7043', edgecolors='#BF360C', linewidth=1.5, zorder=5)
    ax2.text(7.5, 6.9, 'Mono-agent\n(jak model)', ha='center', fontsize=7.5,
            color='#BF360C', fontweight='bold')

    # Single weak agent (cheap but bad)
    ax2.scatter([2.0], [3.5], s=150, c='#FFCC80', edgecolors='#E65100', linewidth=1.5, zorder=5)
    ax2.text(2.0, 2.6, 'Mono-agent\n(slab model)', ha='center', fontsize=7.5,
            color='#E65100', fontweight='bold')

    # Multi-agent all strong (expensive)
    ax2.scatter([8.5], [8.8], s=150, c='#CE93D8', edgecolors='#6A1B9A', linewidth=1.5, zorder=5)
    ax2.text(8.5, 8.0, 'Multi-agent\n(svi jaki)', ha='center', fontsize=7.5,
            color='#6A1B9A', fontweight='bold')

    ax2.set_xlabel('Relativna cena (tokeni × cena/token)', fontsize=8.5, color='#333333')
    ax2.set_ylabel('Kvalitet izlaza', fontsize=8.5, color='#333333')
    ax2.set_title('Odnos cene i kvaliteta', fontsize=10, fontweight='bold', color='#333333', pad=10)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.tick_params(colors='#333333')

    # Quadrant labels
    ax2.text(2.5, 9.3, 'OPTIMALNO', ha='center', fontsize=7, color='#4CAF50',
            fontstyle='italic', fontweight='bold')
    ax2.text(7.5, 9.3, 'Skupo ali kvalitetno', ha='center', fontsize=7, color='#666666',
            fontstyle='italic')
    ax2.text(2.5, 0.5, 'Jeftino ali loše', ha='center', fontsize=7, color='#666666',
            fontstyle='italic')
    ax2.text(7.5, 0.5, 'NAJGORE', ha='center', fontsize=7, color='#C62828',
            fontstyle='italic', fontweight='bold')

    fig.patch.set_facecolor('white')
    plt.tight_layout()
    path = os.path.join(IMG_DIR, "cost_quality_analysis.png")
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def add_heading(doc, text, level, numbering=None):
    """Add a heading with Times New Roman, Automatic color (black)."""
    if numbering:
        heading = doc.add_heading(f'{numbering} {text}', level=level)
    else:
        heading = doc.add_heading(text, level=level)
    # Force Times New Roman, black color on all heading runs
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Automatic/black
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, italic=False, font_size=12):
    """Add a justified paragraph with Times New Roman 12pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_image(doc, path, width=Inches(5.5)):
    """Add image centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    return p


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5

    # Style headings to use Times New Roman, black color, 1.5 spacing
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_style.paragraph_format.line_spacing = 1.5
        if i == 1:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    # Style list items
    for list_style_name in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = 'Times New Roman'
            ls.font.size = Pt(12)
            ls.paragraph_format.line_spacing = 1.5
            ls.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    # Set margins - standard academic (2.54cm, left 3cm for binding)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    # ========== PAGE 1: Title Page (placeholder) ==========
    for _ in range(8):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Agenti zasnovani na Velikim jezičkim modelima\ni planiranje zadataka')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Seminarski rad')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Predmet: Inteligentni agenti\n[Univerzitet / Fakultet]\n[Ime i prezime studenta]\n[Ime profesora]\n[Godina]')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    add_page_break(doc)

    # ========== PAGE 2: Table of Contents ==========
    add_heading(doc, 'Sadržaj', level=1)
    
    toc_items = [
        ('1.', 'Uvod', '3'),
        ('2.', 'Veliki jezički modeli kao osnova inteligentnih agenata', '4'),
        ('2.1.', 'Arhitektura Transformer modela', '4'),
        ('2.2.', 'Sposobnosti LLM-a relevantne za agente', '5'),
        ('3.', 'Inteligentni agenti zasnovani na LLM', '6'),
        ('3.1.', 'Definicija i komponente LLM agenta', '6'),
        ('3.2.', 'Planiranje i rasuđivanje', '7'),
        ('3.3.', 'Memorija i učenje', '7'),
        ('4.', 'Multi-agent sistemi', '8'),
        ('4.1.', 'Komunikacioni obrasci', '8'),
        ('4.2.', 'Koordinacija i saradnja', '9'),
        ('5.', 'Planiranje i paralelno izvršavanje zadataka', '10'),
        ('5.1.', 'Dekompozicija zadataka', '10'),
        ('5.2.', 'Paralelno vs. sekvencijalno izvršavanje', '11'),
        ('6.', 'Arhitektura Planer-Menadžer-Radnik-Tester', '12'),
        ('6.1.', 'Uloge agenata', '12'),
        ('6.2.', 'Tok izvršavanja', '13'),
        ('6.3.', 'Iterativno poboljšanje kroz povratnu informaciju', '13'),
        ('7.', 'Primeri primene i diskusija', '14'),
        ('7.1.', 'Komparativna analiza: multi-agent vs. mono-agent', '15'),
        ('8.', 'Zaključak', '17'),
        ('', 'Literatura', '18'),
    ]
    
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num} {title}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        # Add dots and page number
        tab_run = p.add_run(f'  {"." * (50 - len(title))}  {page}')
        tab_run.font.size = Pt(12)
        tab_run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5

    add_page_break(doc)

    # ========== CHAPTER 1: Uvod ==========
    add_heading(doc, '1. Uvod', level=1)
    
    add_paragraph(doc, 
        'Razvoj veštačke inteligencije u poslednjih nekoliko godina doživeo je revolucionarne promene, '
        'pre svega zahvaljujući napretku u oblasti velikih jezičkih modela (eng. Large Language Models - LLM). '
        'Modeli poput GPT-4 [14], Claude i LLaMA demonstrirali su izuzetne sposobnosti u razumevanju i '
        'generisanju prirodnog jezika, rešavanju složenih problema i rezonovanju o apstraktnim konceptima. '
        'Ove sposobnosti otvorile su put ka novoj paradigmi u računarstvu — autonomnim agentima zasnovanim '
        'na jezičkim modelima [1, 2].')

    add_paragraph(doc,
        'Inteligentni agenti zasnovani na LLM-u predstavljaju sisteme koji koriste veliki jezički model '
        'kao centralni mehanizam za rasuđivanje i donošenje odluka, proširujući ga sposobnostima percepcije '
        'okruženja, planiranja akcija, korišćenja alata i učenja iz iskustva. Za razliku od tradicionalnih '
        'agenata čije je ponašanje determinističko i unapred definisano, LLM agenti pokazuju emergentno '
        'ponašanje i fleksibilnost u suočavanju sa novim situacijama [2].')

    add_paragraph(doc,
        'Posebno interesantan pravac istraživanja predstavljaju multi-agent sistemi u kojima više '
        'specijalizovanih agenata sarađuje na rešavanju složenih zadataka. Ovaj pristup, inspirisan '
        'principima podele rada i specijalizacije u ljudskim organizacijama, omogućava dekompoziciju '
        'kompleksnih problema na manje podzadatke koji se mogu izvršavati paralelno [3, 10].')

    add_paragraph(doc,
        'U ovom radu detaljno analiziramo arhitekturu i principe funkcionisanja multi-agent sistema '
        'zasnovanih na LLM-u, sa posebnim fokusom na planiranje zadataka i njihovo paralelno '
        'izvršavanje. Predstavljamo hijerarhijsku arhitekturu Planer-Menadžer-Radnik-Tester koja '
        'demonstrira efikasan pristup koordinaciji autonomnih agenata u kompleksnim softverskim '
        'sistemima.')

    add_page_break(doc)

    # ========== CHAPTER 2: LLM kao osnova ==========
    add_heading(doc, '2. Veliki jezički modeli kao osnova inteligentnih agenata', level=1)

    add_heading(doc, '2.1. Arhitektura Transformer modela', level=2)

    add_paragraph(doc,
        'Transformer arhitektura, predstavljena u radu "Attention Is All You Need" [13], '
        'predstavlja temelj svih modernih velikih jezičkih modela. Ova arhitektura uvodi mehanizam '
        'samo-pažnje (eng. self-attention) koji omogućava modelu da istovremeno razmatra sve delove '
        'ulazne sekvence, čime se prevazilaze ograničenja rekurentnih neuronskih mreža u pogledu '
        'paralelizacije i modelovanja dugoročnih zavisnosti.')

    add_paragraph(doc,
        'Veliki jezički modeli trenirani su na enormnim količinama tekstualnih podataka koristeći '
        'princip predviđanja sledećeg tokena. Ovaj naizgled jednostavan cilj treniranja rezultira '
        'modelima koji internalizuju bogato znanje o svetu, jeziku, logici i rasuđivanju. Skaliranje '
        'modela — povećanje broja parametara i količine podataka za treniranje — dovodi do emergentnih '
        'sposobnosti koje nisu prisutne u manjim modelima [12].')

    add_heading(doc, '2.2. Sposobnosti LLM-a relevantne za agente', level=2)

    add_paragraph(doc,
        'Jezički modeli poseduju nekoliko ključnih sposobnosti koje ih čine pogodnim za ulogu '
        'centralnog mehanizma inteligentnih agenata:')

    capabilities = [
        'Razumevanje prirodnog jezika — sposobnost parsiranja i interpretacije složenih instrukcija korisnika.',
        'Rasuđivanje u više koraka (eng. Chain-of-Thought) — sposobnost dekompozicije problema na sekvencijalne korake [7].',
        'Generisanje strukturiranog izlaza — sposobnost produkcije koda, JSON-a, planova akcija.',
        'Kontekstualno učenje (eng. In-context learning) — sposobnost adaptacije ponašanja na osnovu primera datih u promptu [12].',
        'Korišćenje alata (eng. Tool use) — sposobnost pozivanja eksternih funkcija i API-ja na osnovu opisa.',
    ]
    
    for cap in capabilities:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(cap)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    add_paragraph(doc,
        'Ove sposobnosti, kombinovane sa odgovarajućom arhitekturom sistema, omogućavaju '
        'konstrukciju autonomnih agenata sposobnih za planiranje, izvršavanje i evaluaciju '
        'složenih zadataka bez stalnog nadzora čoveka.')

    add_page_break(doc)

    # ========== CHAPTER 3: LLM Agenti ==========
    add_heading(doc, '3. Inteligentni agenti zasnovani na LLM', level=1)

    add_heading(doc, '3.1. Definicija i komponente LLM agenta', level=2)

    add_paragraph(doc,
        'LLM agent se definiše kao autonomni sistem koji koristi veliki jezički model kao '
        'kognitivno jezgro za percepciju okruženja, donošenje odluka i preduzimanje akcija '
        'u cilju ostvarivanja zadatog cilja [1]. Za razliku od jednostavnog korišćenja LLM-a '
        'kao alata za generisanje teksta, agent poseduje sposobnost interakcije sa okruženjem '
        'i akumuliranja iskustva tokom vremena.')

    add_paragraph(doc,
        'Na slici 1 prikazane su ključne komponente LLM agenta. Centralni element je sam '
        'jezički model koji služi kao mehanizam rasuđivanja, okružen modulima za percepciju, '
        'planiranje, memoriju, korišćenje alata i izvršavanje akcija.')

    # Insert LLM agent components diagram
    img_path = create_llm_agent_components_diagram()
    add_image(doc, img_path, width=Inches(4.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Slika 1: Komponente LLM agenta [1, 2]')
    run.italic = True
    run.font.size = Pt(10)

    add_paragraph(doc,
        'Percepcija omogućava agentu da prima informacije iz okruženja — tekst korisnika, '
        'sadržaj fajlova, izlaz komandi ili poruke drugih agenata. Modul za planiranje '
        'koristi LLM za razbijanje složenih ciljeva na izvršive korake. Memorija obuhvata '
        'kratkoročno pamćenje (kontekst konverzacije) i dugoročno pamćenje (vektorske baze, '
        'baze znanja). Alati proširuju sposobnosti agenta izvan čistog generisanja teksta, '
        'omogućavajući mu izvršavanje koda, pretraživanje interneta ili manipulaciju fajlovima.')

    add_heading(doc, '3.2. Planiranje i rasuđivanje', level=2)

    add_paragraph(doc,
        'Planiranje predstavlja jednu od najvažnijih sposobnosti LLM agenata. Moderni pristupi '
        'planiranju uključuju:')

    planning_approaches = [
        'Chain-of-Thought (CoT) — rasuđivanje korak po korak pre davanja konačnog odgovora [7].',
        'ReAct — iterativno smenjivanje rasuđivanja i akcija, gde agent na osnovu opservacija iz okruženja prilagođava plan [6].',
        'Reflexion — mehanizam samo-refleksije gde agent evaluira sopstveni učinak i uči iz grešaka [8].',
        'Tree-of-Thought — istovremeno razmatranje više strategija rešavanja sa evaluacijom svake.',
    ]
    
    for approach in planning_approaches:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(approach)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    add_heading(doc, '3.3. Memorija i učenje', level=2)

    add_paragraph(doc,
        'LLM agenti koriste različite mehanizme memorije za akumuliranje i korišćenje znanja. '
        'Kratkoročna memorija realizuje se kroz kontekstni prozor jezičkog modela, dok se dugoročna '
        'memorija implementira kroz eksterne sisteme skladištenja — vektorske baze podataka za '
        'semantičko pretraživanje, strukturirane baze za činjenično znanje, ili čak same LLM-ove '
        'za kompresiju i sumarizaciju prošlih iskustava [1].')

    add_paragraph(doc,
        'Posebno je značajan koncept episodičke memorije — sposobnost agenta da zapamti konkretne '
        'epizode rada (uspehe i neuspehe) i koristi ih za poboljšanje budućih performansi. Ovaj '
        'mehanizam je analogan ljudskom iskustvenom učenju i predstavlja ključni faktor u '
        'postizanju konzistentnog poboljšanja performansi tokom vremena [8].')

    add_page_break(doc)

    # ========== CHAPTER 4: Multi-agent sistemi ==========
    add_heading(doc, '4. Multi-agent sistemi', level=1)

    add_paragraph(doc,
        'Multi-agent sistemi zasnovani na LLM-u predstavljaju paradigmu u kojoj više specijalizovanih '
        'agenata sarađuje na rešavanju zadataka koji prevazilaze sposobnosti pojedinačnog agenta. '
        'Ovaj pristup inspirisan je principima organizacije ljudskih timova — specijalizacijom, '
        'podelom rada i koordinacijom [3, 10].')

    add_paragraph(doc,
        'Prednosti multi-agent pristupa u odnosu na jednog monolitnog agenta su višestruke: '
        '(1) svaki agent može biti specijalizovan za uži domen, čime se povećava kvalitet; '
        '(2) zadaci se mogu izvršavati paralelno, čime se smanjuje ukupno vreme; '
        '(3) sistem je modularan i lakše se održava; '
        '(4) greške su izolovane — neuspeh jednog agenta ne mora kompromitovati ceo sistem [4].')

    add_heading(doc, '4.1. Komunikacioni obrasci', level=2)

    add_paragraph(doc,
        'Komunikacija između agenata može biti organizovana prema različitim obrascima, '
        'prikazanim na slici 2:')

    img_path = create_communication_patterns_diagram()
    add_image(doc, img_path, width=Inches(5.0))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Slika 2: Komunikacioni obrasci u multi-agent sistemima')
    run.italic = True
    run.font.size = Pt(10)

    add_paragraph(doc,
        'U centralizovanom obrascu, jedan koordinatorski agent upravlja komunikacijom između svih '
        'ostalih. Decentralizovani obrazac omogućava svim agentima direktnu komunikaciju, što '
        'povećava fleksibilnost ali i kompleksnost. Hijerarhijski obrazac, koji koristimo u '
        'arhitekturi Planer-Menadžer-Radnik-Tester, kombinuje prednosti oba pristupa — jasnu '
        'strukturu upravljanja sa mogućnošću paralelnog rada na nižim nivoima [9, 10].')

    add_heading(doc, '4.2. Koordinacija i saradnja', level=2)

    add_paragraph(doc,
        'Efektivna koordinacija u multi-agent sistemima zahteva rešavanje nekoliko ključnih izazova: '
        'raspodela zadataka, sinhronizacija rezultata, rešavanje konflikata i upravljanje zavisnostima '
        'između podzadataka. Sistemi poput AutoGen [3] i CAMEL [9] predlažu različite mehanizme '
        'koordinacije, od rigidnih protokola do emergentne saradnje kroz slobodnu konverzaciju.')

    add_paragraph(doc,
        'MetaGPT [4] uvodi koncept "standardnih operativnih procedura" (SOP) za multi-agent saradnju, '
        'gde su uloge agenata i tokovi rada formalno definisani, što smanjuje redundanciju i povećava '
        'efikasnost. Ovaj pristup je posebno efektivan u domenima gde su procesi dobro definisani, '
        'poput razvoja softvera.')

    add_page_break(doc)

    # ========== CHAPTER 5: Planiranje i paralelno izvršavanje ==========
    add_heading(doc, '5. Planiranje i paralelno izvršavanje zadataka', level=1)

    add_heading(doc, '5.1. Dekompozicija zadataka', level=2)

    add_paragraph(doc,
        'Dekompozicija zadataka predstavlja proces razbijanja složenog cilja na manje, '
        'upravljive podzadatke koji se mogu nezavisno izvršavati. U kontekstu LLM agenata, '
        'ovaj proces se realizuje korišćenjem samog jezičkog modela koji analizira zahtev '
        'korisnika i generiše strukturiran plan izvršavanja [1, 11].')

    add_paragraph(doc,
        'Na slici 3 prikazan je tok planiranja zadataka od zahteva korisnika do agregacije '
        'rezultata. Ključni koraci uključuju: analizu zahteva, identifikaciju podzadataka, '
        'utvrđivanje zavisnosti među njima, alokaciju resursa (agenata) i konačno izvršavanje '
        'sa agregacijom rezultata.')

    img_path = create_task_planning_flowchart()
    add_image(doc, img_path, width=Inches(5.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Slika 3: Tok planiranja i izvršavanja zadataka u multi-agent sistemu')
    run.italic = True
    run.font.size = Pt(10)

    add_paragraph(doc,
        'Efikasna dekompozicija zahteva identifikaciju zavisnosti između podzadataka. Zadaci '
        'koji nemaju međusobne zavisnosti mogu se izvršavati paralelno, dok se zavisni zadaci '
        'moraju sekvencijalizovati. Ova analiza zavisnosti je ključna za maksimizaciju paralelizma '
        'i minimizaciju ukupnog vremena izvršavanja.')

    add_heading(doc, '5.2. Paralelno vs. sekvencijalno izvršavanje', level=2)

    add_paragraph(doc,
        'Jedna od fundamentalnih prednosti multi-agent sistema je mogućnost paralelnog izvršavanja '
        'nezavisnih zadataka. Na slici 4 ilustrovana je razlika između sekvencijalnog i paralelnog '
        'pristupa izvršavanju.')

    img_path = create_parallel_execution_diagram()
    add_image(doc, img_path, width=Inches(5.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Slika 4: Poređenje sekvencijalnog i paralelnog izvršavanja zadataka')
    run.italic = True
    run.font.size = Pt(10)

    add_paragraph(doc,
        'U sekvencijalnom režimu, ukupno vreme izvršavanja jednako je zbiru vremena svih '
        'zadataka: T_total = T(A) + T(B) + T(C). U paralelnom režimu, ukupno vreme je '
        'određeno najsporijim zadatkom: T_total = max(T(A), T(B), T(C)). Za N nezavisnih '
        'zadataka približno istog trajanja, paralelno izvršavanje pruža ubrzanje faktora N.')

    add_paragraph(doc,
        'U praksi, potpuna paralelizacija retko je moguća jer postoje zavisnosti između zadataka, '
        'potreba za sinhronizacijom i ograničeni resursi. Efikasan planer mora balansirati između '
        'maksimalnog paralelizma i poštovanja zavisnosti, koristeći tehnike poput topološkog '
        'sortiranja grafa zavisnosti za određivanje optimalnog rasporeda izvršavanja [10].')

    add_page_break(doc)

    # ========== CHAPTER 6: Arhitektura Planer-Menadžer-Radnik-Tester ==========
    add_heading(doc, '6. Arhitektura Planer-Menadžer-Radnik-Tester', level=1)

    add_paragraph(doc,
        'U ovom poglavlju detaljno opisujemo hijerarhijsku multi-agent arhitekturu koja se '
        'sastoji od četiri specijalizovane uloge: Planer, Menadžer, Radnik i Tester. Ova '
        'arhitektura je inspirisana realnim razvojnim timovima i implementirana je u savremenim '
        'alatima za automatizaciju softverskog razvoja.')

    add_heading(doc, '6.1. Uloge agenata', level=2)

    add_paragraph(doc,
        'Svaki agent u sistemu ima jasno definisanu ulogu i odgovornost:', bold=True)

    add_paragraph(doc,
        'Planer Agent — Prima zahtev korisnika i vrši inicijalnu analizu i dekompoziciju. '
        'Koristi LLM za razumevanje konteksta, identifikaciju podzadataka i kreiranje '
        'strukturiranog plana izvršavanja. Plan uključuje opis svakog zadatka, procenu složenosti, '
        'zavisnosti i prioritete. Planer ne izvršava zadatke direktno, već delegira menadžeru.')

    add_paragraph(doc,
        'Menadžer Agent — Koordinira izvršavanje plana. Prima plan od Planera i vrši alokaciju '
        'zadataka radnicima, vodeći računa o zavisnostima i mogućnostima paralelizacije. Menadžer '
        'prati status izvršavanja, upravlja redosledom i donosi odluke o ponovnim pokušajima u '
        'slučaju neuspeha.')

    add_paragraph(doc,
        'Radnik Agent — Izvršava konkretne zadatke. Svaki radnik je specijalizovan za određeni '
        'tip operacija (editovanje koda, kreiranje fajlova, izvršavanje komandi). Više radnika '
        'može raditi paralelno na nezavisnim zadacima, čime se značajno ubrzava proces. Radnik '
        'prima precizne instrukcije od menadžera i vraća rezultat po završetku.')

    add_paragraph(doc,
        'Tester Agent — Verifikuje rezultate rada radnika. Po završetku zadataka, tester '
        'proverava korektnost implementacije, pokreće testove, analizira konzistentnost i '
        'identifikuje potencijalne probleme. Ukoliko tester pronađe greške, šalje povratnu '
        'informaciju menadžeru koji može ponovo delegirati zadatak radniku sa korigovanim '
        'instrukcijama.')

    add_heading(doc, '6.2. Tok izvršavanja', level=2)

    add_paragraph(doc,
        'Na slici 5 prikazan je dijagram toka izvršavanja u arhitekturi Planer-Menadžer-Radnik-Tester. '
        'Sistem funkcioniše kroz sledeće faze:')

    img_path = create_multi_agent_architecture_diagram()
    add_image(doc, img_path, width=Inches(4.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('Slika 5: Arhitektura Planer-Menadžer-Radnik-Tester')
    run.italic = True
    run.font.size = Pt(10)

    phases = [
        'Faza 1 — Planiranje: Korisnik postavlja zahtev. Planer analizira zahtev i kreira detaljan plan sa listom zadataka, njihovim opisima i zavisnostima.',
        'Faza 2 — Koordinacija: Menadžer prima plan i identifikuje zadatke koji se mogu paralelno izvršiti. Kreira graf zavisnosti i određuje optimalan raspored.',
        'Faza 3 — Izvršavanje: Menadžer delegira zadatke radnicima. Nezavisni zadaci se izvršavaju paralelno. Zavisni zadaci čekaju završetak prethodnika.',
        'Faza 4 — Verifikacija: Po završetku zadataka, tester proverava rezultate. Ako su rezultati korektni, proces se nastavlja. U suprotnom, generiše se povratna informacija.',
        'Faza 5 — Iteracija: Na osnovu povratne informacije testera, menadžer može ponovo delegirati zadatke sa korigovanim instrukcijama. Ciklus se ponavlja do postizanja zadovoljavajućeg rezultata.',
    ]
    
    for phase in phases:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(phase)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    add_heading(doc, '6.3. Iterativno poboljšanje kroz povratnu informaciju', level=2)

    add_paragraph(doc,
        'Ključna karakteristika ove arhitekture je ciklus povratne informacije između testera '
        'i menadžera. Ovaj mehanizam, inspirisan principima iz Reflexion pristupa [8], omogućava '
        'sistemu da iterativno poboljšava kvalitet izlaza. Proces prati obrazac:')

    add_paragraph(doc,
        'Write → Review → Fix → Review → Fix (ako je potrebno) → Review...', bold=True)

    add_paragraph(doc,
        'Ovaj iterativni pristup obezbeđuje konvergenciju ka korektnom rešenju i značajno '
        'smanjuje verovatnoću isporuke nekvalitetnog rezultata. U praksi, većina zadataka '
        'zahteva jednu do dve iteracije korekcije, dok se složeniji zadaci mogu korigovati '
        'u tri do pet iteracija pre nego što se postigne zadovoljavajući kvalitet.')

    add_paragraph(doc,
        'Prednost ovog pristupa u odnosu na monolitnog agenta je jasna separacija odgovornosti: '
        'radnik se fokusira na implementaciju, a tester na verifikaciju. Ova podela sprečava '
        'pristrasnost — agent koji je kreirao rešenje skloniji je da previdi sopstvene greške, '
        'dok nezavisni tester pristupa evaluaciji objektivnije.')

    add_page_break(doc)

    # ========== CHAPTER 7: Primeri primene ==========
    add_heading(doc, '7. Primeri primene i diskusija', level=1)

    add_paragraph(doc,
        'Multi-agent arhitektura opisana u prethodnom poglavlju nalazi primenu u različitim '
        'domenima. U nastavku razmatramo konkretne primere i diskutujemo prednosti i ograničenja.')

    add_paragraph(doc, 'Automatizacija razvoja softvera', bold=True)

    add_paragraph(doc,
        'Jedan od najrazvijenijih domena primene multi-agent sistema je automatizacija '
        'softverskog razvoja. Sistemi poput MetaGPT [4] i AutoGen [3] demonstriraju kako '
        'tim agenata (arhitekta, programer, tester, reviewer) može autonomno razvijati '
        'softverske komponente. Planer razlaže funkcionalne zahteve na implementacione zadatke, '
        'radnici paralelno implementiraju nezavisne module, a tester verifikuje korektnost '
        'kroz automatsko pokretanje testova.')

    add_paragraph(doc, 'Obrada složenih korisničkih zahteva', bold=True)

    add_paragraph(doc,
        'U kontekstu korisničke podrške i interakcije sa kompleksnim sistemima, multi-agent '
        'pristup omogućava efikasnu obradu zahteva koji zahtevaju pristup različitim izvorima '
        'podataka ili sistemima. Na primer, zahtev koji uključuje pretraživanje baze znanja, '
        'izvršavanje koda i generisanje izveštaja može se paralelizovati tako da svaki agent '
        'obrađuje svoj deo nezavisno, a rezultati se na kraju agregiraju [11].')

    add_paragraph(doc, 'Ograničenja i izazovi', bold=True)

    add_paragraph(doc,
        'Uprkos značajnim prednostima, multi-agent sistemi suočavaju se sa nekoliko izazova: '
        '(1) povećana cena — svaki agent zahteva zasebne pozive jezičkom modelu; '
        '(2) latencija koordinacije — komunikacija između agenata unosi dodatno kašnjenje; '
        '(3) propagacija grešaka — greška u ranoj fazi može se propagirati kroz lanac agenata; '
        '(4) složenost debagovanja — interakcije između više agenata teško je pratiti i analizirati. '
        'Rešavanje ovih izazova predstavlja aktivan pravac istraživanja u zajednici [2, 10].')

    add_page_break(doc)

    # ========== SECTION 7.1: Komparativna analiza ==========
    add_heading(doc, '7.1. Komparativna analiza: multi-agent vs. mono-agent pristup', level=2)

    add_paragraph(doc,
        'Ključno pitanje pri projektovanju sistema zasnovanih na LLM-u je da li koristiti '
        'jednog monolitnog agenta koji obavlja sve zadatke, ili multi-agent sistem sa '
        'specijalizovanim ulogama. Ova analiza razmatra dva osnovna aspekta: potrošnju tokena '
        '(koja direktno determiniše cenu) i kvalitet izlaza.')

    add_paragraph(doc, 'Model rasporeda u multi-agent arhitekturi', bold=True)

    add_paragraph(doc,
        'U predloženoj hijerarhijskoj arhitekturi koristi se princip raslojavanja modela '
        '(eng. model tiering): Planer koristi najjači dostupni model (npr. GPT-4, Claude Opus) '
        'jer je kvalitet dekompozicije kritičan za celokupan proces. Menadžer koristi model '
        'srednje klase (npr. Claude Sonnet, GPT-4o) koji je dovoljan za koordinaciju i alokaciju. '
        'Radnici koriste slabije, jeftinije modele (npr. Claude Haiku, GPT-4o-mini) jer dobijaju '
        'precizne, dobro definisane instrukcije od menadžera — kvalitet ulaza kompenzuje slabiji model.')

    add_paragraph(doc, 'Analiza potrošnje tokena', bold=True)

    add_paragraph(doc,
        'Razmotrimo scenario sa 5 nezavisnih podzadataka. Mono-agent (jak model) mora da '
        'održava celokupan kontekst u jednom prozoru: akumulira sve prethodne rezultate, '
        'instrukcije i međurezultate. Sa svakim novim podzadatkom, kontekst raste — za peti '
        'zadatak, agent procesira ~40-50K tokena ulaza jer mora da „vidi" sve prethodno. '
        'Ukupna potrošnja: ~45.000 tokena na skupom modelu.')

    add_paragraph(doc,
        'Multi-agent pristup: Planer troši ~4.000 tokena (jedna dekompozicija na jakom modelu). '
        'Menadžer troši ~3.000 tokena za koordinaciju (srednji model). Svaki od 5 radnika '
        'dobija izolovan kontekst od ~2.500 tokena (ukupno 12.500 na jeftinom modelu). '
        'Tester troši ~3.500 tokena za verifikaciju (srednji model). Ukupno: ~23.000 tokena, '
        'od čega samo 4.000 na skupom modelu.')

    add_paragraph(doc, 'Ekonomska kalkulacija', bold=True)

    add_paragraph(doc,
        'Uzmimo reprezentativne cene (jun 2025): jak model ≈ $15/M tokena ulaz, '
        'srednji model ≈ $3/M tokena, slab model ≈ $0.25/M tokena. '
        'Mono-agent: 45K × $15/M = $0.675 po zadatku. '
        'Multi-agent: (4K × $15 + 6.5K × $3 + 12.5K × $0.25) / 1M = $0.060 + $0.020 + $0.003 = $0.083 po zadatku. '
        'Multi-agent pristup je u ovom scenariju ~8× jeftiniji, uz uporediv kvalitet.')

    # Generate and add cost analysis diagram
    cost_img = create_cost_quality_analysis_diagram()
    add_image(doc, cost_img, width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run('Slika 6. Komparativna analiza potrošnje tokena i odnosa cene i kvaliteta')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True

    add_paragraph(doc, 'Analiza kvaliteta', bold=True)

    add_paragraph(doc,
        'Paradoksalno, multi-agent pristup sa slabijim radnicima često postiže kvalitet '
        'uporediv ili čak bolji od mono-agenta sa jakim modelom. Razlozi su sledeći:')

    quality_points = [
        'Fokusirani kontekst — radnik dobija samo relevantne informacije za svoj podzadatak, '
        'bez šuma celokupnog konteksta. Kraći kontekst = manje halucinacija i "izgubljene niti".',
        'Specijalizacija promptova — svaki radnik ima prompt optimizovan za jednu vrstu zadatka, '
        'što poboljšava kvalitet u odnosu na generički prompt koji pokriva sve.',
        'Nezavisna verifikacija — tester objektivno evaluira rezultat, dok mono-agent sam sebi '
        'proverava rad (eng. self-bias problem).',
        'Izolacija grešaka — greška jednog radnika ne kontaminira kontekst ostalih. '
        'U mono-agentu, jedna greška u koraku 2 može propagirati konfuziju kroz korake 3-5.',
        'Iterativna korekcija — feedback petlja menadžer-radnik-tester konvergira brže '
        'nego self-correction kod mono-agenta jer koristi različite "perspektive".',
    ]

    for point in quality_points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(point)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    add_paragraph(doc,
        'Na slici 6 (desno) prikazan je odnos relativne cene i kvaliteta za četiri konfiguracije. '
        'Multi-agent pristup sa raslojavanjem modela zauzima optimalni kvadrant (nizak trošak, '
        'visok kvalitet). Mono-agent sa jakim modelom postiže sličan kvalitet ali po znatno višoj '
        'ceni. Korišćenje slabog modela kao mono-agenta je jeftino ali nedovoljno kvalitetno za '
        'složene zadatke. Multi-agent sa svim jakim modelima daje marginalno bolji kvalitet, ali '
        'uz neopravdano visok trošak.')

    add_paragraph(doc, 'Kada koristiti koji pristup?', bold=True)

    add_paragraph(doc,
        'Multi-agent pristup je superioran za: zadatke koji se prirodno dekompozuju na nezavisne '
        'delove, projekte gde je budžet ograničen, scenarije gde je verifikacija kritična, '
        'i dugoročne agentske sesije gde bi akumulacija konteksta kod mono-agenta dovela do '
        'degradacije kvaliteta. Mono-agent je bolji izbor za: jednostavne zadatke koji ne '
        'zahtevaju dekompoziciju, scenarije gde je latencija kritičnija od cene, i zadatke '
        'koji zahtevaju duboko razumevanje celokupnog konteksta bez fragmentacije.')

    add_page_break(doc)

    # ========== CHAPTER 8: Zaključak ==========
    add_heading(doc, '8. Zaključak', level=1)

    add_paragraph(doc,
        'U ovom radu predstavili smo sveobuhvatan pregled multi-agent sistema zasnovanih na '
        'velikim jezičkim modelima, sa posebnim fokusom na planiranje i paralelno izvršavanje '
        'zadataka. Analizirali smo kako LLM-ovi služe kao kognitivno jezgro agenata, omogućavajući '
        'im sposobnosti rasuđivanja, planiranja i komunikacije na nivou prirodnog jezika.')

    add_paragraph(doc,
        'Hijerarhijska arhitektura Planer-Menadžer-Radnik-Tester demonstrira efikasan pristup '
        'organizaciji multi-agent sistema koji kombinuje jasnu podelu odgovornosti sa fleksibilnošću '
        'paralelnog izvršavanja. Iterativni ciklus verifikacije i korekcije obezbeđuje konvergenciju '
        'ka kvalitetnim rešenjima.')

    add_paragraph(doc,
        'Budući pravci istraživanja uključuju: unapređenje mehanizama koordinacije za smanjenje '
        'komunikacionog overhead-a, razvoj efikasnijih strategija dekompozicije zadataka, '
        'integraciju naprednih mehanizama memorije za dugoročno učenje, i formalizaciju '
        'protokola za multi-agent interakciju u heterogenim okruženjima.')

    add_paragraph(doc,
        'Multi-agent sistemi zasnovani na LLM-u predstavljaju jedan od najdinamičnijih pravaca '
        'u razvoju veštačke inteligencije, sa potencijalom da transformišu način na koji se '
        'pristupa rešavanju složenih problema u softverskom inženjerstvu i šire.')

    add_page_break(doc)

    # ========== LITERATURA (References at the end) ==========
    add_heading(doc, 'Literatura', level=1)

    references = [
        '[1] Wang, L., Ma, C., Feng, X., et al. (2024). "A Survey on Large Language Model based Autonomous Agents." Frontiers of Computer Science, 18(6), 186345.',
        '[2] Xi, Z., Chen, W., Guo, X., et al. (2023). "The Rise and Potential of Large Language Model Based Agents: A Survey." arXiv preprint arXiv:2309.07864.',
        '[3] Wu, Q., Bansal, G., Zhang, J., et al. (2023). "AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation." arXiv preprint arXiv:2308.08155.',
        '[4] Hong, S., Zhuge, M., Chen, J., et al. (2023). "MetaGPT: Meta Programming for A Multi-Agent Collaborative Framework." arXiv preprint arXiv:2308.00352.',
        '[5] Park, J.S., O\'Brien, J.C., Cai, C.J., et al. (2023). "Generative Agents: Interactive Simulacra of Human Behavior." UIST 2023.',
        '[6] Yao, S., Zhao, J., Yu, D., et al. (2023). "ReAct: Synergizing Reasoning and Acting in Language Models." ICLR 2023.',
        '[7] Wei, J., Wang, X., Schuurmans, D., et al. (2022). "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models." NeurIPS 2022.',
        '[8] Shinn, N., Cassano, F., Gopinath, A., et al. (2023). "Reflexion: Language Agents with Verbal Reinforcement Learning." NeurIPS 2023.',
        '[9] Li, G., Hammoud, H.A.A.K., Itani, H., et al. (2023). "CAMEL: Communicative Agents for Mind Exploration of Large Language Model Society." NeurIPS 2023.',
        '[10] Talebirad, Y., Nadiri, A. (2023). "Multi-Agent Collaboration: Harnessing the Power of Intelligent LLM Agents." arXiv preprint arXiv:2306.03314.',
        '[11] Shen, Y., Song, K., Tan, X., et al. (2024). "HuggingGPT: Solving AI Tasks with ChatGPT and its Friends in Hugging Face." NeurIPS 2023.',
        '[12] Brown, T., Mann, B., Ryder, N., et al. (2020). "Language Models are Few-Shot Learners." NeurIPS 2020.',
        '[13] Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). "Attention Is All You Need." NeurIPS 2017.',
        '[14] OpenAI. (2023). "GPT-4 Technical Report." arXiv preprint arXiv:2303.08774.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ========== Save ==========
    output_path = os.path.join(OUTPUT_DIR, "LLM_Agenti_Planiranje_Zadataka_v5.docx")
    doc.save(output_path)
    print(f"Dokument sačuvan: {output_path}")
    return output_path


if __name__ == '__main__':
    build_document()
    print("Generisanje završeno!")
