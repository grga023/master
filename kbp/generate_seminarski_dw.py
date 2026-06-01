from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / 'Seminarski_rad_DW_ITKompanije.docx'

IMAGES = {
    'cdm': BASE_DIR / 'diagrami' / 'slika 1.png',
    'pdm': BASE_DIR / 'diagrami' / 'slika2 relation.png',
    'db': BASE_DIR / 'diagrami' / 'diag baze pod.png',
    'view_trosak': BASE_DIR / 'diagrami' / 'anal troskova rada.png',
    'view_prihod': BASE_DIR / 'diagrami' / 'analiza prihoda.png',
    'view_tehnologija': BASE_DIR / 'diagrami' / 'analiza tehnologija.png',
    'select': BASE_DIR / 'diagrami' / 'korisniSelectModeli.png',
}

DEPT_DATA = [
    ['Design', '16', '416', '78'],
    ['Development', '177', '4922', '664'],
    ['Management', '22', '770', '175'],
    ['QA', '52', '1276', '172'],
]

PROJECT_PRODUCTIVITY = [
    ['Analytics Dashboard', '38', '1070'],
    ['Cloud Migration', '49', '1509'],
    ['CRM sistem', '57', '1681'],
    ['E-Commerce platforma', '67', '1796'],
    ['Mobile Banking App', '56', '1328'],
]

EMPLOYEE_PRODUCTIVITY = [
    ['Ana Stojanović', '42', '1050', '6'],
    ['Ivan Marković', '30', '660', '4'],
    ['Jelena Nikolić', '16', '416', '3'],
    ['Maja Ilić', '22', '616', '3'],
    ['Marko Petrović', '54', '1620', '7'],
    ['Milica Đorđević', '21', '378', '3'],
    ['Nikola Todorović', '37', '1184', '5'],
    ['Petar Pavlović', '23', '690', '3'],
    ['Stefan Jovanović', '22', '770', '5'],
]

CLIENT_REVENUE = [
    ['AppDev Inc', '10200', '340', '2'],
    ['CloudNet d.o.o.', '18000', '600', '3'],
    ['DataSys GmbH', '27300', '910', '5'],
    ['TechCorp Solutions', '23650', '885', '5'],
    ['WebPro Ltd', '26550', '885', '5'],
]

PROJECT_REVENUE = [
    ['Analytics Dashboard', '10200', '120000', '2'],
    ['Cloud Migration', '18000', '285000', '3'],
    ['CRM sistem', '27300', '600000', '5'],
    ['E-Commerce platforma', '23650', '425000', '5'],
    ['Mobile Banking App', '26550', '750000', '5'],
]

YEAR_REVENUE = [
    ['2024', '77500', '15'],
    ['2025', '28200', '5'],
]

QUARTER_REVENUE = [
    ['Q1', '48300', '1675'],
    ['Q2', '31150', '1070'],
    ['Q3', '21300', '710'],
    ['Q4', '4950', '165'],
]

LITERATURE = [
    'Inmon, W. H. (2005). Building the Data Warehouse. Wiley.',
    'Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling. Wiley.',
    'Ponniah, P. (2010). Data Warehousing Fundamentals for IT Professionals. Wiley.',
    'Chaudhuri, S., & Dayal, U. (1997). An Overview of Data Warehousing and OLAP Technology. ACM SIGMOD Record, 26(1), 65-74.',
    'Harrington, J. L. (2008). Relational Database Design and Implementation. Morgan Kaufmann.',
    'Microsoft. (2024). SQL Server Analysis Services Documentation. Microsoft Learn.',
    'Golfarelli, M., & Rizzi, S. (2009). Data Warehouse Design: Modern Principles and Methodologies. McGraw-Hill.',
    'Elmasri, R., & Navathe, S. B. (2016). Fundamentals of Database Systems. Pearson.',
    'Vassiliadis, P. (2010). A Survey of Extract-Transform-Load Technology. International Journal of Data Warehousing and Mining, 5(3), 1-27.',
]


APSTRAKT = (
    'U ovom radu prikazano je projektovanje i implementacija skladišta podataka namenjenog analizi poslovanja IT kompanije koja istovremeno upravlja zaposlenima, projektima, klijentima, tehnologijama i procesom fakturisanja. Polazna tačka rada jeste relaciona baza podataka sa osam međusobno povezanih tabela, nad kojom su definisani analitički pogledi i dve OLAP kocke. Posebna pažnja posvećena je razdvajanju transakcionog i analitičkog sloja, izboru odgovarajućih dimenzija i mera, kao i oblikovanju izveštaja koji omogućavaju pouzdano donošenje poslovnih odluka. U radu su analizirani troškovi rada po odeljenjima, produktivnost po projektima i zaposlenima, kao i prihodi po klijentima, projektima, godinama i kvartalima. Dobijeni rezultati ukazuju da skladište podataka obezbeđuje jedinstven, konzistentan i vremenski orijentisan pogled na poslovne performanse kompanije. Time se menadžmentu omogućava da preciznije prati profitabilnost, opterećenje resursa i dinamiku fakturisanja, uz osnovu za buduće proširenje sistema ka naprednim prediktivnim analizama i automatizovanom izveštavanju.'
)

KEYWORDS = 'Ključne reči: skladište podataka, OLAP, dimenzionalno modeliranje, SQL Server, analiza produktivnosti, analiza prihoda.'


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.bold = bold


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, space_after=Pt(6), first_line=0.75):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    p.paragraph_format.first_line_indent = Cm(first_line) if first_line else Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.bold = True
    return p


def add_caption(doc, text, kind='Slika'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    caption_text = text if text.startswith(kind) else f'{kind} {text}'
    run = p.add_run(caption_text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.italic = True


def add_image(doc, image_path, caption, analysis_paragraphs, width=6.2):
    if not image_path.exists():
        raise FileNotFoundError(f'Nedostaje slika: {image_path}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption, kind='Slika')
    for text in analysis_paragraphs:
        add_paragraph(doc, text)


def add_table(doc, headers, rows, caption, analysis_paragraphs):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, bold=False, align=align)
    add_caption(doc, caption, kind='Tabela')
    for text in analysis_paragraphs:
        add_paragraph(doc, text)


def set_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for line in [
        'UNIVERZITET U NOVOM SADU',
        'Tehnički fakultet "Mihajlo Pupin" Zrenjanin',
        'Master akademske studije',
    ]:
        run = p.add_run(line + '\n')
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        run.bold = True

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    r1 = title.add_run('SEMINARSKI RAD\n')
    r1.font.name = 'Times New Roman'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r1.font.size = Pt(16)
    r1.bold = True
    r2 = title.add_run('Projektovanje skladišta podataka i OLAP analize za IT kompaniju')
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2.font.size = Pt(16)
    r2.bold = True

    for _ in range(8):
        doc.add_paragraph()

    info_lines = [
        'Student: Ognjen Grgur',
        'Broj indeksa: MIT 37/24',
        'Predmet: Koncepti baza podataka',
    ]
    for line in info_lines:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_info.paragraph_format.line_spacing = 1.5
        run = p_info.add_run(line)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        run.bold = True

    p_city = doc.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_city.paragraph_format.space_before = Pt(36)
    p_city.paragraph_format.line_spacing = 1.5
    run = p_city.add_run('Zrenjanin, 2025.')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    run.bold = True

    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, 'Sadržaj')
    contents = [
        '1. Uvod',
        '2. Teorijske osnove skladišta podataka',
        '2.1 Definicija i karakteristike skladišta podataka',
        '2.2 OLTP i OLAP sistemi',
        '2.3 Dimenzionalno modeliranje',
        '3. Projektovanje baze podataka',
        '3.1 Opis poslovnog domena',
        '3.2 Konceptualni model',
        '3.3 Fizički model',
        '3.4 Implementacija u SQL Server okruženju',
        '3.5 Pogledi za analizu',
        '4. Projektovanje OLAP kocke',
        '4.1 Dimenzije',
        '4.2 Mere',
        '4.3 Šema zvezde',
        '4.4 Kreiranje i procesiranje kocke',
        '5. Analiza podataka',
        '5.1 Analiza produktivnosti',
        '5.2 Analiza prihoda',
        '6. Zaključak',
        '7. Literatura',
    ]
    for item in contents:
        add_paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0), first_line=0)
    doc.add_page_break()


def build_document():
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc)

    add_heading(doc, 'Apstrakt')
    add_paragraph(doc, APSTRAKT)
    add_paragraph(doc, KEYWORDS, italic=True)
    doc.add_page_break()

    add_contents(doc)

    add_heading(doc, '1. Uvod')
    uvod_paragraphs = [
        'Savremeno poslovanje IT kompanija zasniva se na kontinuiranom stvaranju podataka o zaposlenima, projektima, klijentima, radnim satima i fakturisanim uslugama. U operativnim sistemima ovi podaci nastaju kroz veliki broj pojedinačnih transakcija, ali njihova neposredna upotreba za strateško odlučivanje često nije dovoljna. Upravo iz tog razloga skladište podataka predstavlja ključnu komponentu savremenog informacionog sistema, jer omogućava objedinjavanje podataka iz različitih izvora, njihovu istorijsku stabilizaciju i analitičku obradu u skladu sa poslovnim ciljevima.',
        'Ukoliko menadžment kompanije želi da proceni produktivnost pojedinih odeljenja, profitabilnost projekata ili dinamiku prihoda po klijentima, potrebno je raspolagati tačnim i konzistentnim pokazateljima. Operativna baza podataka jeste pogodna za unos i ažuriranje podataka, ali nije optimizovana za složene upite nad velikim skupovima istorijskih zapisa. Skladište podataka rešava navedeni problem tako što organizuje podatke u analitički orijentisanu strukturu, prilagođenu agregaciji, upoređivanju i praćenju trendova kroz vreme.',
        'Predmet ovog rada jeste projektovanje modela skladišta podataka za IT kompaniju koja posluje kroz više odeljenja i realizuje projekte za različite klijente. Polazni model obuhvata osam relacionih tabela: Odeljenje, Zaposleni, Klijent, Projekat, Tehnologija, ProjekatTehnologija, EvidencijaSati i Faktura. Nad ovim strukturama formirani su pogledi za analizu, a zatim i OLAP kocke koje omogućavaju višedimenzionalno posmatranje podataka prema zaposlenom, odeljenju, projektu, klijentu i vremenu.',
        'Cilj rada jeste da se prikaže celovit postupak prelaska od operativnog modela ka analitičkom sistemu, uz obrazloženje teorijskih osnova, modelovanja baze podataka, definisanja dimenzija i mera, kao i tumačenja rezultata dobijenih iz OLAP analiza. Posebno su razmatrani troškovi rada, produktivnost zaposlenih i struktura prihoda, jer upravo navedeni pokazatelji imaju neposredan uticaj na planiranje resursa, kontrolu budžeta i vrednovanje poslovnog učinka.',
        'Doprinos rada ogleda se u tome što na konkretnom primeru pokazuje kako dobro projektovano skladište podataka može da poveže operativne procese i menadžersko izveštavanje. Takav pristup omogućava donošenje odluka zasnovanih na podacima, smanjuje rizik od subjektivnih procena i stvara pouzdanu osnovu za buduća proširenja u oblasti poslovne inteligencije, prediktivnih modela i automatizovanog praćenja performansi.'
    ]
    for paragraph in uvod_paragraphs:
        add_paragraph(doc, paragraph)

    add_heading(doc, '2. Teorijske osnove skladišta podataka')
    add_heading(doc, '2.1 Definicija i karakteristike skladišta podataka', level=2)
    teorija_1 = [
        'Skladište podataka predstavlja subjektno orijentisanu, integrisanu, vremenski promenljivu i nepromenljivu kolekciju podataka namenjenu podršci odlučivanju. Za razliku od operativnih baza, koje su usmerene na tekuće poslovne transakcije, skladište podataka objedinjuje informacije iz više izvora i priprema ih za analitičku upotrebu. Takva organizacija podataka omogućava da se poslovni procesi posmatraju kroz duži vremenski period i na različitim nivoima agregacije.',
        'Osnovna karakteristika skladišta podataka jeste integracija heterogenih podataka. U praksi to znači da podaci o zaposlenima, projektima, troškovima, tehnologijama i fakturama mogu poticati iz različitih aplikacija, ali se u analitičkom sistemu prikazuju kroz jedinstven model i jedinstvenu poslovnu terminologiju. Time se smanjuje mogućnost nastanka nedoslednosti i obezbeđuje da svi korisnici donose zaključke na osnovu istih pokazatelja.',
        'Vremenska komponenta skladišta podataka od posebnog je značaja, jer se poslovna analiza ne zasniva samo na trenutnom stanju već i na praćenju promena kroz mesece, kvartale i godine. U analitičkom okruženju istorijski zapisi imaju visoku vrednost, pošto omogućavaju uočavanje trendova, sezonskih odstupanja i posledica ranijih poslovnih odluka. Neizmenljivost podataka dodatno doprinosi pouzdanosti izveštavanja, budući da se već obrađeni i konsolidovani podaci ne menjaju na način karakterističan za operativne sisteme.'
    ]
    for paragraph in teorija_1:
        add_paragraph(doc, paragraph)

    add_heading(doc, '2.2 OLTP i OLAP sistemi', level=2)
    teorija_2 = [
        'OLTP sistemi namenjeni su podršci svakodnevnim poslovnim transakcijama, kao što su unos radnih sati, evidentiranje projekata ili izdavanje faktura. Njihov prioritet jeste brzina i pouzdanost obrade pojedinačnih događaja, uz očuvanje integriteta podataka. Nasuprot tome, OLAP sistemi orijentisani su na kompleksne upite, agregacije i poređenja većeg broja dimenzija radi analize poslovnih performansi.',
        'U kontekstu posmatrane IT kompanije, OLTP sloj obezbeđuje evidenciju zaposlenih, projekata i faktura, dok OLAP sloj pruža odgovore na pitanja koja prevazilaze operativni nivo, kao što su: koje odeljenje ostvaruje najveći trošak rada, koji projekti donose najveći prihod ili kako se prihodi menjaju po kvartalima. Upravo ta razlika u nameni određuje i različite modele podataka, strategije indeksiranja i načine pristupa informacijama.'
    ]
    for paragraph in teorija_2:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        ['Karakteristika', 'OLTP sistem', 'OLAP sistem'],
        [
            ['Primarna svrha', 'Operativna obrada transakcija', 'Analiza i podrška odlučivanju'],
            ['Struktura podataka', 'Visoko normalizovana', 'Dimenzionalno organizovana'],
            ['Tip upita', 'Kratki i brojni', 'Složeni i agregirani'],
            ['Vremenska orijentacija', 'Aktuelno stanje', 'Istorijski i trendovski pogled'],
            ['Korisnici', 'Operateri i administracija', 'Menadžment i analitičari'],
            ['Performanse', 'Optimizovane za unos i izmene', 'Optimizovane za čitanje i analizu'],
        ],
        'Tabela 1. Poređenje OLTP i OLAP sistema',
        [
            'Prikazano poređenje potvrđuje da OLTP i OLAP sistemi imaju različite, ali komplementarne uloge u informacionoj arhitekturi preduzeća. Operativni nivo obezbeđuje tačne i pravovremene podatke, dok analitički nivo omogućava njihovo tumačenje u širem poslovnom kontekstu.',
            'Za potrebe ovog rada bilo je neophodno razdvojiti transakcioni i analitički sloj, jer se jedino na taj način mogu istovremeno ostvariti pouzdan unos podataka i efikasna višedimenzionalna analiza.'
        ]
    )

    add_heading(doc, '2.3 Dimenzionalno modeliranje', level=2)
    teorija_3 = [
        'Dimenzionalno modeliranje predstavlja najzastupljeniji pristup organizovanju podataka u skladištu podataka. Njegova osnovna ideja jeste da se podaci strukturiraju oko centralnih tabela činjenica i povezanih tabela dimenzija. Tabela činjenica sadrži merljive poslovne događaje, kao što su broj sati, trošak rada ili iznos fakture, dok dimenzije daju kontekst tim merama kroz podatke o zaposlenima, projektima, klijentima, odeljenjima i vremenu.',
        'Najčešće korišćena dimenzionalna rešenja jesu šema zvezde i šema pahulje. Šema zvezde podrazumeva direktno povezivanje centralne tabele činjenica sa dimenzijama, čime se postiže jednostavnije razumevanje modela i brže izvršavanje analitičkih upita. Šema pahulje dodatno normalizuje pojedine dimenzije, čime se smanjuje redundansa, ali se povećava složenost modela i upita.',
        'U ovom radu dominantno je primenjen pristup blizak šemi zvezde, pošto je osnovni cilj bio formiranje jasnog analitičkog modela pogodnog za OLAP obradu. Takav pristup obezbeđuje preglednost, jednostavno definisanje mera i dimenzija, kao i efikasno izvođenje agregacija nad podacima o produktivnosti i prihodima.'
    ]
    for paragraph in teorija_3:
        add_paragraph(doc, paragraph)

    add_heading(doc, '3. Projektovanje baze podataka')
    add_heading(doc, '3.1 Opis poslovnog domena', level=2)
    domen = [
        'Poslovni domen rada zasniva se na modelu IT kompanije koja pruža usluge razvoja softvera za više klijenata istovremeno. Organizacija poseduje više odeljenja, među kojima se izdvajaju Development, QA, Design i Management, a svaki zaposleni pripada jednom odeljenju, ima definisanu poziciju, datum zaposlenja i satnicu. Poslovni proces započinje angažovanjem zaposlenih na projektima, nastavlja se evidentiranjem radnih sati i završava se fakturisanjem izvršenih usluga klijentima.',
        'Ključna specifičnost domena jeste potreba da se povežu resursi, vreme rada, tehnološki okvir i finansijski rezultati. Projekti se izvode za konkretne klijente, imaju budžet, vremenski okvir i status realizacije, dok se kroz vezu ProjekatTehnologija evidentira koje su tehnologije korišćene tokom implementacije. Na taj način moguće je sagledati ne samo troškovnu i prihodnu stranu poslovanja, već i tehnološku strukturu svakog projekta.',
        'Analitička vrednost ovakvog domena posebno dolazi do izražaja kada se podaci posmatraju kroz duži period. Menadžment može da proceni da li je određeni klijent stabilan izvor prihoda, koje odeljenje generiše najveće troškove rada, kao i da li postoji nesklad između obima angažovanja i iznosa ostvarenih prihoda. Zbog toga je poslovni domen pogodan za primenu skladišta podataka i OLAP tehnologije.'
    ]
    for paragraph in domen:
        add_paragraph(doc, paragraph)

    add_heading(doc, '3.2 Konceptualni model', level=2)
    add_paragraph(doc, 'Konceptualni model prikazuje osnovne entitete poslovnog sistema i logičke veze među njima, bez ulaska u tehničke detalje fizičke implementacije. Fokus konceptualnog nivoa jeste razumevanje poslovnih pojmova i relacija koje čine jezgro analitičkog rešenja.')
    add_image(
        doc,
        IMAGES['cdm'],
        'Slika 1. Konceptualni model sistema (CDM)',
        [
            'Konceptualni model jasno razdvaja ključne entitete poslovnog sistema: zaposlene, odeljenja, klijente, projekte, tehnologije, evidenciju rada i fakture. Na ovom nivou posebno je važno to što su uočene veze koje omogućavaju kasnije izvođenje i troškovnih i prihodnih analiza u jedinstvenom modelu.',
            'Struktura modela potvrđuje da je poslovni domen dovoljno stabilan za prelazak ka detaljnijem relacijskom i analitičkom modelovanju. Uočava se i centralna uloga entiteta Projekat, jer upravo on povezuje operativne i finansijske informacije.'
        ]
    )

    add_heading(doc, '3.3 Fizički model', level=2)
    add_paragraph(doc, 'Fizički model prevodi konceptualna rešenja u konkretne tabele, atribute, primarne i strane ključeve, uz pravila integriteta koja omogućavaju implementaciju u sistemu za upravljanje bazom podataka. Na ovom nivou određuje se i način povezivanja podataka koji će kasnije biti korišćeni u analitičkim pogledima i OLAP kockama.')
    add_image(
        doc,
        IMAGES['pdm'],
        'Slika 2. Fizički model baze podataka (PDM)',
        [
            'Fizički model potvrđuje da je relacijska struktura projektovana konzistentno, sa jasno definisanim primarnim i stranim ključevima. Posebno je značajna tabela EvidencijaSati, jer predstavlja najvažniji izvor podataka za obračun troškova rada i analizu produktivnosti.',
            'Takođe je uočljivo da tabela Faktura uvodi finansijsku perspektivu sistema, čime je omogućeno povezivanje utrošenog rada sa ostvarenim prihodima. Time je stvorena osnova za izgradnju dve odvojene, ali međusobno povezane OLAP kocke.'
        ]
    )

    add_heading(doc, '3.4 Implementacija u SQL Server okruženju', level=2)
    impl = [
        'Implementacija modela realizovana je u SQL Server okruženju, gde su definisane relacione tabele, ograničenja integriteta i pomoćni upiti za proveru i pripremu podataka. Takvo okruženje pogodno je za kombinovanje operativnog dela sistema sa naknadnim analitičkim proširenjem putem SQL Server Analysis Services alata.',
        'Baza podataka sadrži osam tabela: Odeljenje, Zaposleni, Klijent, Projekat, Tehnologija, ProjekatTehnologija, EvidencijaSati i Faktura. Ovakva struktura omogućava evidentiranje radnih angažovanja na nivou pojedinačnih zaposlenih i projekata, kao i praćenje fakturisanih iznosa kroz vremenske dimenzije meseca i godine.',
        'Posebna vrednost implementacije ogleda se u činjenici da su relacione strukture projektovane tako da se bez dodatnih transformacija mogu iskoristiti za kreiranje analitičkih pogleda i kocki. Time se smanjuje kompleksnost održavanja i povećava pouzdanost celokupnog rešenja.'
    ]
    for paragraph in impl:
        add_paragraph(doc, paragraph)

    add_image(
        doc,
        IMAGES['db'],
        'Slika 3. Dijagram implementirane baze podataka u SQL Server okruženju',
        [
            'Dijagram baze potvrđuje da je implementacija dosledno izvedena u skladu sa relacijskim pravilima i definisanim poslovnim zahtevima. Vizuelni prikaz veza olakšava razumevanje tokova podataka od operativne evidencije ka analitičkim izveštajima.',
            'Na osnovu ovakve strukture moguće je precizno definisati JOIN operacije koje stoje u osnovi pogleda i OLAP analiza. Time se obezbeđuje transparentnost izvora podataka za svaku meru koja se prikazuje u izveštajima.'
        ]
    )

    add_image(
        doc,
        IMAGES['select'],
        'Slika 4. Pomoćni SELECT upiti za proveru i pregled podataka',
        [
            'Pomoćni SELECT upiti imaju važnu ulogu u validaciji poslovnih pravila i proveri kvaliteta unetih podataka. Njihovom primenom moguće je brzo potvrditi da su veze između zaposlenih, projekata, klijenata i evidencija rada pravilno uspostavljene.',
            'Ovakvi upiti predstavljaju neizostavan korak između faze implementacije i faze analitičke obrade, jer omogućavaju rano otkrivanje nelogičnosti i obezbeđuju pouzdaniju osnovu za kasnije agregacije.'
        ]
    )

    add_heading(doc, '3.5 Pogledi za analizu', level=2)
    views_intro = [
        'Radi pojednostavljenja analitičkog pristupa nad relacijskim modelom formirana su tri pogleda: Analiza_troska_rada, Analiza_prihoda i Analiza_tehnologija. Njihova uloga jeste da objedine podatke iz više tabela i pripreme ih za doslednu upotrebu u izveštajima i OLAP okruženju.',
        'Pogled Analiza_troska_rada povezuje podatke iz tabela EvidencijaSati, Zaposleni, Projekat, Klijent i Odeljenje, pri čemu se dodatno izračunava mera TrošakRada kao proizvod broja sati i satnice. Pogled Analiza_prihoda povezuje fakture sa projektima i klijentima, uz izvedenu vrednost kvartala, dok pogled Analiza_tehnologija prikazuje tehnološku strukturu projekata po klijentima.'
    ]
    for paragraph in views_intro:
        add_paragraph(doc, paragraph)

    add_image(
        doc,
        IMAGES['view_trosak'],
        'Slika 5. Rezultat pogleda Analiza_troska_rada',
        [
            'Prikaz pogleda Analiza_troska_rada potvrđuje da su troškovi rada uspešno izvedeni iz operativnih evidencija i podataka o satnicama zaposlenih. Time je obezbeđen jasan most između evidencije utrošenog vremena i finansijskog vrednovanja angažovanja.',
            'Za potrebe menadžerskog izveštavanja ovaj pogled ima naročit značaj, jer omogućava poređenje troškova po odeljenjima, projektima i klijentima bez potrebe za ponavljanjem složenih JOIN operacija u svakom pojedinačnom upitu.'
        ]
    )

    add_image(
        doc,
        IMAGES['view_prihod'],
        'Slika 6. Rezultat pogleda Analiza_prihoda',
        [
            'Pogled Analiza_prihoda integriše podatke o fakturama, projektima i klijentima, čime se prihodna strana poslovanja posmatra kroz jedinstven analitički okvir. Uvođenje kvartala kao izvedene vremenske dimenzije dodatno unapređuje mogućnost praćenja sezonskih obrazaca naplate.',
            'Ovakav pogled omogućava da se isti skup podataka koristi i za klasične SQL izveštaje i za OLAP kocku, što povećava konzistentnost tumačenja rezultata i smanjuje rizik od različitih verzija istine u sistemu.'
        ]
    )

    add_image(
        doc,
        IMAGES['view_tehnologija'],
        'Slika 7. Rezultat pogleda Analiza_tehnologija',
        [
            'Pogled Analiza_tehnologija pruža uvid u raspodelu tehnologija po projektima i klijentima, što je od posebne važnosti za tehnološko planiranje i procenu kompetencija organizacije. Na osnovu ovakvih podataka moguće je identifikovati dominantne tehnološke pravce i stepen standardizacije projektnog portfolija.',
            'Analiza tehnologija ima i strateški značaj, jer ukazuje na to gde je potrebno dalje usavršavanje zaposlenih, a gde organizacija već poseduje stabilnu i tržišno relevantnu ekspertizu.'
        ]
    )

    add_heading(doc, '4. Projektovanje OLAP kocke')
    add_heading(doc, '4.1 Dimenzije', level=2)
    olap_dim = [
        'U okviru rada definisane su dve OLAP kocke: cbProduktivnost i cbPrihodi. Kocka cbProduktivnost obuhvata dimenzije Zaposleni, Odeljenje, Projekat, Klijent, Godina i Kvartal, dok kocka cbPrihodi koristi dimenzije Projekat, Klijent, Godina i Kvartal. Izbor dimenzija izvršen je u skladu sa pitanjima na koja menadžment želi da dobije odgovor, pri čemu svaka dimenzija predstavlja jedan ugao posmatranja poslovnih performansi.',
        'Dimenzija Zaposleni omogućava analizu individualnog angažovanja i troškova rada, dimenzija Odeljenje agregira učinak na organizacionom nivou, dok dimenzije Projekat i Klijent pružaju pogled na tržišnu i komercijalnu stranu poslovanja. Vremenske dimenzije Godina i Kvartal od posebne su važnosti zato što omogućavaju praćenje dinamike promena i sezonskih oscilacija u produktivnosti i prihodima.'
    ]
    for paragraph in olap_dim:
        add_paragraph(doc, paragraph)

    add_heading(doc, '4.2 Mere', level=2)
    olap_measures = [
        'U kocki cbProduktivnost definisane su mere Broj Sati, Trošak Rada, Satnica i Broj Evidencija. Mere Broj Sati i Trošak Rada agregiraju se funkcijom SUM, dok se Broj Evidencija dobija prebrojavanjem zapisa. Agregirana mera Satnica u ovom modelu predstavlja zbir vrednosti satnica na nivou odabranog preseka i koristi se kao dopunski indikator strukture angažovanja.',
        'U kocki cbPrihodi definisane su mere Iznos, Ukupno Sati, Budžet i Broj Faktura. Njihovom kombinovanom analizom moguće je sagledati ne samo ostvaren prihod, već i odnos između naplate, planiranog budžeta i radnog opterećenja koje stoji iza fakturisanih usluga.'
    ]
    for paragraph in olap_measures:
        add_paragraph(doc, paragraph)

    add_heading(doc, '4.3 Šema zvezde', level=2)
    olap_star = [
        'Analitički model ovog rada organizovan je prema principima šeme zvezde. U središtu modela nalaze se činjenice o evidenciji rada i prihodima, dok se oko njih raspoređuju dimenzije koje obezbeđuju poslovni kontekst. Takva organizacija pogodna je za OLAP upite, jer pojednostavljuje logiku povezivanja podataka i ubrzava agregacije po više kriterijuma.',
        'Za kocku cbProduktivnost tabela činjenica praktično nastaje iz pogleda Analiza_troska_rada, dok se oko nje vezuju dimenzije zaposlenog, odeljenja, projekta, klijenta i vremena. Za kocku cbPrihodi centralnu ulogu ima pogled Analiza_prihoda, sa dimenzijama projekta, klijenta i vremena. Ovakav pristup omogućava jednostavno „bušenje“ podataka od ukupnog nivoa ka detaljnom zapisu i obrnuto.'
    ]
    for paragraph in olap_star:
        add_paragraph(doc, paragraph)

    add_heading(doc, '4.4 Kreiranje i procesiranje kocke', level=2)
    olap_process = [
        'Kreiranje OLAP kocke podrazumeva definisanje izvora podataka, data source view sloja, dimenzija, hijerarhija i mera, a zatim i samu konstrukciju kocke u alatu SQL Server Analysis Services. Nakon logičkog definisanja modela, neophodno je izvršiti procesiranje kocke, čime se podaci iz relacione baze učitavaju u multidimenzionalnu strukturu pogodnu za brzo izvršavanje analitičkih upita.',
        'Procesiranjem kocke omogućeno je generisanje pivot analiza nad rezultatima, pri čemu korisnik može da menja perspektivu posmatranja bez izmene same baze podataka. Ovaj korak predstavlja završnu tačku transformacije relacijskih podataka u poslovnu inteligenciju, jer od tog trenutka sistem pruža interaktivne, konzistentne i brzo dostupne uvide.'
    ]
    for paragraph in olap_process:
        add_paragraph(doc, paragraph)

    add_heading(doc, '5. Analiza podataka')
    add_heading(doc, '5.1 Analiza produktivnosti', level=2)
    produkt_intro = [
        'Analiza produktivnosti zasniva se na podacima iz kocke cbProduktivnost i omogućava sagledavanje radnog angažovanja sa više aspekata: po odeljenjima, projektima i zaposlenima. Na osnovu ovih pokazatelja moguće je oceniti raspodelu posla, identifikovati nosioce najvećeg opterećenja i uočiti eventualne neravnoteže u korišćenju resursa.',
        'Ukupan analizirani fond iznosi 267 radnih sati i 7.384 jedinice troška rada, raspoređenih kroz 39 pojedinačnih evidencija. Ovakav obim podataka dovoljan je za ilustraciju načina na koji OLAP pristup omogućava brzo poređenje različitih preseka bez dodatnih ručnih obračuna.'
    ]
    for paragraph in produkt_intro:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        ['Odeljenje', 'Broj sati', 'Trošak rada', 'Satnica'],
        DEPT_DATA,
        'Tabela 2. Produktivnost po odeljenjima',
        [
            'Najveće opterećenje evidentirano je u odeljenju Development sa 177 sati i troškom rada od 4.922, što potvrđuje njegovu centralnu ulogu u realizaciji projekata. QA odeljenje zauzima drugo mesto po obimu angažovanja, dok su Design i Management uključeni u manjem obimu, ali sa jasno vidljivim doprinosom ukupnom procesu rada.',
            'Agregirana mera Satnica pokazuje da se najveći zbir vrednosti satnica takođe vezuje za Development, što je očekivano usled najvećeg broja angažovanih resursa. Menadžment može da iskoristi ovakav pregled za planiranje ravnomernije raspodele zadataka i procenu troškovnih efekata po organizacionim celinama.'
        ]
    )

    add_table(
        doc,
        ['Projekat', 'Broj sati', 'Trošak rada'],
        PROJECT_PRODUCTIVITY,
        'Tabela 3. Produktivnost po projektima',
        [
            'Najveći broj sati zabeležen je na projektu E-Commerce platforma, sa 67 sati i troškom rada od 1.796, što ukazuje na visok nivo angažovanja i potencijalnu složenost implementacije. Odmah iza njega nalazi se CRM sistem sa 57 sati, dok projekat Analytics Dashboard ima najniži evidentirani obim rada.',
            'Poređenje troška rada i broja sati pokazuje da opterećenje nije ravnomerno raspoređeno između projekata. Ovakav uvid omogućava preciznije upravljanje portfoliom i pravovremeno prepoznavanje projekata koji zahtevaju dodatnu kontrolu resursa.'
        ]
    )

    add_table(
        doc,
        ['Zaposleni', 'Broj sati', 'Trošak rada', 'Br. evidencija'],
        EMPLOYEE_PRODUCTIVITY,
        'Tabela 4. Produktivnost po zaposlenima',
        [
            'Najveći individualni angažman ostvario je Marko Petrović sa 54 sata, troškom rada od 1.620 i ukupno 7 evidencija, što ga izdvaja kao najopterećenijeg člana analiziranog skupa. Značajan doprinos ostvaruju i Ana Stojanović i Nikola Todorović, dok se kod ostalih zaposlenih uočava umereniji nivo angažovanja.',
            'Broj evidencija pruža dodatnu dimenziju tumačenja, jer pokazuje ne samo ukupan obim rada već i učestalost angažovanja. Na osnovu ovih podataka moguće je identifikovati nosioce ključnih aktivnosti, ali i prepoznati prostor za uravnoteženje opterećenja među zaposlenima.'
        ]
    )

    add_heading(doc, '5.2 Analiza prihoda', level=2)
    prihod_intro = [
        'Analiza prihoda zasniva se na podacima iz kocke cbPrihodi i omogućava procenu finansijskih rezultata prema klijentima, projektima i vremenskim intervalima. Ukupan evidentirani prihod iznosi 105.700, raspoređen kroz 20 faktura, što pruža reprezentativnu osnovu za sagledavanje komercijalne uspešnosti poslovanja.',
        'Kombinovanjem mera Iznos, Ukupno Sati, Budžet i Broj Faktura moguće je posmatrati ne samo nominalnu vrednost naplate, već i stepen iskorišćenja resursa, intenzitet fakturisanja i relativni položaj pojedinih projekata i klijenata u ukupnoj strukturi prihoda.'
    ]
    for paragraph in prihod_intro:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        ['Klijent', 'Iznos', 'Ukupno sati', 'Br. faktura'],
        CLIENT_REVENUE,
        'Tabela 5. Prihodi po klijentima',
        [
            'Najveći prihod ostvaren je od klijenta DataSys GmbH u iznosu od 27.300, dok je odmah iza njega WebPro Ltd sa 26.550. Klijenti AppDev Inc i CloudNet d.o.o. ostvaruju niže apsolutne iznose, ali i sa manjim brojem faktura i manjim fondom sati.',
            'Poređenje iznosa i utrošenih sati ukazuje da prihodi nisu potpuno proporcionalni angažovanju, što može biti posledica različitih cenovnih politika, složenosti projekata ili ugovorenih uslova naplate. Ovakva analiza posebno je korisna za ocenu profitabilnosti klijentskog portfolija.'
        ]
    )

    add_table(
        doc,
        ['Projekat', 'Iznos', 'Budžet', 'Br. faktura'],
        PROJECT_REVENUE,
        'Tabela 6. Prihodi po projektima',
        [
            'Najveći prihod ostvaren je na projektu CRM sistem, sa 27.300, dok se projekat Mobile Banking App izdvaja najvećim planiranim budžetom od 750.000. Ovaj odnos pokazuje da visina budžeta ne mora neposredno da prati trenutno ostvareni prihod, naročito ukoliko projekat još nije dostigao zreliju fazu realizacije.',
            'Analiza po projektima omogućava menadžmentu da razlikuje projekte sa visokim prihodnim potencijalom od projekata koji zahtevaju dodatnu finansijsku kontrolu. Povezivanjem budžeta, broja faktura i prihoda dobija se realnija slika komercijalne uspešnosti portfolija.'
        ]
    )

    add_table(
        doc,
        ['Godina', 'Iznos', 'Br. faktura'],
        YEAR_REVENUE,
        'Tabela 7. Prihodi po godinama',
        [
            'U 2024. godini ostvaren je prihod od 77.500 kroz 15 faktura, dok je u 2025. godini evidentirano 28.200 kroz 5 faktura. Ovakav odnos ukazuje da je najveći deo analiziranog poslovnog rezultata koncentrisan u 2024. godini.',
            'Vremenska analiza na godišnjem nivou omogućava sagledavanje kontinuiteta naplate i procenu stabilnosti poslovanja. Ukoliko se ovakav obrazac potvrdi i u narednim periodima, menadžment može preciznije planirati prihode i dinamiku ugovaranja novih poslova.'
        ]
    )

    add_table(
        doc,
        ['Kvartal', 'Iznos', 'Ukupno sati'],
        QUARTER_REVENUE,
        'Tabela 8. Prihodi po kvartalima',
        [
            'Najveći prihod ostvaren je u prvom kvartalu, kada je evidentirano 48.300 uz 1.675 sati rada, dok se u četvrtom kvartalu beleži najniža vrednost od 4.950 i 165 sati. Ovakva raspodela ukazuje na izraženu koncentraciju poslovne aktivnosti u prvoj polovini analiziranog perioda.',
            'Poređenje kvartala omogućava uočavanje sezonskih oscilacija u angažovanju i naplati. Takav uvid je od posebnog značaja za planiranje kapaciteta, ugovaranje novih projekata i upravljanje novčanim tokovima kompanije.'
        ]
    )

    add_heading(doc, '6. Zaključak')
    zakljucak = [
        'Na osnovu sprovedene analize može se zaključiti da je skladište podataka efikasan okvir za objedinjavanje operativnih i analitičkih informacija o poslovanju IT kompanije. Projektovani model uspešno povezuje podatke o zaposlenima, odeljenjima, projektima, klijentima, tehnologijama, evidenciji rada i fakturama, čime se obezbeđuje pouzdana osnova za menadžersko izveštavanje i strateško odlučivanje.',
        'Rezultati analize produktivnosti pokazuju da Development odeljenje nosi najveći deo radnog opterećenja, dok su Marko Petrović, Ana Stojanović i Nikola Todorović među najangažovanijim zaposlenima. Na nivou projekata najveći obim rada zabeležen je na projektu E-Commerce platforma, što ukazuje na njegovu operativnu zahtevnost. Istovremeno, prihodna analiza pokazuje da DataSys GmbH i WebPro Ltd predstavljaju najznačajnije izvore prihoda, dok je prvi kvartal najintenzivniji period posmatranog poslovanja.',
        'Posebna vrednost rada ogleda se u tome što su analitički pogledi i OLAP kocke omogućili brzo sagledavanje podataka iz više perspektiva, bez potrebe za ponovnim definisanjem složenih upita. Na taj način potvrđena je praktična korist dimenzionalnog modelovanja i multidimenzionalne analize u poslovnom okruženju u kome je neophodno svakodnevno donositi odluke o raspodeli resursa, troškovima i prioritetima u radu sa klijentima.',
        'Kao preporuka za buduća unapređenja može se navesti uvođenje detaljnijih vremenskih hijerarhija, pokazatelja profitabilnosti po projektu i klijentu, kao i integracija ETL procesa sa spoljnim izvorima podataka. Dodatno unapređenje predstavljalo bi uključivanje prediktivnih modela koji bi na osnovu istorijskih trendova procenjivali buduće opterećenje resursa, verovatni prihod i rizik od kašnjenja projekata. Time bi skladište podataka preraslo iz sistema za retrospektivnu analizu u sistem za proaktivnu podršku odlučivanju.'
    ]
    for paragraph in zakljucak:
        add_paragraph(doc, paragraph)

    add_heading(doc, '7. Literatura')
    for idx, item in enumerate(LITERATURE, start=1):
        add_paragraph(doc, f'{idx}. {item}', align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0), first_line=0)

    if OUTPUT_PATH.exists():
        OUTPUT_PATH.unlink()
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == '__main__':
    for image in IMAGES.values():
        if not image.exists():
            raise FileNotFoundError(f'Nedostaje obavezna slika: {image}')
    output = build_document()
    print(f'Kreiran dokument: {output}')
    print(f'Veličina fajla: {output.stat().st_size} bajta')
