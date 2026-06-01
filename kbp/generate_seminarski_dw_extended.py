# -*- coding: utf-8 -*-
from pathlib import Path
import math
import re
import shutil
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from openpyxl import load_workbook


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / 'Seminarski_rad_DW_ITKompanije.docx'
PIVOT_XLSX = BASE_DIR / 'OLAP_Pivot_Tabele.xlsx'
MDX_PATH = BASE_DIR / 'MDX_upiti_za_kocku.mdx'

IMAGE_ALIASES = {
    'CDM.png': 'slika 1.png',
    'PDM.png': 'slika2 relation.png',
    'DB_diagram.png': 'diag baze pod.png',
    'VIEW_results.png': 'anal troskova rada.png',
    'SELECT_queries_1.png': 'korisniSelectModeli.png',
    'SELECT_queries_2.png': 'analiza prihoda.png',
    'SELECT_queries_3.png': 'analiza tehnologija.png',
}

DISPLAY_REPLACEMENTS = {
    'Ana Stojanovic': 'Ana Stojanović',
    'Ivan Markovic': 'Ivan Marković',
    'Jelena Nikolic': 'Jelena Nikolić',
    'Maja Ilic': 'Maja Ilić',
    'Marko Petrovic': 'Marko Petrović',
    'Milica Djordjevic': 'Milica Đorđević',
    'Nikola Todorovic': 'Nikola Todorović',
    'Petar Pavlovic': 'Petar Pavlović',
    'Stefan Jovanovic': 'Stefan Jovanović',
    'CloudNet d.o.o': 'CloudNet d.o.o.',
    'Trosak Rada': 'Trošak rada',
    'Broj Sati': 'Broj sati',
    'Broj Faktura': 'Broj faktura',
    'Ukupno Sati': 'Ukupno sati',
    'Budzet': 'Budžet',
    'Projekat': 'Projekat',
    'Zaposleni': 'Zaposleni',
    'Klijent': 'Klijent',
    'Odeljenje': 'Odeljenje',
    'Satnica': 'Satnica',
    'Godina': 'Godina',
    'Kvartal': 'Kvartal',
    'Iznos': 'Iznos',
}

PIVOT_METADATA = {
    'Produktivnost_Odeljenja': {
        'title': 'Produktivnost po odeljenjima',
        'analysis': [
            'Pivot analiza po odeljenjima pokazuje da je najveći deo radnog opterećenja koncentrisan u razvojnom delu kompanije. Takav raspored je očekivan, jer poslovni model IT kompanije najveću vrednost stvara upravo kroz implementaciju softverskih rešenja i održavanje projektnog tempa. Istovremeno, vidljivo je da QA, dizajn i menadžment imaju podržavajuću, ali ne i zanemarljivu ulogu u ukupnom trošku rada.',
            'Odnos između broja sati i troška rada omogućava dublju interpretaciju. Odeljenje Development ostvaruje i najveći zbir satnica, što ukazuje na veći broj angažovanih stručnjaka sa višim nivoom odgovornosti. Menadžmentu ovakav pregled omogućava da proceni da li je odnos između organizacionih celina održiv i da li je potrebno uravnotežiti opterećenje između timova.'
        ],
    },
    'Produktivnost_Projekti': {
        'title': 'Produktivnost po projektima',
        'analysis': [
            'Analiza po projektima pokazuje da se radni angažman ne raspoređuje ravnomerno. Projekti E-Commerce platforma i CRM sistem zauzimaju najveći deo radnog fonda, što ukazuje na njihovu funkcionalnu složenost i veći broj poslovnih zahteva. Nasuprot tome, Analytics Dashboard ima najmanji zbir sati, ali to ne mora značiti niži poslovni značaj, već kraći životni ciklus ili bolje definisan opseg rada.',
            'Kada se zajedno posmatraju broj sati i trošak rada, uočava se da isti fond rada može da proizvede različite finansijske efekte u zavisnosti od strukture angažovanih profila. Takav zaključak je od posebne važnosti za planiranje budućih ponuda, jer pomaže u realističnijem procenjivanju obima rada i cene projekta.'
        ],
    },
    'Produktivnost_Zaposleni': {
        'title': 'Produktivnost po zaposlenima',
        'analysis': [
            'Na nivou pojedinačnih zaposlenih jasno se vidi koncentracija angažovanja kod nekoliko ključnih nosilaca posla. Marko Petrović, Ana Stojanović i Nikola Todorović izdvajaju se po broju sati i trošku rada, što ukazuje da upravo ovi stručnjaci nose najveći teret realizacije projekata. Takva raspodela može biti opravdana zbog kompetencija, ali istovremeno nosi i organizacioni rizik ukoliko se oslanjanje na mali broj resursa dodatno poveća.',
            'Kolona broj evidencija daje dodatni uvid u obrazac rada. Zaposleni sa većim brojem evidencija nisu nužno radili najviše sati po pojedinačnom unosu, ali je njihov angažman bio češći i ravnomernije raspoređen. Menadžment može da koristi ovakav prikaz radi uočavanja preopterećenja, planiranja zamena i definisanja programa razvoja kompetencija.'
        ],
    },
    'Produktivnost_Godine': {
        'title': 'Produktivnost po godinama',
        'analysis': [
            'Vremenska raspodela produktivnosti pokazuje da je u 2024. godini ostvaren veći fond sati i veći ukupan trošak rada nego u 2025. godini. Takva raspodela može se objasniti činjenicom da je u 2024. godini istovremeno aktivan veći broj projekata u fazi intenzivne realizacije, dok je 2025. godina delimično obuhvatila nastavak i završne faze pojedinih angažmana.',
            'Analiza po godinama ima veliku upravljačku vrednost, jer omogućava upoređivanje obima rada između poslovnih ciklusa. Time se dobija osnov za procenu rasta, sezonskih opterećenja i dugoročnih trendova u iskorišćenju ljudskih resursa.'
        ],
    },
    'Prihodi_Klijenti': {
        'title': 'Prihodi po klijentima',
        'analysis': [
            'Analiza prihoda po klijentima otkriva da portfelj prihoda nije ravnomerno raspoređen. DataSys GmbH, WebPro Ltd i TechCorp Solutions generišu najveći deo ukupne naplate, što znači da su odnosi sa ovim klijentima od strateškog značaja za kompaniju. Manji klijenti ostaju važni, ali se njihov doprinos više vidi u diverzifikaciji portfelja nego u apsolutnom finansijskom efektu.',
            'Poređenje iznosa i ukupno utrošenih sati ukazuje na to da broj angažovanih sati ne vodi automatski ka proporcionalno većem prihodu. Takva razlika može nastati zbog različitih ugovornih cena, strukture usluge, naplate fiksnih faza ili različitog stepena dodatne vrednosti koju kompanija isporučuje.'
        ],
    },
    'Prihodi_Projekti': {
        'title': 'Prihodi po projektima',
        'analysis': [
            'Na nivou projekata vidi se da CRM sistem i Mobile Banking App ostvaruju najveće prihode, dok Analytics Dashboard ima najniži iznos. Međutim, posmatranje samo ostvarenog prihoda nije dovoljno. Uključivanje budžeta pokazuje da pojedini projekti i dalje poseduju značajan neiskorišćen prihodovni potencijal, što je posebno važno kada projekat nije u završnoj fazi.',
            'Broj faktura dodatno pomaže interpretaciju projektne dinamike. Projekti sa većim brojem faktura pokazuju stabilniji obrazac naplate i verovatno bolje definisane isporuke, dok manji broj faktura može ukazivati na duže obračunske cikluse ili drugačiju ugovornu strukturu.'
        ],
    },
    'Prihodi_Godine': {
        'title': 'Prihodi po godinama',
        'analysis': [
            'Godišnja analiza prihoda potvrđuje da je 2024. godina bila finansijski intenzivnija od 2025. godine. Veći broj faktura i veći iznos ukazuju na to da je naplata bila jače koncentrisana u ranijem delu posmatranog perioda, što može biti posledica završetka velikih faza projekata ili uspešnije dinamike fakturisanja.',
            'Za strateško planiranje posebno je značajno to što se na godišnjem nivou lako uočavaju promene u ritmu poslovanja. Takvi uvidi omogućavaju pravovremeno prilagođavanje prodajnih aktivnosti, planova zapošljavanja i očekivanih novčanih tokova.'
        ],
    },
    'Prihodi_Kvartali': {
        'title': 'Prihodi po kvartalima',
        'analysis': [
            'Kvartalni pregled otkriva jasnu sezonsku neravnomernost. Prvi kvartal ostvaruje najveći prihod i najveći fond sati, dok je četvrti kvartal najslabiji. Takav obrazac može značiti da se početkom godine zatvaraju značajne projektne faze ili da se naplata intenzivira odmah nakon godišnjeg planiranja i ugovaranja.',
            'Za menadžersko odlučivanje ovaj prikaz ima praktičnu vrednost zato što olakšava usklađivanje kapaciteta, dinamike fakturisanja i planiranja likvidnosti. OLAP perspektiva po kvartalima posebno je korisna kada se želi brzo preći sa godišnjeg pregleda na detaljniji vremenski nivo.'
        ],
    },
}

MDX_EXPLANATIONS = {
    1: [
        'Prvi MDX upit agregira mere Broj sati, Trošak rada i Satnica po članovima dimenzije Odeljenje. Na taj način se demonstrira osnovna OLAP operacija roll-up, pošto se pojedinačne evidencije rada sabiraju na nivou organizacionih celina.',
        'Interpretacija rezultata ovog upita omogućava da se utvrdi u kom delu organizacije nastaje najveći radni i troškovni pritisak. U radu je ovaj upit korišćen kao polazna tačka za procenu raspodele kapaciteta između razvojnih, kontrolnih i upravljačkih funkcija.'
    ],
    2: [
        'Drugi MDX upit posmatra cbProduktivnost kroz dimenziju Projekat. Time se ostvaruje poređenje projekata prema obimu rada i ukupnom trošku angažovanja, bez potrebe za pisanjem novih SQL agregacija nad relacijskim slojem.',
        'Rezultati ovakvog upita posebno su korisni prilikom revizije projektnog portfolija. Analizom je moguće utvrditi koji projekti traže najveće resurse, kao i da li su visoki troškovi povezani sa strateški važnim klijentima i isporukama.'
    ],
    3: [
        'Treći MDX upit spušta analizu na nivo zaposlenih i uvodi meru Broj evidencija. Time se omogućava drill-down sa organizacionog ili projektnog nivoa ka pojedinačnom nosiocu rada.',
        'Upravljački značaj ovog upita ogleda se u mogućnosti da se prepoznaju ključni resursi, preopterećenja i neujednačena distribucija angažovanja. Takav uvid je posebno važan za planiranje obuka, rotaciju zaposlenih i smanjenje operativnog rizika.'
    ],
    4: [
        'Četvrti MDX upit prikazuje produktivnost po godinama i time uvodi vremensku dimenziju u analizu rada. Iako je rezultat agregiran samo na dva člana, vremenski pogled je dovoljan da se prepozna trend između poslovnih ciklusa.',
        'U radu se ovakav upit koristi za povezivanje analize produktivnosti sa prihodnom analizom, jer tek zajedničko posmatranje vremena, rada i naplate omogućava ocenu poslovne efikasnosti kroz duži horizont.'
    ],
    5: [
        'Peti MDX upit iz kocke cbPrihodi prikazuje prihodnu strukturu po klijentima, uz mere Iznos, Ukupno sati i Broj faktura. Time se u jednoj tabeli spajaju komercijalni rezultat, opterećenje resursa i dinamika naplate.',
        'Ovaj upit omogućava segmentaciju klijentskog portfelja i identifikaciju kupaca od najvećeg strateškog značaja. U praksi predstavlja osnovu za diskusiju o profitabilnosti saradnje, kvalitetu naplate i potencijalu za dalje širenje usluga.'
    ],
    6: [
        'Šesti MDX upit posmatra prihode po projektima i uključuje meru Budžet. Time se uvodi odnos između realizovanog i planiranog finansijskog obima, što je od ključne važnosti za kontrolu projektne uspešnosti.',
        'Tumačenjem rezultata može se prepoznati projekat koji trenutno ostvaruje visok prihod, ali i projekat kod koga budžet ukazuje na budući potencijal. Upravo zbog toga je OLAP analiza pogodna za periodično praćenje projektne komercijalne slike.'
    ],
    7: [
        'Sedmi MDX upit daje prihodni pregled po godinama. U pitanju je roll-up iznad transakcionih faktura ka godišnjem nivou, čime se veoma brzo dobija sažet strateški pogled na poslovanje.',
        'U radu su rezultati ovakvog upita korišćeni kao osnova za poređenje sa kvartalnim prikazom, jer tek kombinovanje oba nivoa omogućava pravilno razumevanje kada nastaju ključni finansijski vrhovi i padovi.'
    ],
    8: [
        'Osmi MDX upit prikazuje prihode po kvartalima i time predstavlja prirodan drill-down u odnosu na godišnji pregled. Kvartalni nivo je dovoljno detaljan da pokaže sezonske oscilacije, a da pritom ne optereti analizu mesečnim šumom.',
        'Rezultati ovog upita pogodni su za planiranje prodajnih i operativnih aktivnosti. Ukoliko se u više uzastopnih perioda potvrdi isti obrazac, moguće je unaprediti raspored resursa i dinamiku fakturisanja sa većom predvidivošću.'
    ],
}

LITERATURE = [
    'Inmon, W. H. (2005). Building the Data Warehouse. Wiley.',
    'Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling. Wiley.',
    'Ponniah, P. (2010). Data Warehousing Fundamentals for IT Professionals. Wiley.',
    'Golfarelli, M., & Rizzi, S. (2009). Data Warehouse Design: Modern Principles and Methodologies. McGraw-Hill.',
    'Chaudhuri, S., & Dayal, U. (1997). An Overview of Data Warehousing and OLAP Technology. ACM SIGMOD Record, 26(1), 65-74.',
    'Vassiliadis, P. (2010). A Survey of Extract-Transform-Load Technology. International Journal of Data Warehousing and Mining, 5(3), 1-27.',
    'Kimball, R., Reeves, L., Ross, M., & Thornthwaite, W. (1998). The Data Warehouse Lifecycle Toolkit. Wiley.',
    'Elmasri, R., & Navathe, S. B. (2016). Fundamentals of Database Systems. Pearson.',
    'Harrington, J. L. (2008). Relational Database Design and Implementation. Morgan Kaufmann.',
    'Microsoft. (2024). SQL Server Analysis Services Documentation. Microsoft Learn.',
    'Oracle. (2024). Oracle SQL Developer Data Modeler User\'s Guide. Oracle Documentation.',
    'Abelló, A., Samos, J., & Saltor, F. (2006). YAM²: A Multidimensional Conceptual Model Extending UML. Information Systems, 31(6), 541-567.',
]

ABSTRACT = "U radu je prikazano projektovanje skladišta podataka i OLAP okruženja namenjenog analizi poslovanja IT kompanije koja istovremeno upravlja zaposlenima, odeljenjima, projektima, klijentima, tehnologijama, evidencijom rada i fakturisanjem usluga. Polazni operativni model implementiran je u relacijskoj bazi, a zatim je kroz definisane poglede, ETL logiku i SQL Server Analysis Services infrastrukturu transformisan u analitičko rešenje pogodno za višedimenzionalnu obradu. Poseban akcenat stavljen je na izbor pristupa projektovanju skladišta podataka, definisanje dimenzija i mera, modelovanje istorijskih promena, kao i na interpretaciju rezultata dobijenih iz OLAP kocki cbProduktivnost i cbPrihodi. Analizom je potvrđeno da ovakav model omogućava konzistentno sagledavanje troškova rada, opterećenja resursa, prihoda po klijentima i projektima, kao i vremenske dinamike poslovanja. Rezultati rada pokazuju da skladište podataka predstavlja pouzdanu osnovu za menadžersko odlučivanje, unapređenje planiranja kapaciteta i buduće širenje sistema ka naprednoj poslovnoj inteligenciji."
KEYWORDS = 'Ključne reči: skladište podataka, OLAP, ETL, dimenzionalno modeliranje, SQL Server, analiza produktivnosti, analiza prihoda.'


def paragraphs(text):
    return [segment.strip().replace('\n', ' ') for segment in text.strip().split('\n\n') if segment.strip()]


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, first_line=0.75, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(first_line) if first_line else Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=bold, italic=italic)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def add_caption(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)
    return p


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=11, bold=bold)


def add_table(doc, state, headers, rows, title, analysis_paragraphs):
    add_caption(doc, f'Tabela {state["table"]}: {title}')
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell(hdr[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[idx], value, align=align)
    state['table'] += 1
    for item in analysis_paragraphs:
        add_paragraph(doc, item)


def add_image(doc, state, image_path, title, analysis_paragraphs, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, f'Slika {state["figure"]}: {title}')
    state['figure'] += 1
    for item in analysis_paragraphs:
        add_paragraph(doc, item)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text.strip())
    set_run_font(run, size=11)
    return p


def add_section(doc, heading, text):
    add_heading(doc, heading)
    for item in paragraphs(text):
        add_paragraph(doc, item)


def ensure_image_aliases():
    image_dir = BASE_DIR / 'diagrami'
    for alias_name, source_name in IMAGE_ALIASES.items():
        source = image_dir / source_name
        target = image_dir / alias_name
        if not source.exists():
            raise FileNotFoundError(f'Nedostaje izvorna slika: {source}')
        if not target.exists():
            shutil.copyfile(source, target)
    return {name: image_dir / name for name in IMAGE_ALIASES}


def normalize_value(value):
    if value is None:
        return ''
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    return DISPLAY_REPLACEMENTS.get(text, text)


def load_pivot_tables():
    wb = load_workbook(PIVOT_XLSX, data_only=True)
    result = {}
    for ws in wb.worksheets:
        headers = [normalize_value(cell) for cell in next(ws.iter_rows(min_row=3, max_row=3, values_only=True))]
        rows = []
        for row in ws.iter_rows(min_row=4, max_row=ws.max_row, values_only=True):
            if any(cell is not None for cell in row):
                rows.append([normalize_value(cell) for cell in row])
        result[ws.title] = {'headers': headers, 'rows': rows}
    return result


def load_mdx_queries():
    lines = MDX_PATH.read_text(encoding='utf-8').splitlines()
    queries = []
    current = None
    for line in lines:
        match = re.match(r'-- Upit (\d+):\s*(.+)', line)
        if match:
            if current:
                current['query'] = '\n'.join(item for item in current['lines'] if item.strip()).strip()
                queries.append(current)
            current = {
                'number': int(match.group(1)),
                'title': match.group(2).strip(),
                'lines': [],
            }
            continue
        if current is not None and not line.startswith('-- ==========================================') and not line.startswith('-- KOCKA'):
            current['lines'].append(line)
    if current:
        current['query'] = '\n'.join(item for item in current['lines'] if item.strip()).strip()
        queries.append(current)
    return queries


def set_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for line in [
        'UNIVERZITET U NOVOM SADU',
        'Tehnički fakultet "Mihajlo Pupin" Zrenjanin',
        'Master akademske studije',
    ]:
        run = p.add_run(line + '\n')
        set_run_font(run, size=12, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    for line in ['SEMINARSKI RAD', 'Data Warehouse IT kompanije', 'Projektovanje skladišta podataka i OLAP analize za IT kompaniju']:
        run = title.add_run(line + '\n')
        set_run_font(run, size=16, bold=True)

    for _ in range(7):
        doc.add_paragraph()

    for line in [
        'Student: Ognjen Grgur',
        'Broj indeksa: MIT 37/24',
        'Predmet: Koncepti baza podataka',
    ]:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_info.paragraph_format.line_spacing = 1.5
        run = p_info.add_run(line)
        set_run_font(run, size=12, bold=True)

    city = doc.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city.paragraph_format.space_before = Pt(36)
    city.paragraph_format.line_spacing = 1.5
    run = city.add_run('Zrenjanin, 2025.')
    set_run_font(run, size=12, bold=True)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, 'Sadržaj')
    items = [
        '1. Uvod',
        '2. Prošireni teorijski okvir skladišta podataka',
        '2.1 Istorijski razvoj koncepta skladišta podataka',
        '2.2 Definicija i karakteristike skladišta podataka',
        '2.3 Komparativna analiza pristupa top-down i bottom-up',
        '2.4 OLTP i OLAP sistemi',
        '2.5 ETL proces: izdvajanje, transformacija i punjenje',
        '2.6 Tipovi dimenzionalnog modeliranja',
        '2.7 Sporo promenljive dimenzije',
        '3. Metodologija rada',
        '3.1 Korišćeni alati',
        '3.2 Pristup projektovanju',
        '4. Projektovanje i implementacija skladišta podataka',
        '4.1 Opis poslovnog domena',
        '4.2 Konceptualni model',
        '4.3 Fizički model',
        '4.4 Detaljna implementacija relacione baze',
        '4.5 ETL proces punjenja skladišta',
        '4.6 Indeksiranje i optimizacija performansi',
        '4.7 Pogledi i priprema podataka za OLAP',
        '5. Projektovanje OLAP kocke',
        '5.1 Dimenzije i hijerarhije',
        '5.2 Mere i kalkulisani pokazatelji',
        '5.3 Šema zvezde i organizacija činjenica',
        '5.4 Kreiranje i procesiranje kocke',
        '6. OLAP analiza i interpretacija rezultata',
        '6.1 OLAP operacije',
        '6.2 MDX upiti i njihovo tumačenje',
        '6.3 Pivot analize produktivnosti',
        '6.4 Pivot analize prihoda',
        '6.5 Prednosti OLAP pristupa u odnosu na klasično izveštavanje',
        '7. Zaključak',
        '8. Literatura',
    ]
    for item in items:
        add_paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, space_after=0)
    doc.add_page_break()


def build_document():
    images = ensure_image_aliases()
    pivot_tables = load_pivot_tables()
    mdx_queries = load_mdx_queries()

    doc = Document()
    set_document_defaults(doc)
    state = {'table': 1, 'figure': 1}

    add_title_page(doc)
    add_heading(doc, 'Apstrakt')
    add_paragraph(doc, ABSTRACT)
    add_paragraph(doc, KEYWORDS, italic=True)
    doc.add_page_break()
    add_contents(doc)

    add_section(doc, '1. Uvod', """
    Savremeno poslovanje IT kompanija zasniva se na kontinuiranom stvaranju podataka o zaposlenima, projektima, klijentima, tehnologijama, utrošenim radnim satima i fakturisanim uslugama. Iako operativni informacioni sistemi omogućavaju pouzdano evidentiranje pojedinačnih transakcija, njihova struktura nije pogodna za izvođenje složenih analitičkih upita koji se koriste u strateškom upravljanju. Upravo iz tog razloga skladište podataka predstavlja centralni mehanizam za objedinjavanje, istorijsko čuvanje i analitičko tumačenje poslovnih informacija.

    U radu je razmatrana IT kompanija koja istovremeno realizuje više projekata za različite klijente, raspolaže zaposlenima različitih profila i koristi više tehnologija u okviru projektnih angažmana. U takvom poslovnom okruženju menadžment mora da odgovori na pitanja koja prevazilaze operativni nivo, kao što su: koje odeljenje generiše najveći trošak rada, koji projekti ostvaruju najveći prihod, da li postoji neravnomerno opterećenje zaposlenih i u kojim vremenskim periodima nastaju finansijski vrhovi ili padovi.

    Da bi se na navedena pitanja odgovorilo na pouzdan i dosledan način, u radu je polazna relaciona baza podataka proširena analitičkim slojem. Taj sloj obuhvata definisanje pogleda kao pripreme za analizu, modelovanje činjenica i dimenzija, kao i kreiranje OLAP kocki namenjenih višedimenzionalnom ispitivanju poslovanja. Primenom ovakvog pristupa obezbeđeno je da se podaci sagledavaju iz više uglova, bez ponavljanja složenih relacijskih upita i bez narušavanja performansi operativnog sistema.

    Naučni i stručni značaj teme proizlazi iz činjenice da skladišta podataka predstavljaju standardnu infrastrukturu poslovne inteligencije u savremenim organizacijama. Posebno u domenu IT usluga, gde su resursi, znanje i vreme glavni nosioci vrednosti, kvalitativno i kvantitativno sagledavanje radnog angažovanja ima neposredan uticaj na profitabilnost i konkurentnost kompanije. Zbog toga je projektovanje skladišta podataka za IT kompaniju pogodno za demonstraciju veze između relacijskog modelovanja, ETL logike i OLAP analitike.

    Cilj rada jeste da se, na konkretnom primeru, prikaže celokupan tok izgradnje analitičkog sistema: od teorijskih osnova i metodologije rada, preko implementacije relacione šeme i pripreme podataka, do izgradnje OLAP kocke i interpretacije rezultata. Dodatni cilj predstavlja demonstracija toga da dobro osmišljen analitički model može da obezbedi jedinstvenu osnovu za praćenje produktivnosti, troškova rada, prihoda i vremenskih trendova.
    """)

    add_section(doc, '2. Prošireni teorijski okvir skladišta podataka', """
    Teorijski okvir skladišta podataka obuhvata skup koncepata, metoda i arhitekturnih odluka koje određuju način prikupljanja, oblikovanja i korišćenja analitičkih podataka. U literaturi se skladište podataka ne posmatra samo kao tehničko spremište, već kao organizacioni mehanizam za pretvaranje transakcionih zapisa u znanje pogodno za odlučivanje.

    U odnosu na klasične baze podataka, skladište podataka uvodi dodatnu vrednost kroz integraciju više izvora, vremensku stabilnost, kontrolisanu redundantnost i orijentaciju na poslovne subjekte. Iz tog razloga teorijski okvir ne obuhvata samo relacione principe, već i dimenzionalno modeliranje, ETL procese, upravljanje kvalitetom podataka i višedimenzionalnu analizu.

    Za potrebe ovog rada posebna pažnja usmerena je na istorijski razvoj discipline, izbor između Inmonovog i Kimballovog pristupa, razumevanje ETL procesa, vrste dimenzionalnih modela i obradu sporih promena u dimenzijama. Ovi elementi predstavljaju teorijsku osnovu za kasnije projektantske odluke primenjene na primeru IT kompanije.
    """)

    add_section(doc, '2.1 Istorijski razvoj koncepta skladišta podataka', """
    Početak savremenog koncepta skladišta podataka najčešće se vezuje za rad Vilijama Inmona iz osamdesetih godina dvadesetog veka. Inmon je skladište podataka definisao kao subjektno orijentisanu, integrisanu, vremenski promenljivu i nepromenljivu kolekciju podataka namenjenu podršci odlučivanju. Ova definicija ostala je jedna od najuticajnijih u literaturi, jer na sažet način ukazuje na ključnu razliku između operativnih i analitičkih informacionih sistema.

    Tokom devedesetih godina Ralf Kimball je dodatno popularizovao skladišta podataka kroz pragmatičan i poslovno orijentisan pristup. Za razliku od centralno projektovanog korporativnog skladišta, Kimball je predlagao izgradnju skladišta kroz skup povezanih data martova, zasnovanih na konformnim dimenzijama i dimenzionalnom modelu pogodnom za brzu isporuku poslovne vrednosti. Time je disciplina dobila dve snažne škole mišljenja koje su oblikovale praksu do danas.

    Početkom dvehiljaditih godina razvoj skladišta podataka intenzivno je povezan sa rastom alata poslovne inteligencije, pojavom standardizovanih ETL platformi i komercijalnom primenom OLAP servera. Organizacije su počele da prepoznaju potrebu za odvajanjem operativnog i analitičkog sloja, pa skladište podataka postaje sastavni deo korporativne arhitekture, a ne samo eksperimentalno rešenje namenjeno finansijskim izveštajima.

    Savremeni razvoj discipline obeležen je pojavom kolonarnih baza, cloud skladišta, data lakehouse paradigme i napredne analitike zasnovane na mašinskom učenju. Ipak, bez obzira na tehnološke promene, osnovne ideje Inmona i Kimballa ostale su aktuelne: podaci moraju biti integrisani, vremenski dosledni, poslovno razumljivi i organizovani tako da podrže donošenje odluka. U tom smislu, savremena rešenja češće predstavljaju evoluciju klasičnih principa nego njihovu zamenu.

    Na akademskom i praktičnom planu istorijski razvoj skladišta podataka pokazuje da je reč o disciplini koja je nastala kao odgovor na ograničenja transakcionih sistema. Kako se povećavala količina podataka i zahtevi za preciznijim izveštavanjem, skladišta podataka su od pomoćnog sloja prerasla u ključnu infrastrukturu poslovne inteligencije. Ovakav razvoj opravdava njihovu primenu i u analiziranom slučaju IT kompanije.
    """)

    add_section(doc, '2.2 Definicija i karakteristike skladišta podataka', """
    Skladište podataka predstavlja centralizovan i analitički orijentisan skup podataka čija je namena podrška odlučivanju. Za razliku od operativnih baza, koje prioritet daju ažuriranju i integritetu pojedinačnih transakcija, skladište podataka usmereno je na objedinjavanje istorijskih podataka i brzo izvršavanje kompleksnih upita. Njegova osnovna vrednost nije u unosu podataka, već u njihovoj interpretaciji i poslovnoj upotrebljivosti.

    Subjektna orijentacija znači da su podaci organizovani oko poslovno važnih pojmova kao što su klijent, projekat, zaposleni ili prihod. Time se analitički model približava načinu na koji rukovodstvo razume poslovanje. Integracija podataka podrazumeva usklađivanje naziva atributa, formata, šifarnika i pravila vrednosti, kako bi svi pokazatelji imali jedinstveno značenje u celom sistemu.

    Vremenska promenljivost skladišta podataka posebno je značajna, jer analitička vrednost podataka raste kada je moguće pratiti promene kroz duži period. Pored toga, nepromenljivost u smislu ograničenog i kontrolisanog ažuriranja istorijskih zapisa obezbeđuje konzistentnost izveštavanja. Na taj način menadžment može da poredi rezultate različitih perioda bez rizika da će prethodno obrađeni podaci biti naknadno izmenjeni bez jasnog traga.

    U praksi se navedene karakteristike realizuju kroz ETL procese, dimenzionalni model, pravila kvaliteta podataka i strogo definisane mere. U radu je upravo takva interpretacija skladišta podataka primenjena na domenu IT kompanije, gde se operativni podaci transformišu u pouzdan izvor za procenu produktivnosti, troškova i prihoda.
    """)

    add_section(doc, '2.3 Komparativna analiza pristupa top-down i bottom-up', """
    Inmonov top-down pristup polazi od ideje da je najpre potrebno izgraditi centralno korporativno skladište podataka u trećoj normalnoj formi, a zatim iz njega izvoditi analitičke data martove. Prednost ovakvog pristupa ogleda se u snažnoj integraciji na nivou cele organizacije, jasnom upravljanju podacima i visokom stepenu centralne kontrole. Ovakav model posebno je pogodan za složene organizacije sa velikim brojem izvora i potrebom za dugoročnom standardizacijom.

    Kimballov bottom-up pristup zasniva se na tome da se analitička arhitektura gradi postepeno, kroz poslovno relevantne data martove povezane konformnim dimenzijama. Na taj način poslovna vrednost može biti isporučena brže, a pojedinačni moduli se lakše prilagođavaju prioritetima korisnika. Zbog snažne povezanosti sa dimenzionalnim modelovanjem, ovaj pristup je često intuitivniji za analitičare i poslovne korisnike.

    Nedostatak top-down pristupa ogleda se u većem početnom trajanju projekta i većim organizacionim zahtevima. U slučaju da poslovni zahtevi nisu dovoljno stabilni, postoji rizik da obimna početna arhitektura uspori isporuku konkretnih analitičkih koristi. Sa druge strane, bottom-up pristup može dovesti do fragmentacije i nastanka nepovezanih data martova ukoliko se ne vodi računa o standardizaciji dimenzija i poslovnih definicija.

    Za projekat prikazan u ovom radu prikladniji je pristup koji je konceptualno bliži Kimballovoj filozofiji. Razlog za to jeste potreba da se relativno brzo formiraju jasne analitičke celine za produktivnost i prihode, pri čemu su konformne dimenzije Projekat, Klijent i Vreme iskorišćene kao zajednička osnova za više analiza. Ipak, u strukturi rešenja prisutna je i Inmonova ideja integracije, jer su svi podaci objedinjeni iz jednog dosledno projektovanog relacijskog izvora.
    """)

    add_table(
        doc,
        state,
        ['Kriterijum', 'Top-down pristup', 'Bottom-up pristup'],
        [
            ['Početna tačka', 'Korporativno skladište podataka', 'Poslovni data martovi'],
            ['Model', 'Često 3NF i integrisani model', 'Dimenzionalni model i konformne dimenzije'],
            ['Brzina prve isporuke', 'Sporija', 'Brža'],
            ['Stepen centralne kontrole', 'Veoma visok', 'Srednji do visok'],
            ['Rizik fragmentacije', 'Manji', 'Veći bez standardizacije'],
            ['Pogodnost za ovaj rad', 'Delimično prisutna kroz integraciju', 'Dominantna kroz OLAP data mart logiku'],
        ],
        'Poređenje Inmonovog i Kimballovog pristupa',
        [
            'Komparativna analiza pokazuje da ni jedan pristup nije univerzalno najbolji, već da izbor zavisi od obima projekta, dostupnih resursa i zrelosti organizacije. Za potrebe rada odabran je pragmatičan model koji koristi bottom-up logiku isporuke analitičkih celina, ali zadržava visok nivo integracije podataka.',
            'Takva kombinacija je pogodna za akademski i praktični prikaz, jer omogućava jasnu vezu između relacione baze, analitičkih pogleda i konačnih OLAP kocki.'
        ]
    )

    add_section(doc, '2.4 OLTP i OLAP sistemi', """
    OLTP sistemi predstavljaju osnovu svakodnevnog poslovanja. Njihov zadatak je da brzo i pouzdano evidentiraju pojedinačne događaje, kao što su unos radnih sati, evidentiranje klijenta, otvaranje projekta ili izdavanje fakture. Zbog toga su optimizovani za veliki broj kratkih transakcija, visok stepen konkurentnog pristupa i očuvanje integriteta podataka.

    OLAP sistemi imaju drugačiju namenu. Umesto obrade pojedinačnih događaja, oni su usmereni na agregaciju, poređenje i interpretaciju podataka kroz više dimenzija. Tipični OLAP upiti nisu kratki i transakcioni, već složeni i analitički: oni sabiraju sate po odeljenjima, prihode po projektima, troškove po kvartalima i slične pokazatelje potrebne menadžmentu.

    U domenu posmatrane IT kompanije ova razlika je od suštinskog značaja. Operativni sloj služi za evidenciju radnih angažmana i finansijskih događaja, dok analitički sloj omogućava da se isti podaci posmatraju kroz organizacionu, projektnu, klijentsku i vremensku perspektivu. Zbog toga se u radu transakcioni i analitički nivo posmatraju kao komplementarni, ali strukturno odvojeni delovi istog informacionog sistema.
    """)

    add_table(
        doc,
        state,
        ['Karakteristika', 'OLTP', 'OLAP'],
        [
            ['Primarna svrha', 'Operativna obrada transakcija', 'Analiza i podrška odlučivanju'],
            ['Struktura podataka', 'Normalizovana', 'Dimenzionalna ili agregirana'],
            ['Tip upita', 'Kratki i brojni', 'Složeni i agregirani'],
            ['Vremenska orijentacija', 'Aktuelno stanje', 'Istorijski i trendovski pogled'],
            ['Performanse', 'Optimizovane za unos i izmene', 'Optimizovane za čitanje i sabiranje'],
            ['Korisnici', 'Operateri i administratori', 'Analitičari i menadžeri'],
        ],
        'Poređenje OLTP i OLAP okruženja',
        [
            'Tabela potvrđuje da se OLTP i OLAP sistemi razlikuju po cilju, strukturi i načinu upotrebe, ali da zajedno čine jedinstven informacioni lanac. Bez pouzdanog transakcionog izvora nema kvalitetne analitike, dok bez analitike operativni podaci ostaju ograničeni na lokalni i kratkoročni kontekst.',
            'Upravo zato je u radu projektovano rešenje koje povezuje stabilan relacijski model sa OLAP analizom, čime se postiže ravnoteža između operativne pouzdanosti i analitičke fleksibilnosti.'
        ]
    )

    add_section(doc, '2.5 ETL proces: izdvajanje, transformacija i punjenje', """
    ETL proces predstavlja operativno jezgro svakog skladišta podataka. Njegov zadatak je da podatke izdvoji iz izvornih sistema, transformiše ih u ciljnu analitičku strukturu i učita u skladište ili data martove. Kvalitet ETL procesa presudno utiče na pouzdanost svih kasnijih analiza, jer ni najbolji model ne može nadoknaditi nekonzistentne ili nepotpune ulazne podatke.

    Faza izdvajanja podataka podrazumeva identifikaciju relevantnih izvora, definisanje obima podataka i način njihove periodične ili inkrementalne isporuke. U slučaju ovog rada izvorni podaci nalaze se u relacijskim tabelama koje beleže organizacionu strukturu, zaposlene, projekte, tehnologije, radne sate i fakture. Iako je izvor jedinstvena baza, logika izdvajanja i dalje mora da vodi računa o zavisnostima između tabela i redosledu obrade.

    Transformacija obuhvata standardizaciju vrednosti, obogaćivanje zapisa novim atributima i izračunavanje poslovnih mera. U radu se u okviru transformacije formiraju atributi Godina, Mesec i Kvartal, izvodi se mera Trošak rada na osnovu broja sati i satnice, kao i usklađuju nazivi i veze koje omogućavaju kasnije dimenzionalno modelovanje. Transformaciona pravila predstavljaju trenutak u kome operativni zapis dobija analitičko značenje.

    Faza punjenja podataka podrazumeva upis transformisanih zapisa u ciljne strukture, uz kontrolu integriteta i auditabilnosti. U konkretnom rešenju punjenje je ostvareno indirektno kroz analitičke poglede koji služe kao izvor za OLAP kocke. Takav pristup pojednostavljuje arhitekturu rada, ali i jasno pokazuje kako se ETL logika može realizovati i bez zasebne spoljne platforme, kada je obim podataka prilagođen akademskom projektu.

    U naprednijim produkcionim sistemima ETL obično uključuje i obradu grešaka, evidenciju neuspešnih slogova, praćenje vremena poslednjeg učitavanja i inkrementalnu obradu. Iako takvi mehanizmi nisu u potpunosti implementirani u ovom radu, njihova teorijska važnost ostaje velika i predstavljaju prirodan pravac budućeg razvoja analitičkog sistema.
    """)

    add_section(doc, '2.6 Tipovi dimenzionalnog modeliranja', """
    Dimenzionalno modeliranje predstavlja dominantan način organizovanja podataka u skladištima podataka koja služe poslovnoj inteligenciji. Osnovu čine tabele činjenica sa numeričkim merama i tabele dimenzija koje pružaju poslovni kontekst. Način organizacije ovih elemenata zavisi od potrebe za jednostavnošću, performansama i stepenom normalizacije.

    Šema zvezde podrazumeva da se centralna tabela činjenica direktno povezuje sa denormalizovanim dimenzijama. Prednost ovog modela ogleda se u jednostavnosti upita, preglednosti za korisnike i dobroj pogodnosti za OLAP obradu. Upravo iz tih razloga šema zvezde predstavlja osnovni analitički obrazac korišćen u ovom radu.

    Šema pahulje uvodi dodatnu normalizaciju unutar dimenzija. Takav pristup smanjuje redundantnost, ali povećava broj veza i složenost navigacije kroz model. Pahuljasta struktura može biti opravdana kada su dimenzije obimne, hijerarhijski složene ili kada postoje jasni zahtevi za centralizovanjem zajedničkih šifarnika.

    Konstelacija činjenica, odnosno galaxy model, podrazumeva postojanje više tabela činjenica koje dele pojedine dimenzije. Ovaj pristup je naročito pogodan kada različiti poslovni procesi treba da se analiziraju u jedinstvenom okviru. U radu su kocke cbProduktivnost i cbPrihodi konceptualno bliske ovakvom modelu, jer dele više dimenzija, ali prate različite poslovne događaje i mere.
    """)

    add_table(
        doc,
        state,
        ['Tip modela', 'Osnovna osobina', 'Prednost', 'Ograničenje'],
        [
            ['Šema zvezde', 'Denormalizovane dimenzije oko jedne činjenice', 'Jednostavni i brzi upiti', 'Veća redundantnost'],
            ['Šema pahulje', 'Normalizovane dimenzije sa više nivoa', 'Manja redundantnost', 'Složeniji upiti'],
            ['Konstelacija', 'Više tabela činjenica sa zajedničkim dimenzijama', 'Širi analitički obuhvat', 'Veća projektantska složenost'],
        ],
        'Tipovi dimenzionalnog modeliranja',
        [
            'Na osnovu uporedne analize može se zaključiti da je šema zvezde najpogodnija za glavni analitički tok rada, dok konstelacija predstavlja širi okvir za povezivanje više poslovnih procesa. Upravo takva kombinacija je implicitno prisutna u prikazanom projektu IT kompanije.',
            'Ovakvo sagledavanje dimenzionalnih modela olakšava razumevanje razloga zbog kojih su produktivnost i prihodi modelovani kao odvojene, ali povezane analitičke celine.'
        ]
    )

    add_section(doc, '2.7 Sporo promenljive dimenzije', """
    Sporo promenljive dimenzije, poznate pod skraćenicom SCD, predstavljaju jedan od najvažnijih koncepata u održavanju istorijske doslednosti analitičkih podataka. Problem nastaje kada se atribut dimenzije menja tokom vremena, na primer kada zaposleni pređe u drugo odeljenje, projekat promeni status ili klijent promeni grad poslovanja. Način evidentiranja takvih promena direktno utiče na kvalitet istorijskih analiza.

    SCD tip 1 podrazumeva da se prethodna vrednost jednostavno prepiše novom vrednošću. Ovaj pristup je jednostavan i često dovoljan kada istorijska vrednost nije analitički važna, ali istovremeno dovodi do gubitka mogućnosti da se prati kako je atribut izgledao u ranijim periodima.

    SCD tip 2 uvodi novi red u dimenziji pri svakoj značajnoj promeni, uz čuvanje datuma važenja ili oznake aktivnog zapisa. Time se istorija čuva u potpunosti, pa je moguće analizirati poslovne događaje u odnosu na stanje dimenzije koje je važilo u trenutku nastanka činjenice. Za ozbiljnije analitičke sisteme ovaj pristup se najčešće smatra standardom.

    SCD tip 3 zadržava ograničen broj prethodnih vrednosti kroz dodatne kolone, pa omogućava praćenje tekuće i neposredno prethodne vrednosti atributa. Iako ne čuva celokupnu istoriju kao tip 2, ovaj model može biti koristan kada postoji potreba za jednostavnim poređenjem stare i nove vrednosti bez uvećavanja broja redova.

    U posmatranom radu teorijski je najrelevantniji SCD tip 2, naročito za dimenzije Zaposleni, Projekat i Klijent. Na primer, promena odeljenja zaposlenog ili statusa projekta može biti od velikog značaja za pravilno tumačenje istorijskih rezultata. Iako puna implementacija SCD mehanizma nije bila nužna zbog ograničenog obima podataka, koncept predstavlja važnu osnovu za buduće proširenje skladišta podataka.
    """)

    add_section(doc, '3. Metodologija rada', """
    Metodologija rada zasniva se na povezivanju teorijske analize i praktične implementacije. Najpre su definisani poslovni zahtevi analitičkog sistema, zatim je razmotrena teorijska osnova skladišta podataka, a nakon toga je pristupljeno projektovanju relacione baze, pripremi analitičkih pogleda i izgradnji OLAP kocki. Ovakav sled rada omogućava da se svaka tehnička odluka poveže sa jasnim poslovnim ciljem.

    Istraživački pristup je kombinovan: teorijski deo rada oslanja se na relevantnu literaturu iz oblasti skladišta podataka, dimenzionalnog modeliranja i OLAP tehnologija, dok praktični deo koristi konkretne artefakte nastale tokom implementacije. Na taj način rad ne ostaje na nivou apstraktne deskripcije, već prikazuje kako se teorijski koncepti operacionalizuju u realističnom primeru IT kompanije.

    Posebna metodološka vrednost rada ogleda se u tome što su isti podaci posmatrani kroz više faza: od izvornog relacijskog modela, preko pogleda za pripremu podataka, do konačne OLAP analize. Takav pristup omogućava proveru konzistentnosti između slojeva sistema i pruža jasnu sliku o tome kako svaka transformacija utiče na krajnje pokazatelje.
    """)

    add_section(doc, '3.1 Korišćeni alati', """
    Za implementaciju relacione baze i izvršavanje SQL skripti korišćen je Microsoft SQL Server. Izbor ovog alata opravdan je njegovom stabilnošću, širokom zastupljenošću u akademskom i poslovnom okruženju, kao i dobrom integracijom sa alatima za analitiku i administraciju podataka.

    Za projektovanje konceptualnog i fizičkog modela korišćen je Oracle Data Modeler, odnosno ekvivalentni alati za modelovanje koji omogućavaju vizuelno prikazivanje entiteta, atributa i relacija. Upotreba modelerskog alata bila je važna kako bi se pre same implementacije jasno definisala logika poslovnog domena i pravila integriteta.

    Za izgradnju OLAP kocke i definisanje dimenzija, mera i procesa procesiranja korišćen je SQL Server Analysis Services. Ovaj alat omogućava formiranje višedimenzionalnih struktura pogodnih za brzo agregiranje podataka i izvođenje MDX upita, čime predstavlja centralnu komponentu praktičnog dela rada.

    Za prikaz i proveru rezultata korišćen je Microsoft Excel kroz pivot tabele, dok su MDX upiti pripremani i izvršavani u okruženju SQL Server Management Studio. Ovakva kombinacija alata omogućila je jasnu vezu između projektantskog, implementacionog i interpretativnog nivoa rada.
    """)

    add_section(doc, '3.2 Pristup projektovanju', """
    Pristup projektovanju započet je identifikacijom ključnih poslovnih entiteta i analitičkih pitanja. Kao centralni poslovni pojmovi izdvojeni su zaposleni, odeljenja, projekti, klijenti, tehnologije, evidencija rada i fakture. Analitička pitanja grupisana su oko tri tematske celine: produktivnost resursa, trošak rada i prihodna uspešnost poslovanja.

    Nakon definisanja poslovnog domena pristupljeno je modelovanju konceptualnog i fizičkog sloja. Posebna pažnja posvećena je uspostavljanju primarnih i stranih ključeva, jer je njihova doslednost preduslov za kasnije formiranje analitičkih pogleda i OLAP dimenzija. U ovoj fazi uspostavljena je i veza između projektnih, kadrovskih i finansijskih podataka.

    Sledeći korak bio je definisanje pogleda koji objedinjuju podatke iz više tabela i uvode izvedene atribute potrebne za analizu. Time je relacijski model pripremljen za izgradnju OLAP kocke bez dodatnog narušavanja operativnog sloja. Konačno, nad pripremljenim pogledima definisane su OLAP kocke, MDX upiti i pivot tabele, čime je zatvoren ceo ciklus od izvora podataka do menadžerske interpretacije rezultata.
    """)

    add_section(doc, '4. Projektovanje i implementacija skladišta podataka', """
    Praktični deo rada obuhvata modelovanje poslovnog domena, implementaciju relacione baze podataka, pripremu analitičkih pogleda i izgradnju struktura namenjenih OLAP analizi. U ovoj celini teorijski principi iz prethodnog poglavlja prevode se u konkretno tehničko rešenje, prilagođeno zahtevima jedne IT kompanije.

    Posebna pažnja usmerena je na to da model bude dovoljno jednostavan za akademsku demonstraciju, ali i dovoljno bogat da omogući realističnu analizu rada, troškova i prihoda. Zbog toga relacijska osnova sadrži ograničen broj tabela, ali su odnosi među njima projektovani tako da pokrivaju više važnih poslovnih tokova.
    """)

    add_section(doc, '4.1 Opis poslovnog domena', """
    Poslovni domen rada zasniva se na modelu IT kompanije koja realizuje projekte razvoja softvera za više klijenata. Organizacija je podeljena na odeljenja kao što su razvoj, kontrola kvaliteta, dizajn i menadžment, pri čemu svaki zaposleni pripada tačno jednoj organizacionoj celini i poseduje definisanu satnicu, poziciju i datum zaposlenja.

    Projekti se realizuju za konkretne klijente, imaju budžet, period trajanja i status realizacije. U radu je pretpostavljeno da jedan projekat može koristiti više tehnologija, pa je formirana posebna veza između projekta i tehnologije. Time je omogućeno da se poslovna analiza ne zaustavi samo na finansijskim pokazateljima, već da obuhvati i tehnološku strukturu projektnog portfolija.

    Operativni tok podataka obuhvata evidenciju radnih sati zaposlenih po projektima, a zatim i fakturisanje izvršenih usluga. Ovakav domen je pogodan za skladište podataka zato što povezuje kadrovski, operativni i finansijski aspekt poslovanja, pa omogućava višedimenzionalnu analizu učinka kompanije.
    """)

    add_heading(doc, '4.2 Konceptualni model')
    add_paragraph(doc, 'Konceptualni model prikazuje poslovne entitete i njihove logičke veze bez ulaska u tehničke detalje implementacije. Njegova uloga jeste da omogući jasan poslovni pogled na domen i da posluži kao osnova za kasnije fizičko modelovanje.')
    add_image(
        doc,
        state,
        images['CDM.png'],
        'Konceptualni model sistema (CDM)',
        [
            'Na konceptualnom modelu jasno se uočava centralna uloga entiteta Projekat, jer on povezuje klijente, rad zaposlenih, fakturisanje i tehnologije. Takva pozicija projekta odgovara poslovnoj realnosti IT kompanije, u kojoj se najveći deo vrednosti stvara kroz projektne angažmane.',
            'Model je dovoljno bogat da podrži kasniju analitičku obradu, ali istovremeno ostaje pregledan. Upravo ta ravnoteža između jednostavnosti i potpunosti čini konceptualni model pogodnim polazištem za transformaciju u relacijsku i analitičku strukturu.'
        ],
        width=5.9,
    )

    add_heading(doc, '4.3 Fizički model')
    add_paragraph(doc, 'Fizički model prevodi poslovne entitete u konkretne tabele, kolone, ključeve i ograničenja. Na ovom nivou već je moguće precizno definisati kako će podaci biti upisivani, povezivani i kasnije korišćeni za analizu.')
    add_image(
        doc,
        state,
        images['PDM.png'],
        'Fizički model baze podataka (PDM)',
        [
            'Fizički model potvrđuje da su primarni i strani ključevi definisani tako da obezbede referencijalni integritet između organizacionih, projektnih i finansijskih podataka. Poseban značaj ima tabela EvidencijaSati, jer upravo ona predstavlja osnovu za izračunavanje mera koje se kasnije analiziraju u kocki produktivnosti.',
            'Takođe je uočljivo da je finansijski tok predstavljen tabelom Faktura, dok tehnološki aspekt poslovanja ulazi u model preko spojne tabele ProjekatTehnologija. Na taj način relacijska baza obuhvata više poslovnih perspektiva neophodnih za celovito skladište podataka.'
        ],
        width=6.0,
    )

    add_heading(doc, '4.4 Detaljna implementacija relacione baze')
    for item in paragraphs("""
    Implementacija u SQL Server okruženju zasniva se na osam tabela koje obuhvataju referentne šifarnike, poslovne entitete i transakcione zapise. Iako se rad bavi skladištem podataka, pažljivo projektovana operativna baza predstavlja neophodan preduslov za kvalitetnu analitičku obradu, jer upravo iz nje potiču svi podaci koji se kasnije agregiraju.

    Tabela Odeljenje sadrži kolone idOdeljenja INT IDENTITY(1,1) PRIMARY KEY i nazivOdeljenja NVARCHAR(100) NOT NULL. Njen zadatak je da obezbedi organizacionu klasifikaciju zaposlenih i da omogući agregiranje radnog opterećenja na nivou organizacionih jedinica.

    Tabela Zaposleni sadrži identifikator idZaposlenog INT IDENTITY PRIMARY KEY, lične atribute ime NVARCHAR(50) i prezime NVARCHAR(50), opcionu kolonu pozicija NVARCHAR(100), strani ključ idOdeljenja koji referencira tabelu Odeljenje, datumZaposlenja DATE i satnica DECIMAL(10,2). Ovakva struktura omogućava da se individualni radni angažman poveže i sa kadrovskom i sa troškovnom perspektivom.

    Tabela Klijent sadrži kolone idKlijenta INT IDENTITY PRIMARY KEY, nazivKompanije NVARCHAR(200) NOT NULL, kontaktOsoba NVARCHAR(100), grad NVARCHAR(100) i drzava NVARCHAR(100). Ova tabela ima dimenzioni karakter i omogućava segmentaciju prihoda prema tržišnom i geografskom kriterijumu.

    Tabela Projekat sadrži kolone idProjekta INT IDENTITY PRIMARY KEY, nazivProjekta NVARCHAR(200) NOT NULL, opis NVARCHAR(MAX), datumPocetka DATE, datumZavrsetka DATE, budzet DECIMAL(12,2), status NVARCHAR(50) i strani ključ idKlijenta. Njen značaj je dvostruk: projekat predstavlja i poslovni objekat operativnog rada i ključnu dimenziju kasnije OLAP analize.

    Tabela Tehnologija sadrži identifikator idTehnologije INT IDENTITY PRIMARY KEY, nazivTehnologije NVARCHAR(100) NOT NULL i kategorija NVARCHAR(50). Ona omogućava da se portfelj projekata interpretira i sa tehnološkog aspekta, što je od posebnog značaja za planiranje kompetencija i razvojnih pravaca kompanije.

    Spojna tabela ProjekatTehnologija sadrži kolone idProjekta i idTehnologije, obe definisane kao strani ključevi, uz složeni primarni ključ nad oba atributa. Time je realizovana veza više-prema-više između projekata i tehnologija, što sprečava redundantno evidentiranje tehnoloških izbora u okviru projekta.

    Tabela EvidencijaSati sadrži idEvidencije INT IDENTITY PRIMARY KEY, strane ključeve idZaposlenog i idProjekta, kolonu datum DATE NOT NULL, brojSati DECIMAL(5,2) NOT NULL i opis NVARCHAR(500). Pošto beleži pojedinačne radne događaje, ova tabela predstavlja glavnu transakcionu činjenicu iz koje se izvode produktivnost i trošak rada.

    Tabela Faktura sadrži idFakture INT IDENTITY PRIMARY KEY, strani ključ idProjekta, kolone mesec INT, godina INT, ukupnoSati DECIMAL(10,2), iznos DECIMAL(12,2) i status NVARCHAR(50). Kroz ovu tabelu se uvodi prihodna perspektiva sistema, pa ona čini osnovu druge analitičke činjenice u radu.
    """):
        add_paragraph(doc, item)

    add_table(
        doc,
        state,
        ['Tabela', 'Ključni atributi i ograničenja', 'Analitička uloga'],
        [
            ['Odeljenje', 'idOdeljenja PK, nazivOdeljenja NOT NULL', 'Organizaciona dimenzija'],
            ['Zaposleni', 'idZaposlenog PK, idOdeljenja FK, satnica DECIMAL', 'Kadrovska dimenzija i trošak'],
            ['Klijent', 'idKlijenta PK, nazivKompanije NOT NULL', 'Klijentska dimenzija'],
            ['Projekat', 'idProjekta PK, idKlijenta FK, budzet DECIMAL', 'Projektna dimenzija'],
            ['Tehnologija', 'idTehnologije PK, nazivTehnologije NOT NULL', 'Tehnološka dimenzija'],
            ['ProjekatTehnologija', 'Složeni PK, dva FK atributa', 'Veza projekta i tehnologije'],
            ['EvidencijaSati', 'idEvidencije PK, brojSati NOT NULL', 'Osnovna činjenica rada'],
            ['Faktura', 'idFakture PK, iznos DECIMAL, status', 'Osnovna činjenica prihoda'],
        ],
        'Pregled implementiranih tabela i njihove analitičke uloge',
        [
            'Pregled jasno pokazuje da relacijski model sadrži i dimenzione i transakcione elemente potrebne za kasniju izgradnju skladišta podataka. Time je ostvarena čvrsta veza između operativne baze i analitičkog sloja.',
            'Analitička vrednost modela proizlazi iz toga što su mere i dimenzije već prisutne u operativnoj strukturi, pa ih je moguće relativno direktno prevesti u OLAP okruženje.'
        ]
    )

    add_heading(doc, '4.5 ETL proces punjenja skladišta')
    for item in paragraphs("""
    Iako u radu nije primenjen zaseban komercijalni ETL alat, logika punjenja skladišta podataka jasno je definisana i može se razložiti na više međusobno povezanih koraka. Prvi korak predstavlja izdvajanje podataka iz relacijskih tabela SQL Server baze, pri čemu se koriste stabilne šifarske tabele i transakcione tabele koje opisuju radne sate i fakture.

    Drugi korak čini transformacija podataka kroz SQL logiku implementiranu u pogledima. U tom koraku se spajaju podaci iz više izvora, izvode atributi Godina, Mesec i Kvartal, formira puna oznaka zaposlenog, povezuju projekti sa klijentima i izračunava mera TrošakRada kao proizvod broja sati i satnice. Time se iz transakcionih zapisa izvodi poslovno smislen analitički kontekst.

    Treći korak odnosi se na punjenje analitičkog sloja. U prikazanom rešenju to je ostvareno tako što su pogledi iskorišćeni kao stabilan i dosledan izvor za Data Source View i OLAP kocke u SSAS okruženju. Time je izbegnuto direktno izlaganje sirovih operativnih tabela analitičkom servisu, a ujedno je povećana transparentnost pravila koja stoje iza svake mere i dimenzije.

    Važan aspekt ETL procesa predstavlja i kontrola kvaliteta. Pre svakog korišćenja podataka neophodno je proveriti da li postoje zapisi bez odgovarajućih stranih ključeva, negativne ili nerealne vrednosti sati, nelogične vremenske odrednice ili neusaglašeni nazivi. U akademskom okruženju rada ove provere su sprovođene kroz pomoćne SELECT upite i pregled rezultata pogleda, dok bi u produkcionom sistemu bile automatizovane kroz validacione procedure i logove izvršenja.
    """):
        add_paragraph(doc, item)

    add_heading(doc, '4.6 Indeksiranje i optimizacija performansi')
    for item in paragraphs("""
    Performanse analitičkog sistema ne zavise samo od OLAP servisa, već i od kvaliteta relacione osnove iz koje se podaci preuzimaju. Primarni ključevi u svim tabelama predstavljaju osnovu fizičke organizacije podataka, dok strani ključevi obezbeđuju da spajanje tabela bude logički dosledno i tehnički efikasno. U realnim sistemima preporučljivo je dodatno indeksirati atribute koji se često koriste u JOIN uslovima i vremenskim filtrima.

    U slučaju ovog modela posebno su važni indeksi nad kolonama idOdeljenja u tabeli Zaposleni, idKlijenta u tabeli Projekat, idZaposlenog i idProjekta u tabeli EvidencijaSati, kao i idProjekta, godina i mesec u tabeli Faktura. Takvi indeksi ubrzavaju i operativne upite i pripremu podataka za analitičke poglede, posebno kada obim podataka raste kroz duži vremenski period.

    Optimizacija performansi obuhvata i pažljivo definisanje pogleda. Iz pogleda treba ukloniti nepotrebne proračune koji bi se mogli skuplje izvršavati pri svakom pozivu, a izvedene atribute treba definisati tako da budu transparentni i predvidivi. U radu je odabrana razumna ravnoteža: izračunavanje mera kao što je TrošakRada realizovano je u pogledu, jer predstavlja centralni analitički pokazatelj i logično pripada sloju pripreme podataka.

    Dodatno unapređenje u produkcionom okruženju moglo bi uključiti particionisanje činjenica po vremenu, inkrementalno procesiranje kocke i detaljnije upravljanje agregacijama u SSAS-u. Takve tehnike nisu nužne za obim podataka korišćen u radu, ali su teorijski značajne jer pokazuju kako ista arhitektura može da se razvija ka većem broju korisnika i znatno većem obimu podataka.
    """):
        add_paragraph(doc, item)

    add_heading(doc, '4.7 Pogledi i priprema podataka za OLAP')
    for item in paragraphs("""
    Pogledi predstavljaju prelazni sloj između operativne baze i OLAP okruženja. Njihova uloga nije samo tehnička, već i metodološka: kroz poglede se eksplicitno dokumentuje koje se tabele spajaju, koje se izvedene vrednosti računaju i koji skup atributa se proglašava relevantnim za analitičku obradu.

    U radu su formirana tri ključna pogleda: Analiza_troska_rada, Analiza_prihoda i Analiza_tehnologija. Prvi objedinjavanjem evidencije rada, zaposlenih, odeljenja, projekata i klijenata priprema činjenice za kocku produktivnosti. Drugi objedinjavanjem faktura, projekata i klijenata priprema činjenice za kocku prihoda. Treći pogled omogućava pregled tehnološke strukture projekata i predstavlja dopunski izvor za kvalitativnu interpretaciju rezultata.
    """):
        add_paragraph(doc, item)

    add_image(
        doc,
        state,
        images['DB_diagram.png'],
        'Dijagram implementirane baze u SQL Server okruženju',
        [
            'Dijagram baze pokazuje kako su poslovni entiteti povezani na fizičkom nivou i na koji način strani ključevi usmeravaju tok pripreme podataka za analitiku. Vizuelni prikaz odnosa posebno olakšava razumevanje veze između operativnih tabela i pogleda koji se nad njima formiraju.',
            'Za potrebe rada ovakav prikaz služi i kao validacioni alat, jer omogućava proveru da li svaka mera ima jasan put do svojih dimenzija. Time se smanjuje rizik od pogrešnog modelovanja u kasnijim fazama projekta.'
        ],
        width=6.0,
    )
    add_image(
        doc,
        state,
        images['VIEW_results.png'],
        'Rezultat pogleda Analiza_troska_rada',
        [
            'Prikaz rezultata pogleda Analiza_troska_rada potvrđuje da su radni sati uspešno povezani sa zaposlenima, projektima, klijentima i odeljenjima, kao i da je mera troška rada pravilno izvedena. Time je formirana stabilna osnova za izgradnju kocke cbProduktivnost.',
            'Pogled je posebno važan zato što uklanja potrebu da se u OLAP sloju ponavljaju kompleksni relacijski JOIN upiti. Analitički servis dobija već pripremljene i semantički obogaćene podatke, što ubrzava i razvoj i kasnije korišćenje sistema.'
        ],
        width=6.0,
    )
    add_image(
        doc,
        state,
        images['SELECT_queries_1.png'],
        'Pomoćni SELECT upiti za proveru integriteta i sadržaja podataka',
        [
            'Pomoćni SELECT upiti imaju važnu ulogu u verifikaciji da li su veze između tabela pravilno uspostavljene i da li izvedene vrednosti odgovaraju očekivanoj poslovnoj logici. Ovakve provere predstavljaju važan korak između implementacije i konačne analitičke eksploatacije podataka.',
            'U radu su SELECT upiti korišćeni za proveru ispravnosti spajanja, filtriranja i agregiranja, čime je obezbeđeno da svi kasniji OLAP rezultati počivaju na tačnoj relacijskoj osnovi.'
        ],
        width=6.0,
    )
    add_image(
        doc,
        state,
        images['SELECT_queries_2.png'],
        'Rezultat pripreme prihodnih podataka za pogled Analiza_prihoda',
        [
            'Snimak prihodnog pregleda pokazuje kako se fakture povezuju sa projektima i klijentima, kao i na koji način vremenski atributi postaju deo analitičke logike. Ovaj korak je presudan za kasnije poređenje prihoda po projektima, klijentima, godinama i kvartalima.',
            'Kroz ovakav pogled finansijski podaci dobijaju jasan poslovni kontekst. Time se omogućava da prihodna analiza ne bude izolovana od operativnog okruženja, već direktno povezana sa projektima i resursima koji su prihod generisali.'
        ],
        width=6.0,
    )
    add_image(
        doc,
        state,
        images['SELECT_queries_3.png'],
        'Rezultat pripreme tehnoloških podataka za analitički pregled projekata',
        [
            'Pogled Analiza_tehnologija dopunjuje kvantitativne analize kvalitativnim uvidom u to koje su tehnologije angažovane na pojedinim projektima. Ovakav prikaz je značajan zato što menadžmentu omogućava da poveže prihodnu ili troškovnu sliku sa tehnološkim profilom poslovanja.',
            'Na osnovu ovakvih podataka moguće je procenjivati tržišnu orijentaciju portfolija, stepen standardizacije razvojnih alata i pravce u kojima kompanija treba dalje da razvija kompetencije zaposlenih.'
        ],
        width=6.0,
    )

    add_section(doc, '5. Projektovanje OLAP kocke', """
    Nakon pripreme relacijskog sloja sledeći korak predstavlja izgradnju višedimenzionalnog modela u SSAS okruženju. U radu su definisane dve OLAP kocke: cbProduktivnost i cbPrihodi. Ovakva podela odgovara činjenici da se analiziraju dva međusobno povezana, ali konceptualno različita poslovna procesa: radni angažman i finansijska realizacija.

    Projektovanje OLAP kocke podrazumeva izbor dimenzija, definisanje mera, organizovanje činjenica, uspostavljanje hijerarhija i procesiranje podataka. Cilj nije samo tehničko formiranje kocke, već i stvaranje semantički jasnog analitičkog modela koji omogućava brzo odgovaranje na poslovna pitanja.
    """)

    add_section(doc, '5.1 Dimenzije i hijerarhije', """
    Kocka cbProduktivnost zasniva se na dimenzijama Zaposleni, Odeljenje, Projekat, Klijent i Vreme. Ovaj izbor omogućava da se svaki radni zapis posmatra na individualnom, organizacionom, projektnom, tržišnom i vremenskom nivou. Time se obezbeđuje velika fleksibilnost analize bez menjanja same strukture podataka.

    Kocka cbPrihodi koristi dimenzije Projekat, Klijent i Vreme, jer su upravo one najrelevantnije za prihodnu interpretaciju. Iako prihodna kocka sadrži manje dimenzija, ona ostaje dovoljno bogata da podrži i projektni i klijentski i vremenski pogled na poslovanje.

    Poseban značaj imaju hijerarhije u vremenskoj dimenziji. Veza godina–kvartal–mesec omogućava postupni prelaz od grubog strateškog pregleda ka detaljnijoj vremenskoj slici. Takva hijerarhijska organizacija predstavlja osnovu za OLAP operacije roll-up i drill-down koje se kasnije koriste u analizi.
    """)

    add_section(doc, '5.2 Mere i kalkulisani pokazatelji', """
    Mere predstavljaju numeričko jezgro OLAP analize. U kocki cbProduktivnost glavne mere su Broj sati, Trošak rada, Satnica i Broj evidencija. Svaka od njih pokriva poseban aspekt upravljanja resursima: obim rada, finansijsko opterećenje, strukturu angažovanih profila i učestalost evidentiranja posla.

    U kocki cbPrihodi definisane su mere Iznos, Ukupno sati, Budžet i Broj faktura. Njihova kombinacija omogućava ne samo merenje nominalne vrednosti naplate, već i poređenje prihoda sa obimom angažovanja i sa planiranim budžetom projekta. Takav višeslojni pristup čini OLAP analizu znatno bogatijom od prostog spiska faktura.

    Posebna vrednost mera ogleda se u tome što su one definisane na osnovu jasno dokumentovanih pravila. U radu su mere zasnovane na izvornoj relacijskoj logici i pripremljenim pogledima, pa postoji puna sledljivost između transakcionog zapisa i agregirane vrednosti prikazane u kocki ili pivot izveštaju.
    """)

    add_section(doc, '5.3 Šema zvezde i organizacija činjenica', """
    Analitički model u radu dominantno prati logiku šeme zvezde. U središtu su činjenice izvedene iz evidencije rada i faktura, dok su oko njih raspoređene dimenzije koje obezbeđuju poslovni kontekst. Ovakva organizacija olakšava agregiranje podataka, smanjuje složenost upita i čini model pristupačnim analitičarima i menadžerima.

    Za kocku cbProduktivnost činjenice potiču iz pogleda Analiza_troska_rada, koji objedinjeno prikazuje zaposlenog, odeljenje, projekat, klijenta i vreme, uz numeričke mere. Za kocku cbPrihodi činjenice potiču iz pogleda Analiza_prihoda, gde se finansijski pokazatelji povezuju sa projektom, klijentom i vremenom. Na taj način dve kocke ostaju zasebne, ali semantički povezane.

    Upravo takva organizacija potvrđuje da je skladište podataka u radu zasnovano na jasnoj i proverljivoj analitičkoj logici. U praksi to znači da korisnik može da menja perspektivu posmatranja bez promene izvora podataka, jer su veze između mera i dimenzija trajno ugrađene u model.
    """)

    add_section(doc, '5.4 Kreiranje i procesiranje kocke', """
    Kreiranje OLAP kocke u SSAS-u obuhvata više tehničkih koraka: definisanje izvora podataka, formiranje Data Source View sloja, izbor tabela koje predstavljaju činjenice i dimenzije, definisanje hijerarhija i mera, kao i podešavanje agregacija. Iako svaki od ovih koraka ima tehničke detalje, njihova zajednička svrha jeste uspostavljanje doslednog analitičkog modela nad relacijskim podacima.

    Nakon logičkog definisanja modela sledi procesiranje kocke. Procesiranje predstavlja učitavanje podataka iz relacijskog izvora u višedimenzionalnu strukturu spremnu za analitičke upite. Ovaj korak je presudan, jer tek tada mere i dimenzije postaju operativno dostupne za MDX upite, pivot tabele i interaktivnu analizu.

    U radu je procesiranje imalo i verifikacionu ulogu. Svaki uspešan prolaz kroz procesiranje potvrđuje da su veze između dimenzija i činjenica pravilno podešene, da su izvori podataka dostupni i da nema logičkih konflikata u definisanom modelu. Na taj način se tehnički uspeh procesa direktno povezuje sa validnošću celokupnog analitičkog rešenja.
    """)

    add_section(doc, '6. OLAP analiza i interpretacija rezultata', """
    Nakon izgradnje kocke analitički fokus prelazi na ispitivanje rezultata i tumačenje poslovnih implikacija. U ovom poglavlju prikazane su osnovne OLAP operacije, odabrani MDX upiti i osam pivot tabela koje sumiraju podatke o produktivnosti i prihodima. Svrha poglavlja nije samo prikaz numeričkih vrednosti, već i demonstracija načina na koji višedimenzionalna analiza menja kvalitet poslovnog uvida.

    U poređenju sa klasičnim tabelarnim izveštajima, OLAP pristup omogućava da se isti skup podataka veoma brzo sagleda iz više perspektiva. Time se ubrzava analiza, smanjuje potreba za ponovnim pisanjem upita i povećava mogućnost da se otkriju obrasci koji u ravnim izveštajima ostaju skriveni.
    """)

    add_section(doc, '6.1 OLAP operacije', """
    Operacija roll-up predstavlja agregiranje podataka na viši nivo hijerarhije. Na primer, pojedinačne evidencije rada mogu se sabrati po projektu, zatim po odeljenju ili po godini. Ovakva operacija je važna kada je cilj da se iz velikog broja detaljnih zapisa dobije strateški pregled poslovanja.

    Drill-down predstavlja obrnut proces: od agregiranog nivoa prelazi se ka detaljnijem prikazu. Ako godišnji prihod pokaže odstupanje, analiza se može spustiti na nivo kvartala, projekata ili čak pojedinačnih faktura. Time se brzo identifikuje izvor promene bez napuštanja istog analitičkog okruženja.

    Slice je operacija izdvajanja jednog preseka kocke prema određenoj vrednosti dimenzije. Na primer, moguće je posmatrati samo 2025. godinu ili samo jedan projekat, čime se analiza fokusira na uži, ali relevantan segment podataka. Dice ide korak dalje i kombinuje više uslova filtriranja, pa se mogu istovremeno posmatrati, na primer, određeni kvartali, određeni klijenti i određena odeljenja.

    Pivot predstavlja rotiranje analitičke perspektive, odnosno zamenu redova i kolona kako bi se ista mera posmatrala iz drugog ugla. Ova operacija je izuzetno korisna u radu sa pivot tabelama, jer korisniku omogućava da bez promene izvora odmah pređe sa projektne na klijentsku ili vremensku perspektivu.

    Zajednički efekat ovih operacija ogleda se u tome što analitički proces postaje iterativan i istraživački. Korisnik ne mora unapred da zna svako pitanje koje će postaviti sistemu, već kroz interakciju sa kockom postepeno produbljuje razumevanje poslovnih obrazaca. Upravo u tome se ogleda jedna od najvećih prednosti OLAP pristupa.
    """)

    add_heading(doc, '6.2 MDX upiti i njihovo tumačenje')
    add_paragraph(doc, 'MDX predstavlja standardni jezik za upit nad višedimenzionalnim strukturama. Za potrebe rada korišćeni su upiti pripremljeni u posebnoj datoteci, a u nastavku su prikazani svi upiti sa sažetim objašnjenjem njihove analitičke vrednosti.')
    for query in mdx_queries:
        add_paragraph(doc, f'Upit {query["number"]}: {query["title"]}', bold=True, first_line=0)
        add_code_block(doc, query['query'])
        for explanation in MDX_EXPLANATIONS[query['number']]:
            add_paragraph(doc, explanation)

    add_heading(doc, '6.3 Pivot analize produktivnosti')
    add_paragraph(doc, 'Pivot tabele izvedene iz kocke cbProduktivnost prikazuju kako se radni angažman i troškovi raspoređuju po organizacionim, projektnim, individualnim i vremenskim kriterijumima. U radu su uključena četiri pregleda produktivnosti, jer zajedno pružaju dovoljno široku i međusobno uporedivu sliku korišćenja resursa.')
    for sheet_name in ['Produktivnost_Odeljenja', 'Produktivnost_Projekti', 'Produktivnost_Zaposleni', 'Produktivnost_Godine']:
        meta = PIVOT_METADATA[sheet_name]
        table_data = pivot_tables[sheet_name]
        add_table(doc, state, table_data['headers'], table_data['rows'], meta['title'], meta['analysis'])

    add_heading(doc, '6.4 Pivot analize prihoda')
    add_paragraph(doc, 'Druga grupa pivot tabela zasniva se na kocki cbPrihodi i prikazuje raspodelu finansijskih rezultata prema klijentima, projektima i vremenu. Kombinovanjem ovih pregleda moguće je povezati nominalnu naplatu sa obimom rada i planiranim budžetima, čime se dobija znatno pouzdaniji uvid u poslovni rezultat od prostog pregleda faktura.')
    for sheet_name in ['Prihodi_Klijenti', 'Prihodi_Projekti', 'Prihodi_Godine', 'Prihodi_Kvartali']:
        meta = PIVOT_METADATA[sheet_name]
        table_data = pivot_tables[sheet_name]
        add_table(doc, state, table_data['headers'], table_data['rows'], meta['title'], meta['analysis'])

    add_section(doc, '6.5 Prednosti OLAP pristupa u odnosu na klasično izveštavanje', """
    Klasično izveštavanje najčešće podrazumeva unapred definisane i relativno statične tabele ili grafikone. Takvi izveštaji mogu biti korisni za rutinsko praćenje, ali postaju ograničavajući kada se pojavi potreba za brzim menjanjem perspektive ili kombinovanjem više analitičkih kriterijuma. Svaka nova poslovna dilema tada zahteva dodatni upit, novu proceduru ili novi izveštaj.

    OLAP pristup rešava ovaj problem tako što iste podatke organizuje u višedimenzionalni model. Korisnik može da vrši agregiranje, filtriranje, drill-down i pivot bez redefinisanja relacione logike. Time se značajno skraćuje vreme potrebno za analitičko istraživanje i povećava se mogućnost samostalnog rada poslovnih korisnika.

    Još jedna važna prednost ogleda se u doslednosti pokazatelja. Kada se mere i dimenzije jednom definišu u kocki, svi korisnici polaze od istog semantičkog modela. Time se smanjuje rizik da različiti odeljenjski izveštaji daju međusobno kontradiktorne rezultate zbog različitih formula ili različitog načina spajanja tabela.

    U kontekstu ovog rada OLAP je omogućio da se produktivnost i prihodi tumače kroz isti skup poslovnih dimenzija, uz veliku fleksibilnost analize. Takva sposobnost čini OLAP značajno pogodnijim od klasičnog izveštavanja onda kada organizacija želi ne samo da prikaže prošlost, već i da razume obrasce koji oblikuju buduće odluke.
    """)

    add_section(doc, '7. Zaključak', """
    Na osnovu sprovedene analize može se zaključiti da skladište podataka predstavlja adekvatnu i efikasnu osnovu za objedinjavanje operativnih i analitičkih informacija o poslovanju IT kompanije. U radu je pokazano da se relativno kompaktan relacijski model može, uz pažljivo projektovanje pogleda i OLAP kocki, transformisati u okruženje koje podržava višedimenzionalno odlučivanje i pruža znatno dublji uvid od klasičnih transakcionih pregleda.

    Rezultati pivot i MDX analiza potvrđuju da je moguće jasno identifikovati obrasce produktivnosti, troškova rada i prihoda. Uočeno je da razvojno odeljenje nosi najveći deo radnog opterećenja, da su pojedini projekti i zaposleni izrazito dominantni u ukupnom angažovanju, kao i da prihodi nisu ravnomerno raspoređeni između klijenata, projekata i vremenskih perioda. Takvi nalazi imaju neposrednu upravljačku vrednost, jer omogućavaju preciznije planiranje resursa, budžeta i odnosa sa klijentima.

    Poseban doprinos rada ogleda se u tome što su teorijski koncepti, kao što su ETL, dimenzionalno modeliranje, SCD i OLAP operacije, povezani sa konkretnim implementacionim odlukama. Time je potvrđeno da uspešno skladište podataka ne nastaje samo kao tehnički artefakt, već kao rezultat metodološki doslednog povezivanja poslovnih zahteva, relacijskog modela i analitičke logike.

    Ograničenja rada prvenstveno proizlaze iz akademskog karaktera studije slučaja. Obim podataka je ograničen, ETL logika nije automatizovana kroz zasebnu platformu, a napredne tehnike poput pune implementacije sporih promenljivih dimenzija, inkrementalnog učitavanja i kompleksnih bezbednosnih politika nisu razrađene do produkcionog nivoa. Ipak, ova ograničenja ne umanjuju vrednost rada, već jasno određuju njegov opseg i fazu razvoja sistema.

    Kao pravci budućeg razvoja izdvajaju se uvođenje detaljnijih vremenskih hijerarhija, primena SCD tipa 2 nad ključnim dimenzijama, automatizacija ETL procesa, proširenje skupa mera pokazateljima profitabilnosti i marže, kao i povezivanje skladišta podataka sa naprednim analitičkim metodama. Na taj način prikazano rešenje može prerasti iz edukativnog i demonstracionog modela u robustan sistem poslovne inteligencije namenjen realnom upravljanju IT kompanijom.
    """)

    add_heading(doc, '8. Literatura')
    for index, item in enumerate(LITERATURE, start=1):
        add_paragraph(doc, f'{index}. {item}', align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, space_after=0)

    if OUTPUT_PATH.exists():
        OUTPUT_PATH.unlink()
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def verify_document(path):
    doc = Document(str(path))
    text_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    char_count = sum(len(text) for text in text_paragraphs)
    estimated_lines = sum(max(1, math.ceil(len(text) / 78)) for text in text_paragraphs)
    estimated_lines += sum(len(table.rows) + 1 for table in doc.tables)

    with ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith('word/media/')]

    return {
        'paragraphs': len(doc.paragraphs),
        'nonempty_paragraphs': len(text_paragraphs),
        'tables': len(doc.tables),
        'images': len(media),
        'estimated_lines': estimated_lines,
        'characters': char_count,
        'valid_docx': True,
        'exists': path.exists(),
        'size': path.stat().st_size,
    }


if __name__ == '__main__':
    output = build_document()
    verification = verify_document(output)
    print(f'Kreiran dokument: {output}')
    for key, value in verification.items():
        print(f'{key}: {value}')
