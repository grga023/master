# -*- coding: utf-8 -*-
from __future__ import annotations

import ast
import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / 'Seminarski_rad_DW_ITKompanije.docx'
SQL_PATH = BASE_DIR / 'ITKompanijaDW_kreiranje.sql'
MDX_PATH = BASE_DIR / 'MDX_upiti_za_kocku.mdx'

IMAGES = {
    'cdm': BASE_DIR / 'diagrami' / 'CDM.png',
    'pdm': BASE_DIR / 'diagrami' / 'PDM.png',
    'db': BASE_DIR / 'diagrami' / 'DB_diagram.png',
}

PIVOT_TABLES = [
    {
        'title': 'Produktivnost po odeljenjima',
        'headers': ['Odeljenje', 'Broj sati', 'Trošak rada', 'Satnica'],
        'rows': [
            ['Design', '16', '416', '78'],
            ['Development', '177', '4922', '664'],
            ['Management', '22', '770', '175'],
            ['QA', '52', '1276', '172'],
        ],
        'analysis': 'Pivot analiza potvrđuje da se najveći deo radnog opterećenja i troška rada nalazi u razvojnom odeljenju, što je očekivano zbog dominantne uloge programerskih aktivnosti u realizaciji projekata. Istovremeno, prisustvo merljivog troška u QA, dizajnu i menadžmentu pokazuje da analitički model obuhvata i podržavajuće funkcije koje neposredno utiču na kvalitet isporuke i koordinaciju rada.',
    },
    {
        'title': 'Produktivnost po projektima',
        'headers': ['Projekat', 'Broj sati', 'Trošak rada'],
        'rows': [
            ['Analytics Dashboard', '38', '1070'],
            ['Cloud Migration', '49', '1509'],
            ['CRM sistem', '57', '1681'],
            ['E-Commerce platforma', '67', '1796'],
            ['Mobile Banking App', '56', '1328'],
        ],
        'analysis': 'Pregled po projektima ukazuje da E-Commerce platforma i CRM sistem zahtevaju najveći angažman resursa, dok projekat Analytics Dashboard ima manji obim rada i niži trošak. Ovakva struktura pomaže da se proceni gde se troše ključni kapaciteti i kako se projektni portfelj odražava na operativno opterećenje kompanije.',
    },
    {
        'title': 'Produktivnost po zaposlenima',
        'headers': ['Zaposleni', 'Broj sati', 'Trošak rada', 'Broj evidencija'],
        'rows': [
            ['Ana Stojanović', '42', '1050', '6'],
            ['Ivan Marković', '30', '660', '4'],
            ['Jelena Nikolić', '16', '416', '3'],
            ['Maja Ilić', '22', '616', '3'],
            ['Marko Petrović', '54', '1620', '7'],
            ['Milica Đorđević', '21', '378', '3'],
            ['Nikola Todorović', '37', '1184', '5'],
            ['Petar Pavlović', '23', '690', '3'],
            ['Stefan Jovanović', '22', '770', '5'],
        ],
        'analysis': 'Najveći pojedinačni doprinos ostvaruju Marko Petrović, Ana Stojanović i Nikola Todorović, što ukazuje na koncentraciju ključnog stručnog rada u ograničenom broju resursa. Kolona Broj evidencija dodatno pokazuje da se intenzitet angažovanja ne ogleda samo u zbiru sati, već i u kontinuitetu učešća zaposlenih kroz više radnih zapisa.',
    },
    {
        'title': 'Produktivnost po godinama',
        'headers': ['Godina', 'Broj sati', 'Trošak rada'],
        'rows': [
            ['2024', '180', '4805'],
            ['2025', '87', '2579'],
        ],
        'analysis': 'Vremenska agregacija pokazuje da je u 2024. godini evidentiran veći fond sati i veći trošak rada nego u 2025. godini. Takav odnos je logičan jer 2024. obuhvata više projekata u intenzivnoj razvojnoj fazi, dok 2025. prikazuje nastavak i završne radove na delu portfelja.',
    },
    {
        'title': 'Prihodi po klijentima',
        'headers': ['Klijent', 'Iznos', 'Ukupno sati', 'Broj faktura'],
        'rows': [
            ['AppDev Inc', '10200', '340', '2'],
            ['CloudNet d.o.o', '18000', '600', '3'],
            ['DataSys GmbH', '27300', '910', '5'],
            ['TechCorp Solutions', '23650', '885', '5'],
            ['WebPro Ltd', '26550', '885', '5'],
        ],
        'analysis': 'Analiza prihoda po klijentima pokazuje da su DataSys GmbH, WebPro Ltd i TechCorp Solutions najznačajniji izvori naplate. Poređenje iznosa, sati i broja faktura potvrđuje da se poslovna vrednost ne može proceniti samo kroz radni angažman, već i kroz dinamiku fakturisanja i ugovorene komercijalne uslove.',
    },
    {
        'title': 'Prihodi po projektima',
        'headers': ['Projekat', 'Iznos', 'Budžet', 'Broj faktura'],
        'rows': [
            ['Analytics Dashboard', '10200', '120000', '2'],
            ['Cloud Migration', '18000', '285000', '3'],
            ['CRM sistem', '27300', '600000', '5'],
            ['E-Commerce platforma', '23650', '425000', '5'],
            ['Mobile Banking App', '26550', '750000', '5'],
        ],
        'analysis': 'Kada se ostvaren prihod uporedi sa planiranim budžetom, vidi se da pojedini projekti još uvek imaju značajan prostor za dalju monetizaciju. Ovakav pregled je posebno važan za menadžersko praćenje odnosa između realizovanog finansijskog efekta i ukupnog poslovnog potencijala projekta.',
    },
    {
        'title': 'Prihodi po godinama',
        'headers': ['Godina', 'Iznos', 'Broj faktura'],
        'rows': [
            ['2024', '77500', '15'],
            ['2025', '28200', '5'],
        ],
        'analysis': 'Na godišnjem nivou jasno se uočava da je 2024. godina finansijski znatno snažnija od 2025. godine, kako po ukupnom iznosu tako i po broju faktura. Ovaj rezultat pokazuje da je najveći deo naplate koncentrisan u prvoj posmatranoj godini, što je značajno za planiranje novčanih tokova i budućih prodajnih aktivnosti.',
    },
    {
        'title': 'Prihodi po kvartalima',
        'headers': ['Kvartal', 'Iznos', 'Ukupno sati'],
        'rows': [
            ['Q1', '48300', '1675'],
            ['Q2', '31150', '1070'],
            ['Q3', '21300', '710'],
            ['Q4', '4950', '165'],
        ],
        'analysis': 'Kvartalna raspodela prihoda otkriva izraženu sezonalnost: prvi kvartal beleži najviši prihod, dok je četvrti kvartal najslabiji. Analitički značaj ovog prikaza ogleda se u tome što omogućava detaljnije vremensko sagledavanje naplate i lakše usklađivanje operativnih kapaciteta sa prihodnim ciklusima.',
    },
]

LITERATURE = [
    'Inmon, W. H. (2005). Building the Data Warehouse. Wiley.',
    'Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling. Wiley.',
    'Ponniah, P. (2010). Data Warehousing Fundamentals for IT Professionals. Wiley.',
    'Golfarelli, M., & Rizzi, S. (2009). Data Warehouse Design: Modern Principles and Methodologies. McGraw-Hill.',
    'Chaudhuri, S., & Dayal, U. (1997). An Overview of Data Warehousing and OLAP Technology. ACM SIGMOD Record, 26(1), 65-74.',
    'Vassiliadis, P. (2010). A Survey of Extract-Transform-Load Technology. International Journal of Data Warehousing and Mining, 5(3), 1-27.',
    'Kimball, R., Reeves, L., Ross, M., & Thornthwaite, W. (1998). The Data Warehouse Lifecycle Toolkit. Wiley.',
    'Elmasri, R., & Navathe, S. B. (2016). Fundamentals of Database Systems. Pearson.',
    'Harrington, J. L. (2008). Relational Database Design and Implementation. Morgan Kaufmann.',
    'Microsoft. (2024). SQL Server Analysis Services Documentation. Microsoft Learn.',
    'Oracle. (2024). Oracle SQL Developer Data Modeler User Guide. Oracle Documentation.',
    'Abelló, A., Samos, J., & Saltor, F. (2006). YAM²: A Multidimensional Conceptual Model Extending UML. Information Systems, 31(6), 541-567.',
]

SECTION_TEXT = {
    'uvod': """
    Savremeno poslovanje IT kompanija generiše veliku količinu operativnih podataka o projektima, zaposlenima, klijentima, tehnologijama, utrošenim radnim satima i fakturisanim uslugama. Iako relacione baze podataka pouzdano podržavaju svakodnevne transakcije, njihova struktura nije optimalna za složena menadžerska pitanja koja zahtevaju pregled istorijskih trendova, sabiranje podataka i analizu iz više perspektiva.

    U radu je prikazan razvoj skladišta podataka nad modelom poslovanja IT kompanije koja paralelno realizuje više projekata za različite klijente. Polazna tačka predstavlja transakcioni model sa tabelama za odeljenja, zaposlene, klijente, projekte, tehnologije, evidenciju rada i fakture. Nad tim modelom definisani su analitički pogledi i dve OLAP kocke, sa ciljem da se obezbedi dosledno praćenje produktivnosti i prihoda.

    Poseban značaj rada ogleda se u povezivanju teorijskih principa skladišta podataka sa praktičnom implementacijom. Prikazano je kako se iz operativnog sloja izdvajaju relevantni atributi, kako se organizuju činjenice i dimenzije i na koji način OLAP analiza omogućava brzo sagledavanje troškova rada, opterećenja zaposlenih, profitabilnosti projekata i vremenske dinamike fakturisanja.

    Cilj rada jeste da se na pregledan i akademski utemeljen način prikaže potpuni tok izgradnje rešenja: od teorijskog okvira i metodologije, preko projektovanja i realizacije baze, do kreiranja OLAP kocki, izvršavanja MDX upita i interpretacije rezultata dobijenih kroz pivot tabele. Takav pristup potvrđuje da skladište podataka predstavlja pouzdanu osnovu za podršku odlučivanju u savremenim IT organizacijama.
    """,
    'teorija1': """
    Skladište podataka predstavlja subjektno orijentisanu, integrisanu, vremenski promenljivu i relativno nepromenljivu zbirku podataka namenjenu podršci odlučivanju. Njegova osnovna funkcija nije obrada pojedinačnih transakcija, već obezbeđivanje pouzdane osnove za analizu istorijskih obrazaca poslovanja. U tom smislu ono se razlikuje od operativne baze, koja je optimizovana za unos, izmene i održavanje integriteta svakodnevnih poslovnih događaja.

    Ključna vrednost skladišta podataka ogleda se u integraciji više poslovnih entiteta u jedinstven analitički okvir. U posmatranom domenu to znači da se evidencija rada, projekti, klijenti i fakture ne analiziraju izolovano, već kroz zajedničke dimenzije koje omogućavaju poređenje rezultata po vremenu, organizacionim jedinicama i poslovnim inicijativama. Takva integracija je preduslov za konzistentne izveštaje i za smanjenje razlika između različitih interpretacija istih podataka.

    Vremenska komponenta skladišta podataka ima poseban značaj zato što menadžerske odluke retko zavise samo od aktuelnog stanja. Potrebno je utvrditi trendove, porediti periode i identifikovati odstupanja koja se ne vide u pojedinačnim zapisima. Zbog toga je u radu vremenska dimenzija uvedena i kroz operativne poglede i kroz OLAP kocke, čime je omogućeno posmatranje rezultata po godinama, mesecima i kvartalima.
    """,
    'teorija2': """
    Dimenzionalno modeliranje predstavlja dominantan pristup oblikovanju skladišta podataka kada je cilj brza i intuitivna analitika. Osnovu tog pristupa čine tabele činjenica i dimenzije. Tabela činjenica sadrži numeričke mere poslovanja, kao što su broj sati, trošak rada ili iznos fakture, dok dimenzije daju kontekst kroz entitete kao što su zaposleni, projekat, klijent, odeljenje i vreme.

    U praksi se najčešće koriste šema zvezde i šema pahulje. Za potrebe ovog rada primenjena je logika bliska šemi zvezde, jer omogućava pregledno mapiranje poslovnih pokazatelja na ograničen broj dimenzija i veoma dobro odgovara OLAP okruženju. Prednost ovakvog pristupa je u tome što krajnji korisnik lakše razume model, a agregacije nad podacima postaju jednostavnije i efikasnije.

    Dobro projektovan dimenzionalni model ima i metodološku vrednost. On primorava autora da precizno definiše poslovna pitanja, granice analize i značenje svake mere. Upravo zato su u radu odvojene dve analitičke celine: produktivnost zasnovana na evidenciji rada i prihodi zasnovani na fakturama. Takvo razdvajanje olakšava tumačenje rezultata i smanjuje rizik od mešanja pokazatelja različite prirode.
    """,
    'teorija3': """
    OLAP predstavlja skup tehnika koje omogućavaju višedimenzionalno posmatranje podataka kroz operacije kao što su roll-up, drill-down, slice i dice. Za razliku od klasičnih SQL izveštaja, OLAP kocka omogućava da se ista činjenica brzo sagleda na više nivoa agregacije i iz različitih uglova, bez potrebe da se za svako pitanje ponovo formuliše složen upit nad relacijskim slojem.

    U radu je OLAP pristup iskorišćen za dve centralne poslovne teme: produktivnost i prihode. Kroz kocku cbProduktivnost analizirani su broj sati, trošak rada, satnica i broj evidencija, dok je kocka cbPrihodi usmerena na iznos, budžet, ukupan broj sati i broj faktura. Time je demonstrirano kako OLAP podržava i operativno-taktički pogled na radne resurse i širu finansijsku perspektivu poslovanja.

    Posebna prednost OLAP analize ogleda se u tome što korisnik može prelaziti sa sažetog pregleda na detaljniji nivo bez promene metodološkog okvira. U jednoj analizi moguće je identifikovati da je određena godina bila dominantna po prihodu, a zatim odmah preći na kvartalni ili projekatni nivo da bi se utvrdio izvor tog rezultata. Takva fleksibilnost čini OLAP jednim od najvažnijih alata poslovne inteligencije.
    """,
    'teorija4': """
    ETL proces obuhvata izdvajanje podataka iz izvora, njihovu transformaciju i punjenje u analitički model. Iako u primeru rada svi podaci potiču iz jedne relacione baze, princip ETL-a ostaje jednako važan, jer se podaci moraju preoblikovati u strukturu pogodnu za agregacije i višedimenzionalnu obradu. To podrazumeva standardizaciju atributa, izračunavanje izvedenih mera i obogaćivanje vremenskim elementima.

    U implementiranom rešenju ETL logika je delimično realizovana kroz SQL poglede Analiza_troska_rada, Analiza_prihoda i Analiza_tehnologija. Ovi pogledi ne služe samo za prikaz podataka, već predstavljaju pripremni sloj koji spaja operativne tabele, izračunava nove kolone i oblikuje sadržaj za dalje analitičko korišćenje. Na taj način se smanjuje složenost OLAP modela i obezbeđuje veća doslednost rezultata.

    Kvalitet ETL procesa neposredno utiče na kredibilitet završne analize. Ako su dimenzije nekonzistentne, mere pogrešno definisane ili vremenski atributi nepotpuni, menadžerski zaključci mogu biti pogrešni. Zbog toga je u radu poseban akcenat stavljen na jasnu vezu između izvornog SQL modela, pogleda koji pripremaju podatke i OLAP kocki koje te podatke dalje agregiraju.
    """,
    'metodologija': """
    Metodologija rada zasnovana je na kombinaciji teorijske analize i praktične implementacije. Najpre su identifikovani ključni koncepti skladišta podataka, dimenzionalnog modeliranja, ETL procesa i OLAP analize, a zatim je definisan poslovni domen IT kompanije sa odgovarajućim entitetima i relacijama. Na toj osnovi je oblikovan transakcioni model koji predstavlja izvor podataka za analitički sloj.

    Za modelovanje i realizaciju rešenja korišćeni su SQL Server skriptovi za kreiranje baze, Oracle SQL Developer Data Modeler i PowerDesigner za izradu modela, kao i SQL Server Analysis Services za definisanje OLAP kocki. MDX upiti korišćeni su za proveru funkcionalnosti kocke i za formiranje analitičkih scenarija, dok su pivot tabele upotrebljene kao završni sloj interpretacije rezultata iz perspektive krajnjeg korisnika.

    Pristup rada može se opisati kao iterativan. Nakon kreiranja početnog modela i unosa test podataka, definisani su pogledi za analizu, provereni SQL upiti, a zatim su projektovane kocke cbProduktivnost i cbPrihodi. Dobijeni rezultati su interpretirani u skladu sa ciljem rada, uz nastojanje da se teorijski pojmovi neposredno povežu sa konkretnim poslovnim pokazateljima.
    """,
    'implementacija_uvod': """
    Implementacija skladišta podataka započinje definisanjem poslovnog domena i preciznim razgraničenjem između operativnog i analitičkog nivoa. U transakcionoj bazi modelovani su entiteti koji opisuju organizacionu strukturu kompanije, zaposlene, klijente, projekte, tehnologije, evidenciju utrošenih sati i fakturisanje. Ovakva struktura omogućava da se svakodnevni poslovni događaji beleže na normalizovan način, dok se analitički zahtevi preusmeravaju na posebno pripremljen sloj.

    U skladu sa tim, implementacija obuhvata tri međusobno povezana nivoa: konceptualni model koji daje poslovni pregled sistema, fizički model koji definiše konkretne tabele i veze, te realizaciju baze u SQL skriptu. Nakon toga su uvedeni pogledi koji predstavljaju most između relacione baze i OLAP analize, jer već na tom nivou vrše spajanje tabela i formiranje izvedenih atributa.
    """,
    'view_intro': """
    Umesto slikovnih prikaza rezultata, u ovom radu su prikazani stvarni tabelarni rezultati izvedeni iz SQL skripta ITKompanijaDW_kreiranje.sql. Prve tri tabele prikazuju reprezentativne redove iz definisanih pogleda, dok naredne tabele prikazuju rezultate provernih SELECT upita kojima se potvrđuje korektnost analitičke logike nad istim skupom podataka.
    """,
    'olap_uvod': """
    OLAP sloj predstavlja centralni deo analitičkog rešenja, jer omogućava brzo agregiranje i višedimenzionalno sagledavanje poslovnih pokazatelja. U radu su projektovane dve kocke: cbProduktivnost, zasnovana na pogledu Analiza_troska_rada, i cbPrihodi, zasnovana na pogledu Analiza_prihoda. Zajedničke dimenzije kao što su Projekat, Klijent i Vreme omogućavaju dosledno poređenje različitih vrsta mera u jedinstvenom analitičkom okruženju.

    Pri izgradnji kocki posebna pažnja posvećena je izboru mera i hijerarhija. Za produktivnost su ključne mere Broj sati, Trošak rada, Satnica i Broj evidencija, dok su za prihode relevantni Iznos, Budžet, Ukupno sati i Broj faktura. Vremenska dimenzija modelovana je tako da podrži analizu po godinama i kvartalima, čime su omogućene i sažete i detaljne vremenske interpretacije.
    """,
    'olap_mdx_intro': """
    MDX upiti predstavljaju standardni jezik za pristup OLAP kockama i u radu služe kao formalni dokaz da je model pravilno postavljen. U nastavku su izdvojeni reprezentativni upiti iz datoteke MDX_upiti_za_kocku.mdx, zajedno sa objašnjenjima njihove funkcije i tumačenjem rezultata koje daju nad kockama cbProduktivnost i cbPrihodi.
    """,
    'pivot_intro': """
    Najvažniji deo analize predstavljaju pivot tabele, jer one na najpregledniji način objedinjuju rezultate OLAP obrade i pretvaraju ih u pregled pogodan za menadžersko odlučivanje. U nastavku su prikazane sve pivot tabele formirane nad kockama produktivnosti i prihoda, uz kratko tumačenje poslovnog značenja dobijenih vrednosti.
    """,
    'zakljucak': """
    U radu je prikazan celovit postupak izgradnje skladišta podataka za domen poslovanja IT kompanije, počev od relacione baze i SQL pogleda, pa sve do projektovanja OLAP kocki, izvršavanja MDX upita i interpretacije rezultata kroz pivot tabele. Time je potvrđeno da se i na relativno kompaktnom skupu podataka može izgraditi analitičko okruženje koje pruža visok nivo preglednosti, doslednosti i poslovne upotrebljivosti.

    Analizom je utvrđeno da kocka cbProduktivnost uspešno prikazuje raspodelu radnog angažovanja po odeljenjima, projektima, zaposlenima i vremenskim periodima, dok kocka cbPrihodi omogućava jasan uvid u strukturu naplate po klijentima, projektima, godinama i kvartalima. Posebno je značajno to što su rezultati međusobno uporedivi i metodološki konzistentni, jer potiču iz jedinstveno definisanog modela i pripremnih pogleda.

    Dobijeno rešenje može predstavljati osnovu za dalje proširenje ka složenijim analizama, kao što su procena profitabilnosti po kombinaciji klijent–projekat, praćenje realizacije budžeta u realnom vremenu ili uvođenje prediktivnih modela za planiranje resursa. Na taj način skladište podataka prevazilazi ulogu tehničkog repozitorijuma i postaje aktivan instrument za strateško upravljanje poslovanjem IT kompanije.
    """,
}

MDX_EXPLANATIONS = {
    1: 'Ovaj upit sabira mere Broj sati, Trošak rada i Satnica po članovima dimenzije Odeljenje. Rezultat omogućava da se prepozna koje organizacione celine nose najveći operativni teret i gde nastaje najveći trošak ljudskog rada.',
    2: 'Upit prikazuje cbProduktivnost po projektima i time omogućava poređenje radnog angažovanja i troška između različitih poslovnih inicijativa. Na osnovu rezultata može se proceniti koji projekti najviše opterećuju raspoložive resurse kompanije.',
    3: 'Ovim upitom analiza se spušta na nivo zaposlenih i uvodi meru Broj evidencija. Dobijeni rezultat pomaže u identifikaciji ključnih resursa, eventualnih preopterećenja i kontinuiteta angažovanja pojedinačnih članova tima.',
    5: 'Peti upit iz kocke cbPrihodi prikazuje strukturu prihoda po klijentima uz mere Iznos, Ukupno sati i Broj faktura. Time se u jednom pogledu povezuju komercijalni efekat, obim angažovanja i dinamika naplate.',
    8: 'Upit po kvartalima predstavlja detaljniji vremenski pogled na prihodnu kocku. Rezultati jasno pokazuju sezonska kretanja i omogućavaju da se na nivou kvartala lakše planiraju kapaciteti, prodajne aktivnosti i očekivani novčani tokovi.',
}


def paragraphs(text: str) -> list[str]:
    return [part.strip().replace('\n', ' ') for part in text.strip().split('\n\n') if part.strip()]



def set_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic



def add_paragraph(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0.75, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(first_line) if first_line else Cm(0)
    run = p.add_run(text)
    set_font(run, size=12, bold=bold, italic=italic)
    return p



def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p



def add_minor_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p



def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=12, italic=True)
    return p



def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)



def add_table(doc: Document, state: dict, headers: list[str], rows: list[list[str]], caption: str):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell(header_cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[idx], value, align=alignment)
    add_caption(doc, f'Tabela {state["table"]}: {caption}')
    state['table'] += 1
    return table



def add_image(doc: Document, state: dict, image_path: Path, caption: str, width_inches: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_caption(doc, f'Slika {state["figure"]}: {caption}')
    state['figure'] += 1



def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code.strip())
    set_font(run, name='Courier New', size=10)
    return p



def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)



def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    for line in [
        'UNIVERZITET U NOVOM SADU',
        'Tehnički fakultet "Mihajlo Pupin" Zrenjanin',
        'Master akademske studije',
    ]:
        run = p.add_run(line + '\n')
        set_font(run, size=12, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    for line in [
        'SEMINARSKI RAD',
        'Projektovanje skladišta podataka i OLAP analize za IT kompaniju',
    ]:
        run = title.add_run(line + '\n')
        set_font(run, size=16, bold=True)
    for _ in range(8):
        doc.add_paragraph()
    for line in [
        'Student: Ognjen Grgur',
        'Broj indeksa: MIT 37/24',
        'Predmet: Koncepti baza podataka',
        'Profesor: prof. dr Dalibor Dobrilović',
    ]:
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        info.paragraph_format.line_spacing = 1.5
        run = info.add_run(line)
        set_font(run, size=12, bold=True)
    city = doc.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city.paragraph_format.space_before = Pt(42)
    city.paragraph_format.line_spacing = 1.5
    run = city.add_run('Zrenjanin, 2025.')
    set_font(run, size=12, bold=True)
    doc.add_page_break()



def add_contents(doc: Document):
    add_heading(doc, 'Sadržaj')
    items = [
        '1. Uvod',
        '2. Teorijski okvir',
        '2.1. Skladišta podataka i njihove karakteristike',
        '2.2. Dimenzionalno modeliranje',
        '2.3. OLAP pristup i višedimenzionalna analiza',
        '2.4. ETL proces i priprema analitičkih podataka',
        '3. Metodologija rada',
        '4. Implementacija skladišta podataka',
        '4.1. Konceptualni model',
        '4.2. Fizički model',
        '4.3. Realizacija baze podataka',
        '4.4. Rezultati VIEW-ova i SELECT upita',
        '5. OLAP analiza',
        '5.1. Kreiranje OLAP kocki',
        '5.2. MDX upiti',
        '5.3. Pivot analiza produktivnosti',
        '5.4. Pivot analiza prihoda',
        '6. Zaključak',
        '7. Literatura',
    ]
    for item in items:
        add_paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, space_after=0)
    doc.add_page_break()



def split_sql_tuples(text: str) -> list[str]:
    parts = []
    buf = []
    depth = 0
    in_string = False
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "'":
            buf.append(ch)
            if i + 1 < len(text) and text[i + 1] == "'":
                buf.append(text[i + 1])
                i += 1
            else:
                in_string = not in_string
        elif not in_string and ch == '(':
            if depth > 0:
                buf.append(ch)
            depth += 1
        elif not in_string and ch == ')':
            depth -= 1
            if depth == 0:
                parts.append(''.join(buf))
                buf = []
            else:
                buf.append(ch)
        elif depth > 0:
            buf.append(ch)
        i += 1
    return parts



def parse_insert_rows(sql_text: str, table_name: str) -> list[dict]:
    match = re.search(rf'INSERT INTO {re.escape(table_name)} \((.*?)\) VALUES\s*(.*?);', sql_text, re.S)
    if not match:
        raise ValueError(f'Nije pronađen INSERT blok za tabelu {table_name}.')
    columns = [column.strip() for column in match.group(1).split(',')]
    values_block = '\n'.join(line.split('--')[0] for line in match.group(2).splitlines())
    rows = []
    id_map = {
        'Odeljenje': 'idOdeljenja',
        'Zaposleni': 'idZaposlenog',
        'Klijent': 'idKlijenta',
        'Projekat': 'idProjekta',
        'Tehnologija': 'idTehnologije',
        'EvidencijaSati': 'idEvidencije',
        'Faktura': 'idFakture',
    }
    for idx, tuple_text in enumerate(split_sql_tuples(values_block), start=1):
        row = ast.literal_eval('(' + tuple_text.replace('NULL', 'None') + ')')
        if not isinstance(row, tuple):
            row = (row,)
        record = dict(zip(columns, row))
        id_column = id_map.get(table_name)
        if id_column and id_column not in record:
            record[id_column] = idx
        rows.append(record)
    return rows



def fix_diacritics(text: str) -> str:
    replacements = {
        'Petrovic': 'Petrović',
        'Stojanovic': 'Stojanović',
        'Markovic': 'Marković',
        'Nikolic': 'Nikolić',
        'Jovanovic': 'Jovanović',
        'Djordjevic': 'Đorđević',
        'Todorovic': 'Todorović',
        'Ilic': 'Ilić',
        'Pavlovic': 'Pavlović',
        'Nemacka': 'Nemačka',
        'Zavrsen': 'Završen',
        'Placena': 'Plaćena',
        'Na cekanju': 'Na čekanju',
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text



def load_source_data():
    sql_text = SQL_PATH.read_text(encoding='utf-8')
    odeljenja = {row['idOdeljenja']: row for row in parse_insert_rows(sql_text, 'Odeljenje')}
    zaposleni = {row['idZaposlenog']: row for row in parse_insert_rows(sql_text, 'Zaposleni')}
    klijenti = {row['idKlijenta']: row for row in parse_insert_rows(sql_text, 'Klijent')}
    projekti = {row['idProjekta']: row for row in parse_insert_rows(sql_text, 'Projekat')}
    tehnologije = {row['idTehnologije']: row for row in parse_insert_rows(sql_text, 'Tehnologija')}
    projekat_tehnologija = parse_insert_rows(sql_text, 'ProjekatTehnologija')
    evidencije = parse_insert_rows(sql_text, 'EvidencijaSati')
    fakture = parse_insert_rows(sql_text, 'Faktura')

    analiza_troska = []
    for evidencija in evidencije:
        zaposleni_row = zaposleni[evidencija['idZaposlenog']]
        odeljenje_row = odeljenja[zaposleni_row['idOdeljenja']]
        projekat_row = projekti[evidencija['idProjekta']]
        klijent_row = klijenti[projekat_row['idKlijenta']]
        datum = date.fromisoformat(evidencija['datum'])
        analiza_troska.append({
            'idEvidencije': evidencija['idEvidencije'],
            'imeZaposlenog': fix_diacritics(f"{zaposleni_row['ime']} {zaposleni_row['prezime']}"),
            'pozicija': zaposleni_row['pozicija'],
            'nazivOdeljenja': odeljenje_row['nazivOdeljenja'],
            'nazivProjekta': projekat_row['nazivProjekta'],
            'klijent': klijent_row['nazivKompanije'],
            'datum': evidencija['datum'],
            'brojSati': int(evidencija['brojSati']),
            'satnica': int(zaposleni_row['satnica']),
            'trosakRada': int(evidencija['brojSati'] * zaposleni_row['satnica']),
            'godina': datum.year,
            'mesec': datum.month,
            'kvartal': (datum.month - 1) // 3 + 1,
        })

    analiza_prihoda = []
    for faktura in fakture:
        projekat_row = projekti[faktura['idProjekta']]
        klijent_row = klijenti[projekat_row['idKlijenta']]
        analiza_prihoda.append({
            'idFakture': faktura['idFakture'],
            'nazivProjekta': projekat_row['nazivProjekta'],
            'budzet': int(projekat_row['budzet']),
            'statusProjekta': fix_diacritics(projekat_row['status']),
            'klijent': klijent_row['nazivKompanije'],
            'mesec': faktura['mesec'],
            'godina': faktura['godina'],
            'ukupnoSati': int(faktura['ukupnoSati']),
            'iznos': int(faktura['iznos']),
            'statusFakture': fix_diacritics(faktura['status']),
            'kvartal': (faktura['mesec'] - 1) // 3 + 1,
        })

    analiza_tehnologija = []
    for veza in projekat_tehnologija:
        projekat_row = projekti[veza['idProjekta']]
        tehnologija_row = tehnologije[veza['idTehnologije']]
        klijent_row = klijenti[projekat_row['idKlijenta']]
        analiza_tehnologija.append({
            'nazivProjekta': projekat_row['nazivProjekta'],
            'nazivTehnologije': tehnologija_row['nazivTehnologije'],
            'kategorija': tehnologija_row['kategorija'],
            'klijent': klijent_row['nazivKompanije'],
        })

    po_projektima = defaultdict(lambda: {'sati': 0, 'trosak': 0})
    po_zaposlenima = defaultdict(lambda: {'odeljenje': '', 'sati': 0})
    for evidencija in evidencije:
        projekat_row = projekti[evidencija['idProjekta']]
        zaposleni_row = zaposleni[evidencija['idZaposlenog']]
        odeljenje_row = odeljenja[zaposleni_row['idOdeljenja']]
        ime = fix_diacritics(f"{zaposleni_row['ime']} {zaposleni_row['prezime']}")
        po_projektima[projekat_row['nazivProjekta']]['sati'] += int(evidencija['brojSati'])
        po_projektima[projekat_row['nazivProjekta']]['trosak'] += int(evidencija['brojSati'] * zaposleni_row['satnica'])
        po_zaposlenima[ime]['odeljenje'] = odeljenje_row['nazivOdeljenja']
        po_zaposlenima[ime]['sati'] += int(evidencija['brojSati'])

    prihod_po_klijentima = defaultdict(int)
    prihod_po_kvartalima = defaultdict(int)
    for faktura in fakture:
        projekat_row = projekti[faktura['idProjekta']]
        klijent_row = klijenti[projekat_row['idKlijenta']]
        prihod_po_klijentima[klijent_row['nazivKompanije']] += int(faktura['iznos'])
        kvartal = f"Q{(faktura['mesec'] - 1) // 3 + 1}"
        prihod_po_kvartalima[(faktura['godina'], kvartal)] += int(faktura['iznos'])

    view1_rows = []
    for projekat_naziv in ['E-Commerce platforma', 'CRM sistem', 'Mobile Banking App', 'Cloud Migration', 'Analytics Dashboard']:
        for row in analiza_troska:
            if row['nazivProjekta'] == projekat_naziv:
                view1_rows.append([
                    row['idEvidencije'], row['imeZaposlenog'], row['nazivOdeljenja'], row['nazivProjekta'], row['brojSati'], row['trosakRada'], row['godina'], row['kvartal']
                ])
                break
    view1_rows.append([analiza_troska[5]['idEvidencije'], analiza_troska[5]['imeZaposlenog'], analiza_troska[5]['nazivOdeljenja'], analiza_troska[5]['nazivProjekta'], analiza_troska[5]['brojSati'], analiza_troska[5]['trosakRada'], analiza_troska[5]['godina'], analiza_troska[5]['kvartal']])

    selected_projects = ['E-Commerce platforma', 'CRM sistem', 'Mobile Banking App', 'Cloud Migration', 'Analytics Dashboard']
    view2_rows = []
    for projekat_naziv in selected_projects:
        for row in analiza_prihoda:
            if row['nazivProjekta'] == projekat_naziv:
                view2_rows.append([
                    row['idFakture'], row['nazivProjekta'], row['klijent'], row['mesec'], row['godina'], row['ukupnoSati'], row['iznos'], row['statusFakture']
                ])
                break
    view2_rows.append([
        analiza_prihoda[-1]['idFakture'], analiza_prihoda[-1]['nazivProjekta'], analiza_prihoda[-1]['klijent'], analiza_prihoda[-1]['mesec'], analiza_prihoda[-1]['godina'], analiza_prihoda[-1]['ukupnoSati'], analiza_prihoda[-1]['iznos'], analiza_prihoda[-1]['statusFakture']
    ])

    view3_rows = [[row['nazivProjekta'], row['nazivTehnologije'], row['kategorija'], row['klijent']] for row in analiza_tehnologija[:10]]
    select1_rows = [[name, values['sati'], values['trosak']] for name, values in sorted(po_projektima.items(), key=lambda item: item[1]['trosak'], reverse=True)]
    select2_rows = [[name, values['odeljenje'], values['sati']] for name, values in sorted(po_zaposlenima.items(), key=lambda item: item[1]['sati'], reverse=True)]
    select3_rows = [[name, value] for name, value in sorted(prihod_po_klijentima.items(), key=lambda item: item[1], reverse=True)]
    select4_rows = [[godina, kvartal, iznos] for (godina, kvartal), iznos in sorted(prihod_po_kvartalima.items())]

    return {
        'view1_rows': view1_rows,
        'view2_rows': view2_rows,
        'view3_rows': view3_rows,
        'select1_rows': select1_rows,
        'select2_rows': select2_rows,
        'select3_rows': select3_rows,
        'select4_rows': select4_rows,
    }



def load_mdx_queries() -> list[dict]:
    text = MDX_PATH.read_text(encoding='utf-8')
    selected = {1, 2, 3, 5, 8}
    queries = []
    current = None
    for line in text.splitlines():
        match = re.match(r'-- Upit (\d+):\s*(.+)', line)
        if match:
            if current and current['number'] in selected:
                current['query'] = '\n'.join(item for item in current['lines'] if item.strip()).strip()
                queries.append(current)
            current = {'number': int(match.group(1)), 'title': match.group(2).strip(), 'lines': []}
            continue
        if current is not None and not line.startswith('-- ==========================================') and not line.startswith('-- KOCKA'):
            current['lines'].append(line)
    if current and current['number'] in selected:
        current['query'] = '\n'.join(item for item in current['lines'] if item.strip()).strip()
        queries.append(current)
    return queries



def build_document():
    for image in IMAGES.values():
        if not image.exists():
            raise FileNotFoundError(f'Nedostaje slika: {image}')

    source_data = load_source_data()
    mdx_queries = load_mdx_queries()
    doc = Document()
    configure_document(doc)
    state = {'table': 1, 'figure': 1}

    add_title_page(doc)
    add_contents(doc)

    add_heading(doc, '1. Uvod')
    for text in paragraphs(SECTION_TEXT['uvod']):
        add_paragraph(doc, text)

    add_heading(doc, '2. Teorijski okvir')
    add_heading(doc, '2.1. Skladišta podataka i njihove karakteristike')
    for text in paragraphs(SECTION_TEXT['teorija1']):
        add_paragraph(doc, text)
    add_heading(doc, '2.2. Dimenzionalno modeliranje')
    for text in paragraphs(SECTION_TEXT['teorija2']):
        add_paragraph(doc, text)
    add_heading(doc, '2.3. OLAP pristup i višedimenzionalna analiza')
    for text in paragraphs(SECTION_TEXT['teorija3']):
        add_paragraph(doc, text)
    add_heading(doc, '2.4. ETL proces i priprema analitičkih podataka')
    for text in paragraphs(SECTION_TEXT['teorija4']):
        add_paragraph(doc, text)

    add_heading(doc, '3. Metodologija rada')
    for text in paragraphs(SECTION_TEXT['metodologija']):
        add_paragraph(doc, text)

    add_heading(doc, '4. Implementacija skladišta podataka')
    for text in paragraphs(SECTION_TEXT['implementacija_uvod']):
        add_paragraph(doc, text)

    add_heading(doc, '4.1. Konceptualni model')
    add_image(doc, state, IMAGES['cdm'], 'Konceptualni model skladišta podataka za IT kompaniju', width_inches=6.0)
    add_paragraph(doc, 'Konceptualni model prikazuje ključne poslovne entitete i njihove međusobne veze na apstraktnom nivou. U središtu modela nalaze se projekti kao veza između klijenata, zaposlenih, tehnologija i finansijskih tokova, čime se jasno ističe poslovna logika budućeg analitičkog sistema.')
    add_paragraph(doc, 'Na konceptualnom nivou već se uočava pogodnost domena za analizu kroz dimenzije Vreme, Projekat, Klijent i Zaposleni. Ovakav raspored entiteta olakšava prelazak ka dimenzionalnom modelu i potvrđuje da domen prirodno podržava višedimenzionalnu obradu podataka.')

    add_heading(doc, '4.2. Fizički model')
    add_image(doc, state, IMAGES['pdm'], 'Fizički model relacione baze podataka', width_inches=6.0)
    add_paragraph(doc, 'Fizički model prevodi konceptualne entitete u konkretne tabele, primarne i strane ključeve, kao i tipove podataka potrebne za realizaciju baze. Na tom nivou posebno dolazi do izražaja normalizacija operativnog sloja, kojom se obezbeđuju integritet i konzistentnost unosa.')
    add_paragraph(doc, 'Za potrebe kasnije analitike važna je činjenica da fizički model zadržava jasne veze između zaposlenih, projekata, klijenata i faktura. Upravo te veze omogućavaju da se u pripremnim pogledima izvrši spajanje tabela bez gubitka poslovnog značenja podataka.')

    add_heading(doc, '4.3. Realizacija baze podataka')
    add_image(doc, state, IMAGES['db'], 'Dijagram realizovane baze podataka', width_inches=6.1)
    add_paragraph(doc, 'Realizovana baza obuhvata osam centralnih tabela: Odeljenje, Zaposleni, Klijent, Projekat, Tehnologija, ProjekatTehnologija, EvidencijaSati i Faktura. Takav skup tabela pokriva i organizacioni i komercijalni aspekt poslovanja, što predstavlja dobru osnovu za analizu radne produktivnosti i prihoda.')
    add_paragraph(doc, 'SQL skript ne sadrži samo strukturu već i reprezentativne test podatke za 2024. i 2025. godinu. To je posebno važno, jer omogućava proveru analitičke logike kroz stvarne agregacije i olakšava kasniju validaciju OLAP kocki i pivot tabela.')

    add_heading(doc, '4.4. Rezultati VIEW-ova i SELECT upita')
    for text in paragraphs(SECTION_TEXT['view_intro']):
        add_paragraph(doc, text)

    add_table(doc, state, ['ID evidencije', 'Zaposleni', 'Odeljenje', 'Projekat', 'Broj sati', 'Trošak rada', 'Godina', 'Kvartal'], source_data['view1_rows'], 'Reprezentativni rezultati pogleda Analiza_troska_rada')
    add_paragraph(doc, 'Pogled Analiza_troska_rada objedinjuje operativne podatke o radu zaposlenih sa organizacionim i projektnim dimenzijama, pri čemu automatski izračunava trošak rada i vremenske atribute. Time ovaj pogled postaje prirodna osnova za formiranje kocke produktivnosti.')

    add_table(doc, state, ['ID fakture', 'Projekat', 'Klijent', 'Mesec', 'Godina', 'Ukupno sati', 'Iznos', 'Status'], source_data['view2_rows'], 'Reprezentativni rezultati pogleda Analiza_prihoda')
    add_paragraph(doc, 'Pogled Analiza_prihoda povezuje fakture sa projektima i klijentima i omogućava da se prihodna analiza vodi kroz jedinstvenu tabelu činjenica. Uvedeni kvartal i status fakture dodatno proširuju mogućnost vremenskog i poslovnog preseka podataka.')

    add_table(doc, state, ['Projekat', 'Tehnologija', 'Kategorija', 'Klijent'], source_data['view3_rows'], 'Reprezentativni rezultati pogleda Analiza_tehnologija')
    add_paragraph(doc, 'Pogled Analiza_tehnologija nije direktna fakt tabela za kocku, ali ima značajnu interpretativnu vrednost. On omogućava da se projekti povežu sa tehnološkim portfoliom kompanije i predstavlja važan kontekst za razumevanje složenosti angažovanja.')

    add_table(doc, state, ['Projekat', 'Ukupno sati', 'Ukupan trošak'], source_data['select1_rows'], 'Rezultati SELECT upita za ukupno sate i trošak po projektima')
    add_paragraph(doc, 'Prvi proverni upit potvrđuje da je najveći trošak rada koncentrisan na projektima E-Commerce platforma i CRM sistem. Dobijeni rezultati u potpunosti su usaglašeni sa kasnijim pivot analizama produktivnosti po projektima.')

    add_table(doc, state, ['Zaposleni', 'Odeljenje', 'Ukupno sati'], source_data['select2_rows'], 'Rezultati SELECT upita za produktivnost po zaposlenima')
    add_paragraph(doc, 'Drugi upit pokazuje raspodelu ukupnih sati po zaposlenima i odeljenjima. Ovaj pregled je značajan kao relacijska potvrda podataka koji će kasnije biti posmatrani u OLAP kocki kroz detaljniju meru Broj evidencija.')

    add_table(doc, state, ['Klijent', 'Ukupan prihod'], source_data['select3_rows'], 'Rezultati SELECT upita za prihod po klijentima')
    add_paragraph(doc, 'Treći proverni upit potvrđuje koncentraciju prihoda kod nekoliko ključnih klijenata. Time se već na relacijskom nivou vidi potreba za daljim OLAP presekom koji uključuje dodatne mere, kao što su broj faktura i utrošeni sati.')

    add_table(doc, state, ['Godina', 'Kvartal', 'Ukupan prihod'], source_data['select4_rows'], 'Rezultati SELECT upita za prihod po kvartalima')
    add_paragraph(doc, 'Kvartalni prihod potvrđuje da je najveća naplata koncentrisana u prvom i drugom kvartalu. Ovaj rezultat predstavlja dobru pripremu za kasniju interpretaciju vremenskih pivot tabela u okviru kocke prihoda.')

    doc.add_page_break()
    add_heading(doc, '5. OLAP analiza')
    add_heading(doc, '5.1. Kreiranje OLAP kocki')
    for text in paragraphs(SECTION_TEXT['olap_uvod']):
        add_paragraph(doc, text)
    add_paragraph(doc, 'Kocka cbProduktivnost zasniva se na evidenciji rada i omogućava analizu po odeljenjima, zaposlenima, projektima i vremenu. Kocka cbPrihodi koristi fakture kao izvor činjenica i podržava analizu po klijentima, projektima i vremenskim hijerarhijama. Odvajanje ove dve celine doprinosi metodološkoj jasnoći, jer se radni i finansijski pokazatelji mogu analizirati odvojeno, a zatim upoređivati na nivou zaključaka.')

    add_heading(doc, '5.2. MDX upiti')
    for text in paragraphs(SECTION_TEXT['olap_mdx_intro']):
        add_paragraph(doc, text)
    for query in mdx_queries:
        add_minor_heading(doc, f"Upit {query['number']}: {query['title']}")
        add_code_block(doc, query['query'])
        add_paragraph(doc, MDX_EXPLANATIONS[query['number']])

    add_heading(doc, '5.3. Pivot analiza produktivnosti')
    for text in paragraphs(SECTION_TEXT['pivot_intro']):
        add_paragraph(doc, text)
    for pivot in PIVOT_TABLES[:4]:
        add_table(doc, state, pivot['headers'], pivot['rows'], pivot['title'])
        add_paragraph(doc, pivot['analysis'])

    add_heading(doc, '5.4. Pivot analiza prihoda')
    for pivot in PIVOT_TABLES[4:]:
        add_table(doc, state, pivot['headers'], pivot['rows'], pivot['title'])
        add_paragraph(doc, pivot['analysis'])
    add_paragraph(doc, 'Zajedničko posmatranje pivot tabela produktivnosti i prihoda pokazuje da OLAP pristup omogućava veoma brzo povezivanje operativnih i finansijskih pokazatelja. Na taj način rukovodstvo ne dobija samo izolovane izveštaje, već konzistentan okvir za procenu odnosa između uloženog rada, troškova i ostvarenih prihoda.')

    add_heading(doc, '6. Zaključak')
    for text in paragraphs(SECTION_TEXT['zakljucak']):
        add_paragraph(doc, text)

    add_heading(doc, '7. Literatura')
    for index, item in enumerate(LITERATURE, start=1):
        add_paragraph(doc, f'[{index}] {item}', align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, space_after=3)

    doc.save(OUTPUT_PATH)


if __name__ == '__main__':
    build_document()
    print(f'Dokument je uspešno generisan: {OUTPUT_PATH}')
